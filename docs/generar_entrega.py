# -*- coding: utf-8 -*-
"""Genera Manual de Despliegue, Manual de Usuario y ZIPs de entrega."""
from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
ENTREGA = Path(__file__).resolve().parent / "entrega"
CODIGO = ENTREGA / "Codigo_Comprimido"
MANUALES = ENTREGA / "Manuales"
DESPLIEGUE_DIR = MANUALES / "Despliegue"
USUARIO_DIR = MANUALES / "Guia_de_Usuario"
REPO = "https://github.com/Jmmsimon/trazabilidad-taller-I"


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)


def P(doc, text, bold=False, italic=False, size=11, align="justify"):
    p = doc.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)


def B(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        set_run_font(r, size=11)


def shade(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "1E3A5F")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            set_run_font(run, size=9, bold=True)


def T(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
        shade(table.rows[0].cells[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def cover(doc, title, subtitle):
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "UNIVERSIDAD PRIVADA ANTENOR ORREGO", bold=True, size=14, align="center")
    P(doc, "Facultad de Ingeniería", bold=True, size=12, align="center")
    P(
        doc,
        "Escuela Profesional de Ingeniería de Computación y Sistemas",
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, title, bold=True, size=16, align="center")
    P(doc, subtitle, bold=True, size=12, align="center")
    doc.add_paragraph()
    P(
        doc,
        "Plataforma de Trazabilidad Académica con Agentes de IA Generativa",
        italic=True,
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "Autor: Jean Marcos Meneses Simón", size=11, align="center")
    P(doc, "Asesor: Ing. Walter Cueva", size=11, align="center")
    P(doc, "Fecha: Julio 2026", size=11, align="center")
    P(doc, f"Repositorio: {REPO}", size=10, align="center")
    doc.add_page_break()


def build_despliegue():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)

    cover(
        doc,
        "MANUAL DE DESPLIEGUE",
        "Instalación, configuración y puesta en marcha",
    )

    H(doc, "1. Propósito")
    P(
        doc,
        "Este manual describe los requisitos, la configuración de variables de entorno "
        "y los pasos para desplegar la plataforma de trazabilidad académica en un "
        "entorno local o de producción (frontend Next.js + backend FastAPI).",
    )

    H(doc, "2. Arquitectura de despliegue")
    B(
        doc,
        [
            "Frontend: Next.js 16 (App Router) en el puerto 3000 (desarrollo) o 3000/443 (producción).",
            "Backend: FastAPI + Uvicorn en el puerto 8000.",
            "Servicios externos: Firebase Authentication/Firestore, GitHub API, LLM (Google Gemini o Anthropic Claude).",
            "El frontend se comunica con el backend mediante API Routes (variable BACKEND_URL).",
        ],
    )

    H(doc, "3. Requisitos previos")
    T(
        doc,
        ["Componente", "Versión recomendada"],
        [
            ["Node.js", "20 LTS o superior"],
            ["npm", "incluido con Node.js"],
            ["Python", "3.12"],
            ["Git", "2.x"],
            ["Cuenta Firebase", "proyecto con Auth + Firestore"],
            ["API Key LLM", "GOOGLE_API_KEY o ANTHROPIC_API_KEY"],
            ["Token GitHub (opcional)", "para lectura ampliada de repos privados"],
        ],
    )

    H(doc, "4. Estructura de entrega")
    B(
        doc,
        [
            "Backend.zip — código del servicio FastAPI (carpeta backend/).",
            "Frontend.zip — código de la aplicación Next.js (raíz del front).",
            "Repositorio Git: " + REPO,
        ],
    )

    H(doc, "5. Despliegue local — Backend")
    H(doc, "5.1 Extracción e instalación", 2)
    P(doc, "Desde una terminal en la carpeta del backend:")
    P(
        doc,
        "python -m venv .venv\n"
        "# Windows (Git Bash / PowerShell):\n"
        "source .venv/Scripts/activate   # o .venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt",
        size=9,
    )
    H(doc, "5.2 Credenciales Firebase (Admin SDK)", 2)
    P(
        doc,
        "Colocar el archivo serviceAccountKey.json dentro de la carpeta backend/ "
        "(descargado desde Firebase Console → Configuración del proyecto → Cuentas de servicio). "
        "En producción puede usarse la variable de entorno FIREBASE_CREDENTIALS con el JSON completo.",
    )
    H(doc, "5.3 Variables de entorno del backend", 2)
    P(
        doc,
        "El backend carga automáticamente el archivo .env.local ubicado en la raíz del frontend "
        "(un nivel arriba de backend/) o un archivo .env. Variables mínimas:",
    )
    T(
        doc,
        ["Variable", "Descripción"],
        [
            ["GOOGLE_API_KEY", "Clave de Google AI (Gemini). Prioritaria."],
            ["ANTHROPIC_API_KEY", "Alternativa si no hay Gemini."],
            ["FIREBASE_CREDENTIALS", "JSON del service account (producción)."],
        ],
    )
    H(doc, "5.4 Ejecución", 2)
    P(doc, "uvicorn main:app --reload --host 0.0.0.0 --port 8000", size=9)
    P(
        doc,
        "Verificación: abrir http://localhost:8000/docs (documentación OpenAPI de FastAPI).",
    )

    H(doc, "6. Despliegue local — Frontend")
    H(doc, "6.1 Extracción e instalación", 2)
    P(doc, "npm install", size=9)
    H(doc, "6.2 Archivo .env.local", 2)
    P(doc, "Crear .env.local en la raíz del frontend con:")
    P(
        doc,
        "NEXT_PUBLIC_FIREBASE_API_KEY=...\n"
        "NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN=...\n"
        "NEXT_PUBLIC_FIREBASE_PROJECT_ID=...\n"
        "NEXT_PUBLIC_FIREBASE_DATABASE_URL=...\n"
        "NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET=...\n"
        "NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID=...\n"
        "NEXT_PUBLIC_FIREBASE_APP_ID=...\n"
        "NEXT_PUBLIC_FIREBASE_MEASUREMENT_ID=...\n"
        "BACKEND_URL=http://localhost:8000\n"
        "GOOGLE_API_KEY=...",
        size=9,
    )
    H(doc, "6.3 Nota sobre middleware", 2)
    P(
        doc,
        "El archivo middleware.ts permanece deshabilitado (middleware.ts.disabled) para evitar "
        "un conflicto de enrutamiento con Next.js 16 que provocaba respuestas 404 en rutas "
        "protegidas y API. La autenticación se gestiona mediante API Routes y cookies httpOnly.",
    )
    H(doc, "6.4 Ejecución", 2)
    P(doc, "npm run dev", size=9)
    P(doc, "Aplicación: http://localhost:3000")

    H(doc, "7. Orden de arranque recomendado")
    B(
        doc,
        [
            "1) Iniciar backend en :8000 y verificar /docs.",
            "2) Iniciar frontend en :3000.",
            "3) Iniciar sesión con un usuario Firebase (roles: estudiante, docente, administrador).",
        ],
    )

    H(doc, "8. Despliegue en producción (referencia)")
    H(doc, "8.1 Backend (Railway / Render / VPS)", 2)
    B(
        doc,
        [
            "Build: pip install -r requirements.txt",
            "Start: uvicorn main:app --host 0.0.0.0 --port $PORT",
            "Configurar FIREBASE_CREDENTIALS, GOOGLE_API_KEY (o ANTHROPIC_API_KEY) y CORS según origen del front.",
            "El repositorio incluye Procfile de referencia para plataformas PaaS.",
        ],
    )
    H(doc, "8.2 Frontend (Vercel u otro host Node)", 2)
    B(
        doc,
        [
            "Build: npm run build",
            "Start: npm run start",
            "Definir todas las variables NEXT_PUBLIC_FIREBASE_* y BACKEND_URL apuntando al backend público (HTTPS).",
            "Habilitar cookie secure en producción (NODE_ENV=production).",
        ],
    )

    H(doc, "9. Verificación post-despliegue")
    T(
        doc,
        ["Prueba", "Resultado esperado"],
        [
            ["GET backend /docs", "Swagger UI visible"],
            ["Login en /", "Redirección según rol"],
            ["Estudiante → Analizar", "Tracking completed + Kanban actualizado"],
            ["Docente → Analítica", "Score, commits y competencias visibles"],
            ["Docente → Auditoría CSV", "Semáforo solo-reporte (no altera Kanban)"],
        ],
    )

    H(doc, "10. Solución de problemas frecuentes")
    T(
        doc,
        ["Síntoma", "Causa probable", "Acción"],
        [
            ["404 en /dashboard o /api/*", "middleware conflictivo", "Mantener middleware.ts.disabled"],
            ["Agentes en modo simulación", "Falta GOOGLE_API_KEY", "Configurar clave Gemini"],
            ["Error Firebase Admin", "Falta serviceAccountKey.json", "Colocar JSON o FIREBASE_CREDENTIALS"],
            ["Front no habla con API", "BACKEND_URL incorrecta", "Revisar .env.local"],
            ["Repo GitHub vacío en análisis", "URL inválida o rate limit", "Verificar repo público / token"],
        ],
    )

    H(doc, "11. Seguridad")
    B(
        doc,
        [
            "No versionar .env.local ni serviceAccountKey.json.",
            "Los ZIP de entrega no incluyen secretos ni node_modules / .venv.",
            "Rotar claves si se compartieron accidentalmente.",
        ],
    )

    out = DESPLIEGUE_DIR / "Manual_Despliegue.docx"
    doc.save(out)
    return out


def build_usuario():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)

    cover(
        doc,
        "MANUAL DE USUARIO",
        "Guía de uso por roles (Estudiante, Docente, Administrador)",
    )

    H(doc, "1. Introducción")
    P(
        doc,
        "La plataforma de trazabilidad académica permite vincular el avance declarado "
        "(backlog e hitos) con evidencia real del repositorio GitHub, mediante agentes "
        "de IA generativa. Este manual describe las operaciones cotidianas de cada rol.",
    )

    H(doc, "2. Acceso al sistema")
    B(
        doc,
        [
            "Abrir la URL de la aplicación (local: http://localhost:3000).",
            "Iniciar sesión con el correo y contraseña provisionados por el administrador (Firebase Auth).",
            "El sistema redirige automáticamente al panel del rol asignado: Estudiante, Docente o Administrador.",
            "Si se solicita cambio de contraseña obligatorio, completar el formulario antes de continuar.",
        ],
    )

    H(doc, "3. Rol Estudiante")
    H(doc, "3.1 Panel principal", 2)
    P(
        doc,
        "Desde el dashboard del estudiante se gestiona el ciclo de vida del proyecto académico: "
        "descubrimiento de roadmap, configuración de entregables, tablero Kanban, análisis de "
        "trazabilidad y analítica de avance.",
    )
    H(doc, "3.2 Discovery (co-creación de roadmap)", 2)
    B(
        doc,
        [
            "Iniciar un nuevo proyecto e indicar tema/descripción.",
            "El sistema ejecuta el grafo Discovery (agentes drafter → validator → product owner).",
            "Revisar la propuesta de hitos y backlog generada.",
            "Aceptar o solicitar ajustes según la interfaz disponible.",
        ],
    )
    H(doc, "3.3 Configurar repositorio y demo", 2)
    B(
        doc,
        [
            "Ingresar la URL del repositorio GitHub del proyecto.",
            "Ingresar, si aplica, la URL de la demo desplegada (p. ej. Vercel).",
            "Guardar la configuración antes de ejecutar el análisis.",
        ],
    )
    H(doc, "3.4 Analizar trazabilidad", 2)
    B(
        doc,
        [
            "Pulsar el botón Analizar.",
            "Esperar a que el estado del Tracking pase a completed (el sistema consulta el estado en segundo plano).",
            "Revisar el Kanban: los ítems con evidencia de código pueden actualizarse automáticamente a Hecho (AUTO-KANBAN).",
            "Consultar la analítica: score de integridad, commits analizados y porcentaje de competencias.",
        ],
    )
    H(doc, "3.5 Kanban y hitos", 2)
    B(
        doc,
        [
            "El tablero muestra estados: backlog, por hacer, en progreso y hecho.",
            "Los hitos pueden consultarse en modo lectura; la edición se habilita bajo demanda cuando el flujo lo permite.",
            "No modificar manualmente ítems solo para “aparentar” avance: el docente puede auditar con evidencia Git.",
        ],
    )
    H(doc, "3.6 Exportaciones", 2)
    P(
        doc,
        "Cuando esté disponible en la interfaz, exportar el reporte del proyecto en PDF o JSON "
        "para respaldo o entrega académica.",
    )

    H(doc, "4. Rol Docente")
    H(doc, "4.1 Listado de proyectos", 2)
    P(
        doc,
        "El panel docente muestra los proyectos asignados o visibles según la configuración del curso. "
        "Seleccionar un proyecto para revisar hitos, Git, analítica y auditoría.",
    )
    H(doc, "4.2 Aprobar o rechazar roadmap", 2)
    B(
        doc,
        [
            "Revisar la propuesta generada en Discovery.",
            "Aprobar para habilitar el seguimiento, o rechazar con observaciones para que el estudiante itere.",
        ],
    )
    H(doc, "4.3 Analizar / abrir Analítica", 2)
    B(
        doc,
        [
            "Ejecutar Analizar para lanzar el mismo pipeline de Tracking usado por el estudiante (fuente única de verdad).",
            "Abrir Analítica para ver score de integridad, commits, resumen de competencias y estado del repositorio.",
            "Contrastar el Kanban con la evidencia reportada (ítems Hecho vs commits/archivos).",
        ],
    )
    H(doc, "4.4 Validar hitos y tareas", 2)
    P(
        doc,
        "Desde la vista de hitos, el docente puede validar o solicitar correcciones sobre tareas específicas, "
        "dejando trazabilidad del feedback académico.",
    )
    H(doc, "4.5 Auditoría de avances (CSV)", 2)
    B(
        doc,
        [
            "Cargar un archivo CSV (p. ej. exportado desde Notion u hoja de backlog).",
            "Ejecutar la auditoría: el sistema compara ítems declarados contra evidencia en GitHub y genera un semáforo.",
            "Importante: la auditoría es solo reporte; no modifica el Kanban del estudiante.",
            "Usar el resultado como apoyo a la evaluación, no como sustituto del criterio docente.",
        ],
    )

    H(doc, "5. Rol Administrador")
    B(
        doc,
        [
            "Gestionar usuarios: listar, invitar y actualizar estado de cuenta.",
            "Asignar o cambiar roles (estudiante, docente, administrador).",
            "Supervisar que las cuentas del curso estén activas antes del inicio de clases.",
        ],
    )

    H(doc, "6. Buenas prácticas")
    B(
        doc,
        [
            "Mantener el repositorio GitHub actualizado con commits significativos.",
            "Vincular correctamente repo_url y demo_url antes de Analizar.",
            "Interpretar el score de integridad junto con el contexto del proyecto (no como nota automática final).",
            "Ante resultados inesperados, re-ejecutar Analizar y verificar conectividad del backend.",
        ],
    )

    H(doc, "7. Preguntas frecuentes")
    T(
        doc,
        ["Pregunta", "Respuesta"],
        [
            [
                "¿Por qué el análisis tarda?",
                "Lee el repositorio (deep) y consulta el LLM; puede tomar varios minutos.",
            ],
            [
                "¿La auditoría CSV cambia mi Kanban?",
                "No. Es un informe para el docente.",
            ],
            [
                "¿Qué significa score 100?",
                "Alta consistencia entre avance declarado y evidencia detectada en el piloto; no reemplaza la calificación humana.",
            ],
            [
                "No veo mis proyectos",
                "Verificar rol y que el administrador haya vinculado la cuenta al curso/proyecto.",
            ],
        ],
    )

    H(doc, "8. Soporte")
    P(
        doc,
        "Para incidencias técnicas del piloto académico, contactar al autor del sistema o al docente del curso. "
        f"Código fuente: {REPO}.",
    )

    out = USUARIO_DIR / "Manual_Usuario.docx"
    doc.save(out)
    return out


BACKEND_EXCLUDE_NAMES = {
    ".venv",
    "__pycache__",
    "serviceAccountKey.json",
    ".env",
    ".env.local",
}
BACKEND_EXCLUDE_SUFFIXES = {".pyc", ".pyo"}

FRONTEND_EXCLUDE_DIRS = {
    "node_modules",
    ".next",
    ".git",
    ".venv",
    "backend",
    "docs",
    "student-repo-test",
    "langgraph-test",
    ".vscode",
}
FRONTEND_EXCLUDE_FILES = {
    ".env.local",
    ".env",
    "clean_diff.txt",
    "diff.txt",
    "diff_utf8.txt",
    "page_history.txt",
    "reconstructed_page.tsx",
    "reconstructed_section.txt",
    "tsconfig.tsbuildinfo",
}


def write_env_examples():
    front_example = CODIGO / "_plantillas"
    front_example.mkdir(parents=True, exist_ok=True)
    (front_example / "env.local.example").write_text(
        "\n".join(
            [
                "NEXT_PUBLIC_FIREBASE_API_KEY=",
                "NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN=",
                "NEXT_PUBLIC_FIREBASE_PROJECT_ID=",
                "NEXT_PUBLIC_FIREBASE_DATABASE_URL=",
                "NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET=",
                "NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID=",
                "NEXT_PUBLIC_FIREBASE_APP_ID=",
                "NEXT_PUBLIC_FIREBASE_MEASUREMENT_ID=",
                "BACKEND_URL=http://localhost:8000",
                "GOOGLE_API_KEY=",
                "# ANTHROPIC_API_KEY=",
                "",
            ]
        ),
        encoding="utf-8",
    )
    (front_example / "LEEME_SECRETOS.txt").write_text(
        "Los ZIP no incluyen .env.local ni serviceAccountKey.json por seguridad.\n"
        "Copiar env.local.example a .env.local y completar valores.\n"
        "Colocar serviceAccountKey.json dentro de backend/ (o usar FIREBASE_CREDENTIALS).\n",
        encoding="utf-8",
    )


def zip_backend(dest: Path) -> Path:
    backend = ROOT / "backend"
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in backend.rglob("*"):
            if not path.is_file():
                continue
            rel_parts = path.relative_to(backend).parts
            if any(p in BACKEND_EXCLUDE_NAMES for p in rel_parts):
                continue
            if path.name in BACKEND_EXCLUDE_NAMES:
                continue
            if path.suffix in BACKEND_EXCLUDE_SUFFIXES:
                continue
            arc = Path("backend") / path.relative_to(backend)
            zf.write(path, arc.as_posix())
        # plantilla de entorno
        zf.writestr(
            "backend/.env.example",
            "GOOGLE_API_KEY=\nANTHROPIC_API_KEY=\n# FIREBASE_CREDENTIALS={...json...}\n",
        )
        zf.writestr(
            "backend/README_ENTREGA.txt",
            "1) python -m venv .venv && activar\n"
            "2) pip install -r requirements.txt\n"
            "3) Colocar serviceAccountKey.json aquí\n"
            "4) Configurar GOOGLE_API_KEY (o en ../.env.local)\n"
            "5) uvicorn main:app --reload --port 8000\n",
        )
    return dest


def zip_frontend(dest: Path) -> Path:
    include_roots = [
        "app",
        "components",
        "lib",
        "public",
        "package.json",
        "package-lock.json",
        "tsconfig.json",
        "next.config.ts",
        "next-env.d.ts",
        "postcss.config.mjs",
        "eslint.config.mjs",
        "middleware.ts.disabled",
        "AGENTS.md",
        "CLAUDE.md",
        "README.md",
        "langgraph.json",
        "Procfile",
        "requirements.txt",
        "prisma",
    ]
    with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
        for name in include_roots:
            path = ROOT / name
            if not path.exists():
                continue
            if path.is_file():
                if path.name in FRONTEND_EXCLUDE_FILES:
                    continue
                zf.write(path, Path("frontend") / path.name)
                continue
            for f in path.rglob("*"):
                if not f.is_file():
                    continue
                rel = f.relative_to(ROOT)
                if any(part in FRONTEND_EXCLUDE_DIRS for part in rel.parts):
                    continue
                if f.name in FRONTEND_EXCLUDE_FILES:
                    continue
                zf.write(f, (Path("frontend") / rel).as_posix())
        zf.writestr(
            "frontend/.env.local.example",
            "\n".join(
                [
                    "NEXT_PUBLIC_FIREBASE_API_KEY=",
                    "NEXT_PUBLIC_FIREBASE_AUTH_DOMAIN=",
                    "NEXT_PUBLIC_FIREBASE_PROJECT_ID=",
                    "NEXT_PUBLIC_FIREBASE_DATABASE_URL=",
                    "NEXT_PUBLIC_FIREBASE_STORAGE_BUCKET=",
                    "NEXT_PUBLIC_FIREBASE_MESSAGING_SENDER_ID=",
                    "NEXT_PUBLIC_FIREBASE_APP_ID=",
                    "NEXT_PUBLIC_FIREBASE_MEASUREMENT_ID=",
                    "BACKEND_URL=http://localhost:8000",
                    "GOOGLE_API_KEY=",
                    "",
                ]
            ),
        )
        zf.writestr(
            "frontend/README_ENTREGA.txt",
            "1) npm install\n"
            "2) Copiar .env.local.example → .env.local y completar\n"
            "3) Asegurar backend en :8000\n"
            "4) npm run dev → http://localhost:3000\n"
            "Nota: middleware.ts.disabled es intencional (ADR enrutamiento Next 16).\n",
        )
    return dest


def main():
    for d in (ENTREGA, CODIGO, DESPLIEGUE_DIR, USUARIO_DIR):
        d.mkdir(parents=True, exist_ok=True)

    write_env_examples()
    d1 = build_despliegue()
    d2 = build_usuario()
    z1 = zip_backend(CODIGO / "Backend.zip")
    z2 = zip_frontend(CODIGO / "Frontend.zip")

    (ENTREGA / "README_ENTREGA.txt").write_text(
        f"""ENTREGA — Plataforma de Trazabilidad Académica
Repositorio: {REPO}

Estructura:
  Codigo_Comprimido/
    Backend.zip
    Frontend.zip
    _plantillas/          (ejemplos de variables, sin secretos)
  Manuales/
    Despliegue/Manual_Despliegue.docx
    Guia_de_Usuario/Manual_Usuario.docx

Importante:
- Los ZIP NO incluyen node_modules, .venv, .env.local ni serviceAccountKey.json.
- Configurar secretos localmente según los manuales.
""",
        encoding="utf-8",
    )

    print("OK", d1)
    print("OK", d2)
    print("OK", z1, "size", z1.stat().st_size)
    print("OK", z2, "size", z2.stat().st_size)
    print("DIR", ENTREGA)


if __name__ == "__main__":
    main()
