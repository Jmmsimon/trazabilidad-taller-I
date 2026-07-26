# -*- coding: utf-8 -*-
"""Regenera TDDR con refs DOI reales, métricas piloto y anexos claros."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "TDDR_Trazabilidad_Academica.docx"
MERMAID = "\n".join(
    line
    for line in (ROOT / "arquitectura_drawio.mmd").read_text(encoding="utf-8").splitlines()
    if not line.strip().startswith("%%")
).strip()


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)


def P(doc, text, bold=False, italic=False, size=11, align="justify"):
    p = doc.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(6)


def B(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(it), size=11)


def shade(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "1E3A5F")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            set_run_font(run, size=8, bold=True, color=RGBColor(255, 255, 255))


def T(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=8, bold=True, color=RGBColor(255, 255, 255))
        shade(table.rows[0].cells[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=8)
    doc.add_paragraph()


def build():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)

    for _ in range(2):
        doc.add_paragraph()
    P(doc, "UNIVERSIDAD PRIVADA ANTENOR ORREGO", bold=True, size=14, align="center")
    P(doc, "Facultad de Ingeniería", bold=True, size=12, align="center")
    P(
        doc,
        "Escuela Profesional de Ingeniería de Computación y Sistemas",
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(
        doc,
        "TECHNICAL DESIGN & DEVELOPMENT REPORT (TDDR)",
        bold=True,
        size=16,
        align="center",
    )
    P(
        doc,
        "Plataforma de Trazabilidad Académica con Agentes de IA Generativa\n"
        "para Gestión de Proyectos Universitarios de Software",
        bold=True,
        size=13,
        align="center",
    )
    doc.add_paragraph()
    P(
        doc,
        "Informe Técnico de Diseño, Desarrollo y Validación",
        italic=True,
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "Autor: Jean Marcos Meneses Simón", size=11, align="center")
    P(doc, "Asesor / Docente revisor: Ing. Walter Cueva", size=11, align="center")
    P(doc, "Fecha: Julio 2026", size=11, align="center")
    doc.add_page_break()

    H(doc, "SECCIÓN 0 — Metadatos del proyecto")
    T(
        doc,
        ["Campo", "Descripción"],
        [
            ["Tipo de solución", "Híbrida: Web + IA Generativa multi-agente"],
            ["Dominio", "Educación superior / Ingeniería de software"],
            [
                "Palabras clave",
                "Academic Traceability; Multi-Agent Systems; Generative AI; "
                "Software Engineering Education; Learning Analytics; Human-in-the-Loop",
            ],
            ["Repositorio", "https://github.com/Jmmsimon/trazabilidad-taller-I"],
            [
                "Dataset",
                "Commits GitHub + backlog CSV/Notion + Firestore (sin dataset ML externo)",
            ],
            [
                "Stack",
                "Next.js 16.2.6 · React 19.2.4 · FastAPI · LangGraph · Firebase · Gemini/Claude",
            ],
            ["Versión", "Commit ddf1890 (main)"],
        ],
    )

    H(doc, "SECCIÓN 1 — Problema y motivación técnica")
    H(doc, "1.1 Descripción del problema real", 2)
    P(
        doc,
        "En proyectos universitarios de software, el docente debe verificar si el avance "
        "declarado (backlog/hitos) existe realmente en el repositorio Git. La revisión manual "
        "no escala con varios proyectos concurrentes: evaluación subjetiva, ítems Done sin "
        "evidencia y feedback tardío.",
    )
    P(
        doc,
        "Evidencia de contexto (literatura): Garousi et al. documentan el desajuste universidad–industria "
        "en SE [8][13]. GitHub Classroom mejora visibilidad del proceso, pero no automatiza trazabilidad "
        "backlog↔código [2][5][14]. La IA generativa agrava riesgos de integridad sin verificación de evidencia [17][18][19].",
    )
    H(doc, "1.2 Brecha tecnológica", 2)
    P(
        doc,
        "Falta una plataforma académica que integre co-creación de backlog con agentes, lectura profunda "
        "del repo, mapeo commits↔hitos↔Kanban evidenciable y auditoría CSV vs código (semáforo) con HITL docente.",
    )
    H(doc, "1.3 Pregunta de investigación", 2)
    P(
        doc,
        "¿En qué medida una plataforma web con orquestación multi-agente (LangGraph) y lectura profunda "
        "de GitHub permite mejorar la trazabilidad entre backlog/hitos académicos y evidencia de código, "
        "reduciendo desviaciones entre avance declarado y real, en proyectos universitarios de software?",
        italic=True,
    )
    H(doc, "1.4 Objetivos", 2)
    P(
        doc,
        "General: Diseñar, implementar y validar una plataforma de trazabilidad académica basada en "
        "agentes de IA generativa que vincule backlog, hitos y repositorio Git.",
    )
    B(
        doc,
        [
            "OE1. Discovery multi-agente (drafter–validator–PO).",
            "OE2. Tracking deep + mapeo commits↔hitos↔Kanban.",
            "OE3. Dashboards por rol + analítica.",
            "OE4. Auditoría CSV/Notion vs código (solo reporte).",
            "OE5. Validación exploratoria funcional en escenario real.",
        ],
    )
    H(doc, "1.5 Alcance y limitaciones", 2)
    B(
        doc,
        [
            "Incluye: Next+FastAPI+LangGraph+Firestore+GitHub; auto-Kanban; auditoría CSV solo-reporte; PDF/JSON.",
            "Excluye: LMS completo; repos privados sin token; plagio forense; estudio estadístico multi-cohorte.",
        ],
    )
    H(doc, "1.6 Contribución técnica", 2)
    P(
        doc,
        "Arquitectura híbrida: grafos LangGraph tipados + evidencia deep GitHub + Kanban compartido "
        "alumno–docente + auditoría CSV desacoplada (solo reporte).",
    )

    H(doc, "SECCIÓN 2 — Revisión de literatura técnica")
    H(doc, "2.1 Marco conceptual", 2)
    B(
        doc,
        [
            "Trazabilidad requisitos/código [6][7][12][16]",
            "Learning analytics [1][4]",
            "Git en educación SE [2][5][8][14][15]",
            "LLM para code review/análisis [3][9][10][20]",
            "Multi-agentes generativos en educación SE [11]",
            "Integridad académica e IA [17][18][19]",
        ],
    )
    H(doc, "2.2 Estado del arte (≥15 trabajos con DOI)", 2)
    T(
        doc,
        ["Ref", "Año", "Tipo", "Técnica", "Contexto", "Hallazgo/Métrica", "Limitación", "DOI"],
        [
            ["[1]", "2018", "Survey LA", "LA + Learning Design", "IEEE TLT", "43 estudios", "Falta framework LD", "10.1109/TLT.2018.2868673"],
            ["[2]", "2021", "Edu tool", "GitHub Classroom", "SEENG/IEEE", "Apoyo project-based", "Sin hitos académicos", "10.1109/SEENG53126.2021.00013"],
            ["[3]", "2024", "LLM SE", "ChatGPT refinement", "ICSE", "EM/BLEU > baseline", "No Kanban académico", "10.1145/3597503.3639101"],
            ["[4]", "2020", "Survey LA", "Constructs LA", "IEEE TLT", "Revisión constructs", "Poco foco Git", "10.1109/TLT.2020.2999970"],
            ["[5]", "2022", "Edu Git", "GitHub lessons", "ACE/ACM", "Lecciones VCS", "Eval. superficial commits", "10.1145/3511861.3511879"],
            ["[6]", "2020", "Traceability", "Semantic graphs", "IEEE RE", "TLR 5 datasets", "Sin UI docente", "10.1109/RE48521.2020.00028"],
            ["[7]", "2022", "Traceability", "IR mapping", "SEKE", "34 estudios VSM/LSI", "Sin agentes GenAI", "10.18293/SEKE2022-098"],
            ["[8]", "2020", "SE Education", "SLR industry gap", "IEEE Software", "33 estudios", "Sin plataforma trazabilidad", "10.1109/MS.2018.2880823"],
            ["[9]", "2024", "LLM SE", "Code-change LLMs", "ACM TOSEM", "ICL/PEFT review/msgs", "Sin Kanban académico", "10.1145/3709358"],
            ["[10]", "2023", "LLM Review", "LLaMA-Reviewer", "ISSRE", "SOTA <1% params", "Sin auditoría backlog", "10.1109/ISSRE59848.2023.00026"],
            ["[11]", "2024", "Multi-agent", "DevCoach", "ACM L@S", "Mejora aprendizaje SDLC", "No verifica backlog↔código", "10.1145/3657604.3664663"],
            ["[12]", "2022", "Traceability", "Interactive GA", "Springer ISSE", "Prometedora reuse", "Sin dashboard académico", "10.1007/s11334-021-00418-2"],
            ["[13]", "2019", "SE Education", "Meta-analysis", "JSS", ">4000 datapoints", "Herramientas genéricas", "10.1016/j.jss.2019.06.044"],
            ["[14]", "2023", "Classroom", "GitHub autograde", "SIGCSE", "Distribución tareas", "Sin semáforo", "10.1145/3545947.3569627"],
            ["[15]", "2019", "Git Classroom", "Outcomes GitHub", "SIGCSE", "Asoc. learning outcomes", "Sin agentes IA", "10.1145/3287324.3287460"],
            ["[16]", "2020", "Traceability", "ML + rules", "Appl. Sci.", "Mejora vs IR", "Sin HITL web", "10.3390/app10207253"],
            ["[17]", "2023", "Integrity", "ChatGPT CS assess.", "ACM TOCE", "IA aprueba exams", "Detección insuficiente", "10.1145/3633287"],
            ["[18]", "2023", "Integrity", "GPT-4 detection", "J. Acad. Ethics", "54.5% reportados", "Tools imperfectas", "10.1007/s10805-023-09492-6"],
            ["[19]", "2023", "Integrity", "AI detectors", "Int.J.Educ.Integr.", "No confiables", "No trazabilidad", "10.1007/s40979-023-00146-z"],
            ["[20]", "2024", "LLM codegen", "ClassEval", "ICSE", "Benchmark class-level", "No eval. proyectos curso", "10.1145/3597503.3639219"],
        ],
    )
    H(doc, "2.3 Gaps", 2)
    T(
        doc,
        ["Capacidad", "LMS", "GitHub Classroom", "Jira/Notion", "Code LLM", "Esta solución"],
        [
            ["Co-creación backlog IA", "No", "No", "Parcial", "Parcial", "Sí"],
            ["Lectura profunda repo", "No", "Básica", "No", "Sí", "Sí"],
            ["Commits↔hitos", "No", "No", "No", "No", "Sí"],
            ["Auto-Kanban evidenciado", "No", "No", "Manual", "No", "Sí"],
            ["CSV vs código", "No", "No", "No", "No", "Sí"],
            ["HITL docente", "Sí", "Parcial", "Sí", "No", "Sí"],
        ],
    )
    H(doc, "2.4 Justificación tecnológica", 2)
    T(
        doc,
        ["Capa", "Elegida", "Descartada", "Motivo"],
        [
            ["Agentes", "LangGraph", "AutoGen lineal", "Grafos + estado + HITL"],
            ["Backend", "FastAPI", "Django", "Async + BackgroundTasks"],
            ["Frontend", "Next.js 16", "CRA/Vue", "App Router + proxies"],
            ["DB", "Firestore", "Postgres", "Docs flexibles"],
            ["LLM", "Gemini/Claude", "Un vendor", "Resiliencia"],
        ],
    )

    H(doc, "SECCIÓN 3 — Diseño de la solución")
    H(doc, "3.1 Arquitectura (Figura 1)", 2)
    P(
        doc,
        "La Figura 1 describe la arquitectura en capas de la plataforma: interfaces por rol "
        "(estudiante, docente, administrador), BFF mediante API Routes de Next.js, backend FastAPI "
        "con tareas en segundo plano, orquestación multi-agente en LangGraph (Discovery, Tracking "
        "y Backlog Audit) e integración con GitHub, proveedores LLM y Firebase Auth/Firestore. "
        "El código fuente Mermaid utilizado para generar la Figura 1 se presenta en el Anexo A.",
    )
    P(
        doc,
        "Flujo lógico: UI Next.js → API Routes → FastAPI → LangGraph (Discovery | Tracking | Audit) "
        "→ Firestore / GitHub / LLM.",
        italic=True,
    )
    H(doc, "3.2 Requerimientos funcionales", 2)
    T(
        doc,
        ["ID", "Requerimiento", "Pri", "Obj"],
        [
            ["RF01", "Auth por roles", "Alta", "OE3"],
            ["RF02", "Discovery multi-agente", "Alta", "OE1"],
            ["RF03", "Aprobar/rechazar roadmap", "Alta", "OE3"],
            ["RF04", "Vincular repo/demo", "Alta", "OE2"],
            ["RF05", "Tracking deep", "Alta", "OE2"],
            ["RF06", "Auto-Kanban", "Alta", "OE2"],
            ["RF07", "Analítica commits↔hitos", "Alta", "OE3"],
            ["RF08", "Validar hitos/HU", "Alta", "OE3"],
            ["RF09", "Auditoría CSV solo-reporte", "Alta", "OE4"],
            ["RF10", "Export PDF/JSON", "Media", "OE5"],
        ],
    )
    H(doc, "3.3 Artefactos de diseño", 2)
    P(
        doc,
        "El diseño se documenta mediante los siguientes artefactos, alineados al modelo "
        "documental de Firestore (Anexo C) y a los endpoints del Anexo B:",
    )
    B(
        doc,
        [
            "Casos de uso principales: autenticación por rol; co-creación/aprobación de roadmap "
            "(Discovery); vinculación de repositorio y demo; análisis de trazabilidad (Tracking); "
            "consulta de Kanban y analítica; validación de hitos; auditoría de backlog CSV "
            "(solo reporte); exportación PDF/JSON. Actores: Estudiante, Docente, Administrador y Sistema.",
            "Modelo de datos: esquema documental Firestore (usuarios, proyectos, tracking, "
            "backlog_scrum, backlog_audit), detallado en el Anexo C.",
            "Secuencia del análisis de trazabilidad: la UI invoca POST /proyectos/{id}/tracking/iniciar; "
            "FastAPI ejecuta el grafo Tracking en BackgroundTasks (devops deep → competency → "
            "analyst → reporter + AUTO-KANBAN); el cliente consulta GET /tracking/status hasta "
            "estado completed y actualiza Kanban/analítica.",
            "Interfaces implementadas: portales de acceso por rol; discovery/roadmap del estudiante; "
            "tablero Kanban; analítica de estudiante y docente; módulo de auditoría CSV del docente "
            "(Anexo F).",
        ],
    )
    H(doc, "3.4 Stack tecnológico y decisiones de arquitectura (ADR)", 2)
    B(
        doc,
        [
            "Stack: Next.js 16.2.6, React 19.2.4, FastAPI, LangGraph, Firebase, Gemini/Claude.",
            "ADR-01: pipeline de Tracking compartido entre estudiante y docente (una sola fuente de verdad).",
            "ADR-02: auditoría CSV desacoplada y solo-reporte (no modifica el Kanban del alumno).",
            "ADR-03: lectura profunda (deep) del repositorio también en el Tracking del estudiante.",
            "ADR-04: middleware de Next.js deshabilitado en despliegue local por interferencia "
            "con rutas App Router (bug de enrutamiento 404); autenticación vía API Routes y cookies.",
        ],
    )
    H(doc, "3.5 Seguridad y consentimiento informado", 2)
    B(
        doc,
        [
            "Firebase Authentication con cookie httpOnly, control de roles, CORS restrictivo, "
            "secretos en variables de entorno y sello SHA-256 en exportaciones PDF.",
            "Consentimiento informado: no aplica (Anexo G). La evaluación es técnico-exploratoria "
            "sobre el sistema en operación, sin reclutamiento de participantes humanos externos "
            "ni recolección de PII de terceros.",
        ],
    )

    H(doc, "SECCIÓN 4 — Desarrollo e implementación")
    B(
        doc,
        [
            "Metodología: prototipado evolutivo, sprints cortos y Design Science Research (DSR).",
            "Módulos implementados: Discovery, Tracking con AUTO-KANBAN, Backlog Audit, "
            "dashboards de estudiante y docente.",
            "Repositorio: https://github.com/Jmmsimon/trazabilidad-taller-I (commit ddf1890).",
            "Entorno de ejecución del piloto: backend Python 3.12 (uvicorn :8000) y frontend "
            "Next.js (npm run dev :3000).",
        ],
    )

    H(doc, "SECCIÓN 5 — Evaluación exploratoria (sistema funcional)")
    H(doc, "5.1 Estrategia de evaluación", 2)
    P(
        doc,
        "La evaluación se declara técnico-exploratoria (prueba de concepto sobre un proyecto piloto). "
        "No se aplican pruebas inferenciales (Wilcoxon / t-test) porque el tamaño muestral es N = 1 "
        "proyecto; dichos contrastes requieren al menos dos grupos o una cohorte mayor "
        "(p. ej. ≥10 proyectos) para reportar p-value y tamaño del efecto de forma válida. "
        "Como trabajo futuro se contempla una cohorte ampliada.",
    )
    H(doc, "5.2 Protocolo de prueba ejecutado", 2)
    B(
        doc,
        [
            "Despliegue local del backend (uvicorn :8000) y del frontend (npm run dev).",
            "Rol estudiante: ejecución de Analizar con repositorio GitHub y demo desplegada.",
            "Rol docente: ejecución de Analizar y consulta de Analítica (score, commits, resumen).",
            "Verificación del Kanban tras AUTO-KANBAN (ítems con evidencia GitHub en estado Hecho).",
            "Ejecución opcional de Auditoría de Avances con CSV (modo solo-reporte).",
        ],
    )
    H(doc, "5.3 Resultados medidos (piloto proj-6c87a9d7)", 2)
    T(
        doc,
        ["Métrica", "Valor", "Fuente"],
        [
            ["tracking_status", "completed", "GET /tracking/status"],
            ["score_integridad", "100/100", "UI + tracking"],
            ["% competencias", "80%", "reporte_competencias"],
            ["Commits analizados", "15", "estado_repo.commits"],
            ["Ítems backlog", "29", "backlog_scrum"],
            ["Ítems done post auto-Kanban", "25/29 (86.2%)", "conteo estados"],
            ["HU-002 / EN-003 GitHub", "done", "kanban_updates + UI"],
            ["CI status", "pass", "estado_repo"],
            ["CSV afecta Kanban", "No", "solo_reporte"],
            ["Rutas Next (post ajuste middleware)", "200 OK", "logs Next"],
        ],
    )
    P(
        doc,
        "Media ± desviación estándar: no reportable con N = 1 proyecto. En una cohorte ≥10 "
        "proyectos se reportarán media ± DE por métrica.",
    )
    H(doc, "5.4 Comparación cualitativa", 2)
    T(
        doc,
        ["Método", "Commits↔hitos", "Auto-Kanban", "CSV vs código", "HITL"],
        [
            ["Manual docente", "Baja", "No", "Manual", "Sí"],
            ["GitHub+Notion", "Media", "No", "No integrada", "Parcial"],
            ["Solo chat LLM", "Variable", "Ad hoc", "Ad hoc", "Débil"],
            ["Propuesto", "Alta", "Sí (86.2% done)", "Sí (semáforo)", "Sí"],
        ],
    )
    H(doc, "5.5 Discusión", 2)
    P(
        doc,
        "La lectura profunda del repositorio, el prompt de análisis con criterios estrictos de "
        "evidencia y el refuerzo por entregables ya vinculados (repo_url / demo_url) reducen "
        "falsos estados in_progress: ítems con integración GitHub evidenciada pasan a done. "
        "La analítica docente expone score de integridad, commits y resumen de competencias. "
        "La principal limitación del piloto es el tamaño muestral (N = 1) y la dependencia de "
        "APIs LLM y GitHub.",
    )

    H(doc, "SECCIÓN 6 — Discusión integradora")
    B(
        doc,
        [
            "Respecto a la pregunta de investigación: en el piloto, la plataforma vincula "
            "backlog/hitos con evidencia Git de forma automatizada y auditables.",
            "Contribuciones: grafos Discovery + Tracking + Audit; Kanban compartido; "
            "auditoría CSV solo-reporte.",
            "Limitaciones: N = 1; variabilidad LLM; middleware Next deshabilitado en local; "
            "auditoría nativa CSV (no xlsx).",
            "Amenazas a la validez: sesgo de autoría del piloto; validez externa limitada; "
            "constructo de integridad basado en heurísticas + LLM.",
            "Trabajo futuro: cohorte ≥10 proyectos; GitHub App; estabilización del proxy Next; "
            "validador visual de CSV.",
        ],
    )

    H(doc, "SECCIÓN 7 — Conclusiones")
    P(
        doc,
        "Plataforma operativa de trazabilidad académica con agentes IA. Piloto: score 100, "
        "15 commits, 80% competencias, 25/29 ítems done evidenciados, auditoría CSV sin alterar Kanban. "
        "Código en GitHub.",
    )

    H(doc, "SECCIÓN 8 — Referencias (DOI)")
    refs = [
        "[1] K. Mangaroska and M. Giannakos, “Learning Analytics for Learning Design…,” IEEE Trans. Learn. Technol., doi: 10.1109/TLT.2018.2868673.",
        "[2] M. A. Nelson and L. Ponciano, “Experiences… Github Classroom…,” SEENG, 2021, doi: 10.1109/SEENG53126.2021.00013.",
        "[3] Empirical study ChatGPT automated code refinement, ICSE 2024, doi: 10.1145/3597503.3639101.",
        "[4] “Where is the Learning in Learning Analytics?…,” IEEE TLT, 2020, doi: 10.1109/TLT.2020.2999970.",
        "[5] Y.-C. Tu et al., “GitHub in the Classroom: Lessons Learnt,” ACE, 2022, doi: 10.1145/3511861.3511879.",
        "[6] A. Schlutter and A. Vogelsang, “Trace Link Recovery…,” IEEE RE, 2020, doi: 10.1109/RE48521.2020.00028.",
        "[7] IR approaches for requirements trace recovery (mapping), SEKE 2022, doi: 10.18293/SEKE2022-098.",
        "[8] V. Garousi et al., “Closing the Gap… SE Education…,” IEEE Software, 2020, doi: 10.1109/MS.2018.2880823.",
        "[9] J. Liu et al., “Exploring… LLMs for Code-Change Related Tasks,” ACM TOSEM, 2024, doi: 10.1145/3709358.",
        "[10] LLaMA-Reviewer, ISSRE 2023, doi: 10.1109/ISSRE59848.2023.00026.",
        "[11] DevCoach, L@S 2024, doi: 10.1145/3657604.3664663.",
        "[12] Hamdi et al., “Requirements traceability recovery…,” Innov. Syst. Softw. Eng., 2022, doi: 10.1007/s11334-021-00418-2.",
        "[13] Garousi et al., “Aligning SE education… meta-analysis,” JSS, 2019, doi: 10.1016/j.jss.2019.06.044.",
        "[14] Hecht et al., “Distributing… GitHub Classroom,” SIGCSE 2023, doi: 10.1145/3545947.3569627.",
        "[15] Hsing and Gennarelli, “Using GitHub in the Classroom Predicts…,” SIGCSE 2019, doi: 10.1145/3287324.3287460.",
        "[16] Wang et al., “Combining ML and Logical Reasoning… Traceability,” Appl. Sci., 2020, doi: 10.3390/app10207253.",
        "[17] Richards et al., “Bob or Bot… ChatGPT,” ACM TOCE, 2023, doi: 10.1145/3633287.",
        "[18] Perkins et al., “Detection of GPT-4 Generated Text…,” J. Acad. Ethics, 2023, doi: 10.1007/s10805-023-09492-6.",
        "[19] Weber-Wulff et al., “Testing of detection tools for AI-generated text,” Int. J. Educ. Integrity, 2023, doi: 10.1007/s40979-023-00146-z.",
        "[20] Du et al., “Evaluating LLMs in Class-Level Code Generation,” ICSE 2024, doi: 10.1145/3597503.3639219.",
    ]
    for r in refs:
        P(doc, r, size=9)

    H(doc, "SECCIÓN 9 — Anexos")
    H(doc, "Anexo A — Fuente Mermaid de la Figura 1 (arquitectura)", 2)
    P(
        doc,
        "A continuación se incluye el código Mermaid a partir del cual se generó la Figura 1 "
        "(exportación vectorial SVG).",
    )
    P(doc, MERMAID, size=8)
    H(doc, "Anexo B — Extracto de API", 2)
    T(
        doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["POST", "/proyectos/{id}/tracking/iniciar", "Inicia Tracking (deep + AUTO-KANBAN)"],
            ["GET", "/proyectos/{id}/tracking/status", "Consulta estado, score y backlog"],
            ["POST", "/profesor/proyectos/{id}/backlog-audit", "Inicia auditoría CSV (solo reporte)"],
            ["GET", "/profesor/proyectos/{id}/backlog-audit/status", "Consulta semáforo de auditoría"],
        ],
    )
    H(doc, "Anexo C — Diccionario Firestore", 2)
    T(
        doc,
        ["Campo", "Descripción"],
        [
            ["backlog_scrum", "Ítems Kanban del proyecto (estados backlog|todo|in_progress|done)"],
            ["tracking", "Resultado del análisis: score, commits, competencias, kanban_updates"],
            ["backlog_audit", "Resultado semáforo CSV vs código; afecta_kanban_alumno = false"],
        ],
    )
    H(doc, "Anexo D — Instalación del entorno de desarrollo", 2)
    P(
        doc,
        "Backend: crear entorno virtual Python 3.12, instalar requirements.txt y ejecutar "
        "uvicorn main:app --reload --port 8000. Frontend: npm install && npm run dev "
        "(puerto 3000). El archivo middleware de Next.js permanece deshabilitado "
        "(middleware.ts.disabled) por el ADR-04.",
    )
    H(doc, "Anexo E — Datos del piloto", 2)
    P(
        doc,
        "Repositorio público del piloto y archivo de ejemplo docs/backlog.csv utilizados "
        "en la auditoría de avances.",
    )
    H(doc, "Anexo F — Interfaces del sistema (piloto)", 2)
    P(
        doc,
        "Las interfaces validadas en el piloto corresponden a los módulos en producción "
        "funcional del sistema:",
    )
    B(
        doc,
        [
            "Analítica docente: score de integridad 100/100, 80% de competencias y 15 commits analizados.",
            "Kanban con ítems vinculados a GitHub en estado Hecho tras AUTO-KANBAN (25/29).",
            "Confirmación de análisis completado en la interfaz de Tracking.",
            "Módulo de Auditoría de Avances en modo solo-reporte (no altera el Kanban del alumno).",
            "Registro de ejecución backend con lectura deep y aplicación de AUTO-KANBAN.",
        ],
    )
    H(doc, "Anexo G — Consentimiento informado", 2)
    P(
        doc,
        "No aplica. El presente informe no incluye estudio con participantes humanos externos; "
        "la validación corresponde a una evaluación técnico-exploratoria del sistema "
        "desarrollado por el autor, sin recolección de datos personales de terceros.",
        bold=True,
    )

    doc.save(OUT)
    print("OK", OUT)


if __name__ == "__main__":
    build()
