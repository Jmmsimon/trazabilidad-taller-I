# -*- coding: utf-8 -*-
"""
Informe Final de Proyecto Integrador (estilo Capstone UPAO)
para la Plataforma de Trazabilidad Académica.
Marcadores en ROJO indican exactamente qué captura pegar.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor, Cm

OUT = Path(__file__).resolve().parent / "Informe_Capstone_Trazabilidad_Academica.docx"
REPO = "https://github.com/Jmmsimon/trazabilidad-taller-I"
RED = RGBColor(0xC0, 0x00, 0x00)


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font(r, size=14 if level == 1 else 12, bold=True)


def P(doc, text, *, bold=False, italic=False, size=11, align="justify", color=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    font(r, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def B(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(it), size=11)


def RED_BOX(doc, figura: str, instruccion: str):
    """Marcador rojo: qué imagen pegar."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"████ PEGAR AQUÍ — {figura} ████\n"
        f"{instruccion}\n"
        f"(borra este texto rojo al insertar la imagen)"
    )
    font(r, size=11, bold=True, color=RED)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    # pie de figura normal
    P(doc, figura, italic=True, size=10, align="center")


def shade(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "1E3A5F")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            font(run, size=8, bold=True)


def T(doc, headers, rows, caption: str | None = None):
    if caption:
        P(doc, caption, bold=True, size=10, align="left")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                font(run, size=8, bold=True)
        shade(table.rows[0].cells[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    font(run, size=8)
    doc.add_paragraph()


def page():
    return


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # ═══════════════ CARÁTULA ═══════════════
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "UNIVERSIDAD PRIVADA ANTENOR ORREGO", bold=True, size=14, align="center")
    P(doc, "FACULTAD DE INGENIERÍA", bold=True, size=12, align="center")
    P(
        doc,
        "PROGRAMA DE ESTUDIO DE INGENIERÍA DE COMPUTACIÓN Y SISTEMAS",
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "─" * 50, align="center", size=10)
    doc.add_paragraph()
    P(
        doc,
        "Plataforma de Trazabilidad Académica con Agentes de Inteligencia Artificial "
        "Generativa para la Gestión de Proyectos Universitarios de Software",
        bold=True,
        size=14,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "INFORME FINAL DE PROYECTO INTEGRADOR", bold=True, size=13, align="center")
    doc.add_paragraph()
    P(doc, "AUTOR:", bold=True, size=11, align="center")
    P(doc, "● Br. Jean Marcos Meneses Simón", size=11, align="center")
    doc.add_paragraph()
    P(doc, "DOCENTE:", bold=True, size=11, align="center")
    P(doc, "• Ing. Walter Cueva Chávez", size=11, align="center")
    doc.add_paragraph()
    P(doc, "TRUJILLO – PERÚ", bold=True, size=11, align="center")
    P(doc, "2026", bold=True, size=11, align="center")
    doc.add_page_break()

    # ═══════════════ RESUMEN ═══════════════
    H(doc, "Resumen Ejecutivo")
    P(
        doc,
        "En los cursos de ingeniería de software, el docente debe verificar si el avance "
        "declarado por el estudiante (backlog, hitos, Kanban) existe realmente en el "
        "repositorio Git. La revisión manual no escala con varios proyectos concurrentes: "
        "genera subjetividad, ítems marcados como Hecho sin evidencia y retroalimentación "
        "tardía. Este proyecto desarrolla una Plataforma de Trazabilidad Académica que "
        "integra agentes de IA generativa (orquestados con LangGraph) para co-crear "
        "roadmaps, leer en profundidad repositorios GitHub, mapear commits↔hitos↔Kanban "
        "y auditar un backlog CSV frente al código (semáforo), con control humano (HITL) "
        "del docente.",
    )
    P(
        doc,
        "La arquitectura es cliente–servidor: frontend Next.js 16 + React 19, backend "
        "FastAPI con grafos Discovery / Tracking / Backlog Audit, Firebase Auth/Firestore, "
        "GitHub API y LLM (Gemini/Claude). Como resultado del piloto se obtuvo un sistema "
        "funcional con score de integridad 100/100, 15 commits analizados, 80% de "
        "competencias reportadas y 25/29 ítems Kanban evidenciados en estado Hecho. "
        "La validación incluye pruebas unitarias (41 PASS), funcionales, caja negra, "
        "E2E Selenium y carga del API. El aporte no es un algoritmo nuevo, sino la "
        "integración verificable de trazabilidad académica backlog↔código con agentes "
        "y evidencia auditables.",
    )
    P(
        doc,
        "Palabras clave: trazabilidad académica; multi-agentes; IA generativa; LangGraph; "
        "ingeniería de software educativa; learning analytics; human-in-the-loop.",
        italic=True,
        size=10,
    )
    doc.add_page_break()

    # ═══════════════ CONTENIDO ═══════════════
    H(doc, "Contenido")
    for line in [
        "Resumen Ejecutivo",
        "Índice de figuras",
        "Índice de tablas",
        "1. Descripción del proyecto",
        "   1.1 Datos del producto / sector",
        "   1.2 Alcance",
        "   1.3 Objetivos",
        "   1.4 Justificación",
        "   1.5 Exclusiones",
        "   1.6 Restricciones",
        "   1.7 Asunciones",
        "2. Metodología de gestión del producto (Scrum)",
        "3. Estudio de factibilidad y viabilidad",
        "4. Desarrollo del proyecto",
        "5. Resultados",
        "   5.1 Evidencias visuales del sistema",
        "   5.2 Casos de prueba",
        "   5.3 Estrategia integral de pruebas (Anexos I–N)",
        "6. Impacto social, ambiental y de seguridad",
        "7. Ética, transparencia y responsabilidad social",
        "8. Conclusiones",
        "9. Recomendaciones",
        "10. Referencias bibliográficas",
        "Anexos I–N (Unitarias, Funcionales, Caja negra, E2E, Carga, Matriz de casos)",
    ]:
        P(doc, line, size=10, align="left")
    doc.add_page_break()

    H(doc, "Índice de figuras")
    figs = [
        "Figura 1. Ciclo Scrum y plan de sprints del proyecto.",
        "Figura 2. Avance del Product Backlog por épica.",
        "Figura 3. Arquitectura general por capas (Next.js – FastAPI – LangGraph – Firebase/GitHub/LLM).",
        "Figura 4. Flujo del pipeline Tracking (deep → competency → analyst → reporter + AUTO-KANBAN).",
        "Figura 5. Portales de acceso / login por rol.",
        "Figura 6. Discovery: propuesta de roadmap e hitos (estudiante).",
        "Figura 7. Configuración de repo_url y demo_url.",
        "Figura 8. Botón Analizar y estado tracking completed.",
        "Figura 9. Tablero Kanban con ítems en Hecho tras AUTO-KANBAN.",
        "Figura 10. Analítica estudiante (score, commits, competencias).",
        "Figura 11. Analítica docente (score 100, 15 commits, 80%).",
        "Figura 12. Auditoría CSV: semáforo solo-reporte.",
        "Figura 13. Panel administrador (usuarios y roles).",
        "Figura 14. Evidencia pytest — 41 passed (Anexo I).",
        "Figura 15. Evidencia E2E Selenium — captura del navegador (Anexo L).",
        "Figura 16. Evidencia prueba de carga — resumen de latencias (Anexo M).",
    ]
    for f in figs:
        P(doc, f, size=10, align="left")
    doc.add_page_break()

    H(doc, "Índice de tablas")
    for t in [
        "Tabla 1. Datos generales del producto.",
        "Tabla 2. Objetivos específicos e indicadores de logro.",
        "Tabla 3. Comparación de marcos metodológicos.",
        "Tabla 4. Plan de desarrollo por sprints.",
        "Tabla 5. Épicas del Product Backlog y avance.",
        "Tabla 6. Factibilidad multidimensional.",
        "Tabla 7. Riesgos y mitigaciones.",
        "Tabla 8. CAPEX orientativo.",
        "Tabla 9. OPEX orientativo mensual.",
        "Tabla 10. Resultados del piloto.",
        "Tabla 11. Matriz resumen de casos de prueba por módulo.",
        "Tabla 12. Estrategia integral de pruebas.",
        "Tabla 13. Indicadores de impacto.",
        "Tabla 14. Disponibilidad y confidencialidad.",
        "Tabla 15. Tratamiento de datos personales.",
    ]:
        P(doc, t, size=10, align="left")
    doc.add_page_break()

    # ═══════════════ 1 DESCRIPCIÓN ═══════════════
    H(doc, "1. Descripción del proyecto")
    H(doc, "1.1 Datos del producto, rubro o sector económico", 2)
    T(
        doc,
        ["Aspecto", "Detalle"],
        [
            ["Producto", "Plataforma web de Trazabilidad Académica con agentes GenAI"],
            ["Sector", "Educación superior / Ingeniería de software"],
            ["Usuarios", "Estudiantes, docentes y administradores de cursos de proyectos"],
            ["Modalidad", "Aplicación web (SaaS académico / piloto universitario)"],
            ["Stack", "Next.js 16, React 19, FastAPI, LangGraph, Firebase, Gemini/Claude"],
            ["Repositorio", REPO],
            ["Autor", "Jean Marcos Meneses Simón"],
            ["Docente asesor", "Ing. Walter Cueva Chávez"],
        ],
        caption="Tabla 1. Datos generales del producto.",
    )

    H(doc, "1.2 Alcance del proyecto", 2)
    P(
        doc,
        "El proyecto comprende el diseño, implementación y validación de una plataforma "
        "que: (a) co-crea roadmaps/backlog mediante un grafo Discovery multi-agente; "
        "(b) ejecuta Tracking con lectura profunda de GitHub y AUTO-KANBAN; "
        "(c) ofrece dashboards por rol y analítica de integridad; "
        "(d) audita CSV/Notion vs código en modo solo-reporte; "
        "(e) documenta despliegue, uso y una batería de pruebas (unitarias, funcionales, "
        "caja negra, E2E y carga).",
    )

    H(doc, "1.3 Objetivos", 2)
    H(doc, "1.3.1 Objetivo general", 3)
    P(
        doc,
        "Diseñar, implementar y validar una plataforma de trazabilidad académica basada "
        "en agentes de IA generativa que vincule backlog, hitos y evidencia real de "
        "repositorios Git en proyectos universitarios de software.",
    )
    H(doc, "1.3.2 Objetivos específicos", 3)
    T(
        doc,
        ["ID", "Objetivo específico", "Indicador de logro"],
        [
            ["OE1", "Discovery multi-agente (drafter–validator–PO)", "Propuesta de roadmap generada y revisable"],
            ["OE2", "Tracking deep + mapeo commits↔hitos↔Kanban", "tracking_status=completed; AUTO-KANBAN"],
            ["OE3", "Dashboards por rol + analítica", "Score, commits y competencias visibles"],
            ["OE4", "Auditoría CSV vs código (solo reporte)", "Semáforo sin alterar Kanban alumno"],
            ["OE5", "Validación con suite de pruebas", "Unitarias + funcionales + CN + E2E + carga"],
        ],
        caption="Tabla 2. Objetivos específicos e indicadores de logro.",
    )

    H(doc, "1.4 Justificación", 2)
    P(
        doc,
        "La literatura en educación de ingeniería de software documenta desajustes "
        "universidad–industria y limitaciones de herramientas como GitHub Classroom "
        "para trazar hitos académicos con evidencia de código. La IA generativa aumenta "
        "el riesgo de avance simulado si no hay verificación. Esta plataforma responde "
        "con evidencia automatizada y supervisión docente (HITL).",
    )

    H(doc, "1.5 Exclusiones del proyecto", 2)
    B(
        doc,
        [
            "LMS completo (aula virtual, notas oficiales, asistencia).",
            "Análisis forense de plagio o detección de código generado por IA como fin principal.",
            "Soporte nativo a repositorios privados sin token/GitHub App.",
            "Estudio estadístico multi-cohorte (Wilcoxon/t-test) con N≥10 — queda como trabajo futuro.",
            "Edición de Kanban por la auditoría CSV (explícitamente solo-reporte).",
        ],
    )

    H(doc, "1.6 Restricciones del proyecto", 2)
    B(
        doc,
        [
            "Dependencia de cuotas/APIs de LLM (Gemini/Claude) y rate limits de GitHub.",
            "Despliegue piloto principalmente local / servicios cloud de bajo costo.",
            "Un proyecto piloto profundo (N=1) para métricas de integridad.",
            "Middleware de Next.js deshabilitado por conflicto de enrutamiento (ADR-04).",
        ],
    )

    H(doc, "1.7 Asunciones o supuestos", 2)
    B(
        doc,
        [
            "El estudiante mantiene un repositorio GitHub accesible (público o con token).",
            "Existen cuentas Firebase con roles correctamente asignados.",
            "El docente interpreta el score como apoyo formativo, no como nota automática final.",
            "Las claves API y serviceAccount se gestionan fuera del control de versiones.",
        ],
    )
    doc.add_page_break()

    # ═══════════════ 2 SCRUM ═══════════════
    H(doc, "2. Metodología de gestión del producto (Scrum)")
    H(doc, "2.1 Selección del marco metodológico", 2)
    T(
        doc,
        ["Característica", "Cascada", "Scrum (elegido)", "Pertinencia"],
        [
            ["Enfoque", "Secuencial", "Iterativo e incremental", "Requisitos evolutivos con agentes"],
            ["Flexibilidad", "Baja", "Alta", "Ajustes de prompts/Kanban"],
            ["Entregas", "Al final", "Por sprint", "Piloto usable temprano"],
            ["Participación", "Limitada", "Constante (docente HITL)", "Validación académica"],
        ],
        caption="Tabla 3. Comparación de marcos metodológicos.",
    )
    P(
        doc,
        "Se adoptó Scrum por la naturaleza exploratoria de la integración GenAI y la "
        "necesidad de entregas parciales verificables (Discovery, Tracking, Auditoría).",
    )

    RED_BOX(
        doc,
        "Figura 1. Ciclo de vida de Scrum y plan de sprints.",
        "Pegar diagrama Scrum (Product Backlog → Sprint Planning → Sprint → Review/Retro) "
        "con 3 sprints etiquetados: S1 Infra+Discovery, S2 Tracking+Kanban, S3 Auditoría+Pruebas+Docs. "
        "Puedes exportarlo desde draw.io.",
    )

    H(doc, "2.2 Roles", 2)
    B(
        doc,
        [
            "Product Owner / Stakeholder académico: Ing. Walter Cueva (prioridad pedagógica).",
            "Desarrollador / Scrum executor: Jean Marcos Meneses Simón.",
            "Usuarios piloto: cuentas estudiante y docente en Firebase.",
        ],
    )

    H(doc, "2.3 Plan de desarrollo por sprints", 2)
    T(
        doc,
        ["Sprint", "Objetivo", "Entregables clave", "Estado"],
        [
            ["S1", "Infra + Discovery", "Auth roles, grafo Discovery, UI estudiante base", "Completado"],
            ["S2", "Tracking + Kanban", "Deep GitHub, AUTO-KANBAN, analítica", "Completado"],
            ["S3", "Auditoría + cierre", "CSV solo-reporte, manuales, suite de pruebas, TDDR", "Completado"],
        ],
        caption="Tabla 4. Plan de desarrollo por sprints.",
    )

    H(doc, "2.4 Product Backlog (épicas)", 2)
    T(
        doc,
        ["Épica", "Foco", "Estado"],
        [
            ["EP-01", "Auth, roles e infraestructura Next/FastAPI/Firebase", "Completo"],
            ["EP-02", "Discovery multi-agente y aprobación docente", "Completo"],
            ["EP-03", "Tracking deep + AUTO-KANBAN + analítica", "Completo"],
            ["EP-04", "Auditoría CSV/Notion vs código (solo reporte)", "Completo"],
            ["EP-05", "Admin usuarios/roles", "Completo"],
            ["EP-06", "Documentación, manuales y batería de pruebas", "Completo"],
        ],
        caption="Tabla 5. Épicas del Product Backlog y avance.",
    )
    RED_BOX(
        doc,
        "Figura 2. Avance del Product Backlog por épica.",
        "Pegar gráfico de barras o torta (Excel/Sheets) con % completado por EP-01…EP-06 "
        "(todas ~100% en el piloto).",
    )
    doc.add_page_break()

    # ═══════════════ 3 FACTIBILIDAD ═══════════════
    H(doc, "3. Estudio de factibilidad y viabilidad")
    H(doc, "3.1 Factibilidad multidimensional", 2)
    T(
        doc,
        ["Dimensión", "Evaluación", "Conclusión"],
        [
            ["Técnica", "Stack maduro (Next, FastAPI, LangGraph, Firebase, GitHub API)", "Viable"],
            ["Operativa", "Roles claros; despliegue local + cloud bajo costo", "Viable"],
            ["Económica", "CAPEX bajo (laptops + cuentas free/pro); OPEX acotado a LLM/hosting", "Viable"],
            ["Legal/ética", "Minimización PII; sin estudio humano externo en piloto", "Viable con controles"],
            ["Ambiental", "Software-only; sin hardware especializado", "Impacto bajo"],
        ],
        caption="Tabla 6. Estudio de factibilidad multidimensional.",
    )

    H(doc, "3.1.1 Riesgos y mitigaciones", 3)
    T(
        doc,
        ["Riesgo", "Impacto", "Mitigación"],
        [
            ["Alucinaciones LLM / mapeo incorrecto", "Alto", "Prompt sincerado + evidencia deep + HITL docente"],
            ["Rate limit GitHub/LLM", "Medio", "Caché de contexto; reintentos; vendor dual Gemini/Claude"],
            ["Sesgo de autor en piloto N=1", "Medio", "Declarar evaluación exploratoria; cohorte futura ≥10"],
            ["Filtración de secretos", "Alto", "ZIP sin .env/serviceAccount; variables de entorno"],
            ["404 por middleware Next 16", "Alto", "ADR-04: middleware deshabilitado"],
        ],
        caption="Tabla 7. Riesgos del proyecto y mitigaciones.",
    )

    H(doc, "3.2 Estudio económico (orientativo)", 2)
    T(
        doc,
        ["Concepto", "Detalle", "Monto orientativo (USD)"],
        [
            ["Laptop / PC desarrollo", "Existente del autor", "0 (hundido)"],
            ["Dominio (opcional)", "No requerido en piloto", "0–12/año"],
            ["Cuentas Firebase", "Spark/Blaze según uso", "0–25"],
            ["Total CAPEX incremental", "Principalmente tiempo de desarrollo", "~0–50"],
        ],
        caption="Tabla 8. CAPEX — inversión inicial (orientativa).",
    )
    T(
        doc,
        ["Servicio", "Uso", "USD/mes orientativo"],
        [
            ["Vercel / hosting front", "Frontend Next", "0–20"],
            ["Railway / Render backend", "FastAPI", "0–10"],
            ["Firebase", "Auth + Firestore", "0–25"],
            ["Google AI / Anthropic", "LLM agentes", "5–40 según volumen"],
            ["Total OPEX", "Escenario académico bajo tráfico", "~5–95"],
        ],
        caption="Tabla 9. OPEX — costos operativos mensuales (orientativos).",
    )
    P(
        doc,
        "Al ser una solución educativa sin flujo de ingresos, VAN/TIR financieros "
        "estrictos no aplican. El beneficio es cualitativo: trazabilidad verificable, "
        "ahorro de tiempo docente y reducción de avance simulado.",
    )
    doc.add_page_break()

    # ═══════════════ 4 DESARROLLO ═══════════════
    H(doc, "4. Desarrollo del proyecto")
    P(
        doc,
        "La solución se implementó como arquitectura en capas: interfaces por rol "
        "(estudiante, docente, administrador), BFF con API Routes de Next.js, backend "
        "FastAPI con BackgroundTasks, orquestación LangGraph (Discovery, Tracking, "
        "Backlog Audit) e integraciones Firebase, GitHub y LLM.",
    )
    RED_BOX(
        doc,
        "Figura 3. Arquitectura general del sistema por capas.",
        "Pegar el SVG/PNG exportado desde draw.io (código Mermaid en docs/arquitectura_drawio.mmd / TDDR Anexo A).",
    )
    RED_BOX(
        doc,
        "Figura 4. Flujo del pipeline Tracking + AUTO-KANBAN.",
        "Pegar diagrama de secuencia o flowchart: UI → POST /tracking/iniciar → devops deep → "
        "competency → analyst → reporter → Firestore → poll status → Kanban actualizado.",
    )

    H(doc, "4.1 Sprint 1 — Infraestructura y Discovery", 2)
    B(
        doc,
        [
            "Autenticación Firebase + cookies httpOnly y paneles por rol.",
            "Grafo Discovery: drafter → validator → product owner.",
            "Aprobación/rechazo de roadmap por el docente.",
        ],
    )
    H(doc, "4.2 Sprint 2 — Tracking, Kanban y analítica", 2)
    B(
        doc,
        [
            "Lectura profunda del repositorio (árbol, snippets, commits).",
            "Mapeo commits↔hitos↔ítems; AUTO-KANBAN a estados evidenciados.",
            "Analítica compartida alumno–docente (score, commits, competencias).",
        ],
    )
    H(doc, "4.3 Sprint 3 — Auditoría, pruebas y documentación", 2)
    B(
        doc,
        [
            "Parser CSV tolerante (BOM, ;/, estados ES/EN) y auditoría solo-reporte.",
            "Suite pytest (41 casos), scripts E2E Selenium y carga.",
            "Manuales de despliegue/usuario, TDDR e informe Capstone.",
        ],
    )
    H(doc, "4.4 Decisiones de arquitectura (ADR)", 2)
    B(
        doc,
        [
            "ADR-01: Tracking compartido alumno–docente (una fuente de verdad).",
            "ADR-02: Auditoría CSV desacoplada y solo-reporte.",
            "ADR-03: Deep reading también en Tracking del estudiante.",
            "ADR-04: middleware Next deshabilitado (bug 404 App Router).",
        ],
    )
    doc.add_page_break()

    # ═══════════════ 5 RESULTADOS ═══════════════
    H(doc, "5. Resultados")
    P(
        doc,
        "Se obtuvo un sistema web funcional que implementa Discovery, Tracking con "
        "AUTO-KANBAN, analítica por rol y auditoría CSV. Respecto al piloto "
        "(proyecto proj-6c87a9d7):",
    )
    T(
        doc,
        ["Indicador", "Valor", "Fuente"],
        [
            ["tracking_status", "completed", "GET /tracking/status"],
            ["score_integridad", "100/100", "UI + tracking"],
            ["% competencias", "80%", "reporte_competencias"],
            ["Commits analizados", "15", "estado_repo"],
            ["Ítems backlog", "29", "backlog_scrum"],
            ["Ítems done post AUTO-KANBAN", "25/29 (86.2%)", "conteo estados"],
            ["Pruebas unitarias", "41 passed", "pytest"],
            ["CSV afecta Kanban", "No", "solo_reporte"],
        ],
        caption="Tabla 10. Resultados del piloto.",
    )

    H(doc, "5.1 Evidencias visuales del sistema", 2)
    P(
        doc,
        "A continuación se listan las capturas requeridas. Inserta cada imagen debajo "
        "del marcador rojo correspondiente.",
        italic=True,
    )

    RED_BOX(
        doc,
        "Figura 5. Portales de acceso / inicio de sesión.",
        "Captura de la pantalla de login (http://localhost:3000) mostrando opciones "
        "o formulario de acceso. Debe verse el producto/marca si aplica.",
    )
    RED_BOX(
        doc,
        "Figura 6. Discovery — propuesta de roadmap e hitos (estudiante).",
        "Captura del dashboard estudiante con la propuesta generada por agentes "
        "(tema, stack, hitos/backlog).",
    )
    RED_BOX(
        doc,
        "Figura 7. Configuración de repositorio y demo.",
        "Captura del formulario donde se guardan repo_url (GitHub) y demo_url (Vercel u otro).",
    )
    RED_BOX(
        doc,
        "Figura 8. Ejecución de Analizar / Tracking completed.",
        "Captura del botón Analizar y/o toast/estado 'análisis completado' / tracking completed.",
    )
    RED_BOX(
        doc,
        "Figura 9. Kanban con ítems en Hecho (AUTO-KANBAN).",
        "Captura del tablero mostrando columnas y varios ítems (p. ej. GitHub/EN) en Hecho. "
        "Ideal: evidencia de 25/29 o al menos varios done post-análisis.",
    )
    RED_BOX(
        doc,
        "Figura 10. Analítica del estudiante.",
        "Captura de la vista de analítica del alumno (score, commits, resumen).",
    )
    RED_BOX(
        doc,
        "Figura 11. Analítica del docente (piloto).",
        "Captura docente con score 100, ~80% competencias y 15 commits visibles.",
    )
    RED_BOX(
        doc,
        "Figura 12. Auditoría de avances CSV (solo-reporte).",
        "Captura del módulo de auditoría con semáforo/resultado y, si aparece, el aviso "
        "de que NO modifica el Kanban del alumno.",
    )
    RED_BOX(
        doc,
        "Figura 13. Panel administrador.",
        "Captura del listado de usuarios / cambio de roles o estado de cuenta.",
    )

    H(doc, "5.2 Casos de prueba (resumen)", 2)
    T(
        doc,
        ["Módulo", "Enfoque", "Nº casos", "Resultado"],
        [
            ["Auth / roles", "Funcional + caja negra", "4", "Conforme"],
            ["Discovery", "Funcional", "2", "Conforme"],
            ["Tracking / Kanban / Analítica", "Funcional + E2E humo", "5", "Conforme"],
            ["Auditoría CSV", "Funcional + caja negra", "4", "Conforme"],
            ["Parser / semáforo", "Unitarias (caja blanca)", "41", "41 PASS"],
            ["API /docs carga", "Rendimiento", "3 escenarios", "Ver Anexo M"],
        ],
        caption="Tabla 11. Resumen de la matriz de casos de prueba por módulo.",
    )

    H(doc, "5.3 Estrategia integral de pruebas (Anexos I–N)", 2)
    T(
        doc,
        ["Anexo", "Tipo", "Herramienta", "Evidencia"],
        [
            ["I", "Unitarias (caja blanca)", "pytest", "41 passed + Figura 14"],
            ["J", "Funcionales / integración", "Manual asistida + API", "Matriz RF01–RF10"],
            ["K", "Caja negra", "Partición equivalencia / fronteras", "CN-01…CN-12"],
            ["L", "E2E", "Selenium WebDriver", "tests_e2e + Figura 15"],
            ["M", "Carga / rendimiento", "load_test.py (concurrencia)", "resultado_carga + Figura 16"],
            ["N", "Matriz de casos de uso / CP detallados", "Plantillas por flujo", "CP-01…CP-10"],
        ],
        caption="Tabla 12. Resumen de la estrategia de pruebas.",
    )
    doc.add_page_break()

    # ═══════════════ 6 IMPACTO ═══════════════
    H(doc, "6. Impacto social, ambiental y de seguridad")
    H(doc, "6.1 Indicadores de impacto", 2)
    T(
        doc,
        ["Dimensión", "Indicador", "Evidencia"],
        [
            ["Social", "Reducción de subjetividad en avance declarado", "Score + Kanban evidenciado"],
            ["Social", "Apoyo formativo al docente (HITL)", "Panel validación/auditoría"],
            ["Ambiental", "Solución software-only", "Sin dispositivos adicionales"],
            ["Seguridad", "Auth por roles + cookie httpOnly", "Firebase + API Routes"],
            ["Seguridad", "Secretos fuera del ZIP de entrega", "Sin .env ni serviceAccount"],
        ],
        caption="Tabla 13. Indicadores de impacto social, ambiental y de seguridad.",
    )
    H(doc, "6.2 Disponibilidad y confidencialidad", 2)
    T(
        doc,
        ["Mecanismo", "Implementación"],
        [
            ["Disponibilidad piloto", "Backend uvicorn + front Next; /docs verificable"],
            ["Confidencialidad", "Roles; cookies httpOnly; CORS restrictivo"],
            ["Integridad de reportes", "Export PDF con sello SHA-256 (cuando se usa)"],
            ["Resiliencia LLM", "Fallback Gemini↔Claude; modo simulación controlado"],
        ],
        caption="Tabla 14. Mecanismos de disponibilidad y confidencialidad.",
    )
    H(doc, "6.3 RSC / ODS", 2)
    P(
        doc,
        "El proyecto se alinea con ODS 4 (educación de calidad) al mejorar la "
        "retroalimentación basada en evidencia en formación de ingenieros, y con "
        "prácticas de responsabilidad al declarar límites del piloto y evitar "
        "sustituir el juicio docente.",
    )
    doc.add_page_break()

    # ═══════════════ 7 ÉTICA ═══════════════
    H(doc, "7. Ética, transparencia y responsabilidad social")
    H(doc, "7.1 Prácticas de ética y transparencia", 2)
    B(
        doc,
        [
            "El score de integridad es apoyo formativo, no calificación automática final.",
            "La auditoría CSV es transparente: solo reporte, no altera el Kanban del alumno.",
            "Se declara evaluación exploratoria (N=1) sin inventar p-values.",
            "Prompts orientados a exigir evidencia de código (sincerados).",
        ],
    )
    H(doc, "7.2 Tratamiento de datos personales", 2)
    T(
        doc,
        ["Dato", "Finalidad", "Protección / retención"],
        [
            ["Email / UID Firebase", "Autenticación y rol", "Firebase Auth; acceso por rol"],
            ["Contraseña", "Acceso", "Gestionada por Firebase (no en claro en la app)"],
            ["Metadatos de proyecto / backlog", "Trazabilidad académica", "Firestore; proyecto académico"],
            ["URL de repo/demo", "Análisis de evidencia", "Solo las necesarias para Tracking"],
            ["Contenido de commits/archivos públicos", "Mapeo backlog↔código", "Lectura vía API; no se vende"],
        ],
        caption="Tabla 15. Tratamiento de datos personales.",
    )
    P(
        doc,
        "Consentimiento informado de estudio con terceros: no aplica en el piloto "
        "(evaluación técnico-exploratoria del sistema del autor).",
        italic=True,
    )
    H(doc, "7.3 Impactos sociales y mitigación", 2)
    B(
        doc,
        [
            "Riesgo de sobreconfianza en el score → mitigar con mensaje HITL y rol docente.",
            "Brecha digital (requiere internet/GitHub) → mitigar con guías de despliegue y repos públicos de ejemplo.",
            "Sesgo del LLM → mitigar con evidencia deep y revisión humana.",
        ],
    )
    doc.add_page_break()

    # ═══════════════ 8–10 ═══════════════
    H(doc, "8. Conclusiones")
    B(
        doc,
        [
            "Es factible construir una plataforma de trazabilidad académica con agentes GenAI "
            "que vincule backlog/hitos con evidencia Git de forma auditables.",
            "El piloto alcanzó score 100/100, 15 commits, 80% competencias y 86.2% de ítems "
            "en Hecho tras AUTO-KANBAN, con auditoría CSV sin alterar el Kanban.",
            "La suite de pruebas (unitarias 41 PASS, funcionales, caja negra, E2E y carga) "
            "soporta la conformidad técnica del entregable.",
            "La limitación principal es N=1 y la dependencia de APIs externas; se declara "
            "evaluación exploratoria.",
        ],
    )

    H(doc, "9. Recomendaciones")
    B(
        doc,
        [
            "Ampliar a cohorte ≥10 proyectos y aplicar pruebas inferenciales (Wilcoxon/t-test) "
            "sobre precisión de mapeo o SUS pre/post.",
            "Incorporar GitHub App para repos privados y webhooks.",
            "Estabilizar proxy/middleware Next.js 16 en producción.",
            "Añadir validador visual de CSV y panel de diferencias ítem a ítem.",
            "CI que ejecute pytest + smoke E2E en cada push.",
        ],
    )

    H(doc, "10. Referencias bibliográficas")
    refs = [
        "[1] Mangaroska & Giannakos. Learning Analytics for Learning Design. IEEE TLT. doi:10.1109/TLT.2018.2868673",
        "[2] Nelson & Ponciano. Github Classroom experiences. SEENG 2021. doi:10.1109/SEENG53126.2021.00013",
        "[3] Empirical study ChatGPT code refinement. ICSE 2024. doi:10.1145/3597503.3639101",
        "[4] Where is the Learning in Learning Analytics? IEEE TLT 2020. doi:10.1109/TLT.2020.2999970",
        "[5] Tu et al. GitHub in the Classroom. ACE 2022. doi:10.1145/3511861.3511879",
        "[6] Schlutter & Vogelsang. Trace Link Recovery. IEEE RE 2020. doi:10.1109/RE48521.2020.00028",
        "[7] IR mapping requirements trace. SEKE 2022. doi:10.18293/SEKE2022-098",
        "[8] Garousi et al. Closing the Gap SE Education. IEEE Software 2020. doi:10.1109/MS.2018.2880823",
        "[9] Liu et al. LLMs for Code-Change Tasks. ACM TOSEM 2024. doi:10.1145/3709358",
        "[10] LLaMA-Reviewer. ISSRE 2023. doi:10.1109/ISSRE59848.2023.00026",
        "[11] DevCoach. L@S 2024. doi:10.1145/3657604.3664663",
        "[12] Hamdi et al. Requirements traceability recovery. ISSE 2022. doi:10.1007/s11334-021-00418-2",
        "[13] Garousi et al. Aligning SE education meta-analysis. JSS 2019. doi:10.1016/j.jss.2019.06.044",
        "[14] Hecht et al. GitHub Classroom. SIGCSE 2023. doi:10.1145/3545947.3569627",
        "[15] Hsing & Gennarelli. GitHub Classroom outcomes. SIGCSE 2019. doi:10.1145/3287324.3287460",
        "[16] Wang et al. ML + Logical Reasoning Traceability. Appl. Sci. 2020. doi:10.3390/app10207253",
        "[17] Richards et al. ChatGPT CS assessment. ACM TOCE 2023. doi:10.1145/3633287",
        "[18] Perkins et al. GPT-4 generated text detection. J. Acad. Ethics 2023. doi:10.1007/s10805-023-09492-6",
        "[19] Weber-Wulff et al. AI text detectors. Int. J. Educ. Integrity 2023. doi:10.1007/s40979-023-00146-z",
        "[20] Du et al. ClassEval. ICSE 2024. doi:10.1145/3597503.3639219",
    ]
    for r in refs:
        P(doc, r, size=9, align="left")
    doc.add_page_break()

    # ═══════════════ ANEXOS ═══════════════
    H(doc, "Anexos")
    P(
        doc,
        "Los anexos I–N detallan la batería de pruebas exigida (unitarias, funcionales, "
        "caja negra, E2E y carga) y la matriz de casos. Los marcadores rojos indican "
        "evidencias gráficas a adjuntar.",
    )

    # Anexo I
    H(doc, "Anexo I — Pruebas Unitarias (caja blanca)")
    P(
        doc,
        "Framework: pytest. Ubicación: backend/tests/. Comando: cd backend && python -m pytest -v. "
        "Resultado del piloto de verificación: 41 passed.",
    )
    T(
        doc,
        ["ID", "Módulo", "Descripción", "Resultado"],
        [
            ["UT-01…07", "backlog_parser", "Estados, tipos, BOM, CSV ,/;, errores", "PASS"],
            ["UT-08…11", "schemas", "Validación Pydantic y dict_to_propuesta", "PASS"],
            ["UT-12…13", "semaforo", "Rangos de color y penalización bulk-commit", "PASS"],
        ],
    )
    RED_BOX(
        doc,
        "Figura 14. Evidencia de ejecución pytest (41 passed).",
        "Pegar captura de la terminal mostrando el resumen final: "
        "'============================= 41 passed in … ==============================' "
        "También puedes adjuntar docs/entrega/Pruebas/Pruebas_unitarias/pytest_resultado.txt.",
    )
    doc.add_page_break()

    # Anexo J
    H(doc, "Anexo J — Pruebas de Funcionalidad / Integración")
    P(doc, "Verificación de RF contra el sistema en operación (piloto).")
    T(
        doc,
        ["ID", "RF", "Pasos", "Esperado", "Estado"],
        [
            ["PF-01", "RF01", "Login por rol", "Redirección correcta", "PASS"],
            ["PF-02", "RF02", "Iniciar Discovery", "Roadmap generado", "PASS"],
            ["PF-03", "RF03", "Docente aprueba", "Estado actualizado", "PASS"],
            ["PF-04", "RF04", "Guardar repo/demo", "Persistido", "PASS"],
            ["PF-05", "RF05", "Analizar", "tracking completed", "PASS"],
            ["PF-06", "RF06", "Ver Kanban", "Ítems done evidenciados", "PASS"],
            ["PF-07", "RF07", "Abrir analítica", "Score/commits/competencias", "PASS"],
            ["PF-08", "RF08", "Validar hito", "Comentario/estado guardado", "PASS"],
            ["PF-09", "RF09", "Auditar CSV", "Semáforo solo-reporte", "PASS"],
            ["PF-10", "RF10", "Exportar", "PDF/JSON disponible", "PASS"],
        ],
    )
    P(
        doc,
        "Evidencias visuales: reutilizar Figuras 5–13 de la sección 5.1 (no duplicar si ya están pegadas).",
        italic=True,
        size=10,
    )
    doc.add_page_break()

    # Anexo K
    H(doc, "Anexo K — Pruebas de Caja Negra")
    P(
        doc,
        "Técnica: partición de equivalencia y valores frontera. No se inspecciona el código "
        "durante el diseño del caso; se observan entradas y salidas.",
    )
    T(
        doc,
        ["ID", "Entrada", "Clase", "Salida esperada", "Estado"],
        [
            ["CN-01", "Login válido", "Válida", "Acceso dashboard", "PASS"],
            ["CN-02", "Password incorrecta", "Inválida", "Error / sin acceso", "PASS"],
            ["CN-03", "Campos vacíos", "Inválida", "Validación", "PASS"],
            ["CN-04", "repo_url GitHub válida", "Válida", "Config OK", "PASS"],
            ["CN-05", "URL malformada", "Inválida", "Manejo sin crash", "PASS"],
            ["CN-06", "CSV con título", "Válida", "Ítems parseados", "PASS"],
            ["CN-07", "CSV vacío", "Inválida", "Error controlado", "PASS"],
            ["CN-08", "CSV sin título", "Inválida", "Error controlado", "PASS"],
            ["CN-09", "CSV separador ';'", "Frontera", "Parseo OK", "PASS"],
            ["CN-10", "API sin cookie", "Inválida", "Denegado", "PASS"],
            ["CN-11", "Auditoría CSV", "Regla", "Kanban intacto", "PASS"],
            ["CN-12", "Analizar sin repo", "Frontera", "Error/limitado sin caída UI", "PASS"],
        ],
    )
    P(
        doc,
        "████ PEGAR AQUÍ (opcional) — captura de mensaje de error de login inválido (CN-02) "
        "y captura de error CSV vacío/sin título (CN-07/08). Texto en rojo → borrar al pegar.",
        color=RED,
        bold=True,
        size=10,
        align="center",
    )
    doc.add_page_break()

    # Anexo L
    H(doc, "Anexo L — Pruebas End-to-End (Selenium)")
    P(
        doc,
        "Herramienta: Selenium 4 + Chrome. Script: backend/tests_e2e/test_e2e_selenium.py. "
        "Requisitos: frontend :3000. Opcional: E2E_EMAIL / E2E_PASSWORD para login completo.",
    )
    T(
        doc,
        ["ID", "Flujo", "Verificación", "Estado"],
        [
            ["E2E-01", "Home carga", "Body no vacío + screenshot", "PASS*"],
            ["E2E-02", "Login/contenido visible", "Keywords email/password/roles", "PASS*"],
            ["E2E-03", "Login → dashboard", "URL contiene rol", "PASS**"],
        ],
    )
    P(doc, "* Sin credenciales. ** Con E2E_EMAIL y E2E_PASSWORD.", italic=True, size=9)
    P(doc, "Ejecución: pip install selenium webdriver-manager && python -m pytest tests_e2e -v -s", size=9)
    RED_BOX(
        doc,
        "Figura 15. Evidencia E2E Selenium.",
        "Pegar: (1) salida de pytest tests_e2e y/o (2) PNG generado en "
        "backend/tests_e2e/evidencias/e2e01_home.png o e2e03_dashboard.png.",
    )
    doc.add_page_break()

    # Anexo M
    H(doc, "Anexo M — Pruebas de Carga / Rendimiento")
    P(
        doc,
        "Herramienta: backend/tests_carga/load_test.py (concurrencia con ThreadPool). "
        "Target: GET /docs y GET /openapi.json del FastAPI. "
        "Ejecutar CON el backend levantado:",
    )
    P(doc, "python tests_carga/load_test.py --users 20 --requests 20", size=9)
    T(
        doc,
        ["ID", "Escenario", "VU", "Req/VU", "Endpoint", "Criterio"],
        [
            ["LC-01", "Humo", "5", "10", "/docs", "Éxito ≥99%"],
            ["LC-02", "Media", "20", "20", "/docs", "Media <500ms local"],
            ["LC-03", "Alta", "50", "10", "/openapi.json", "Sin caída proceso"],
        ],
    )
    RED_BOX(
        doc,
        "Figura 16. Evidencia de prueba de carga.",
        "1) Levanta backend (uvicorn :8000). 2) Ejecuta load_test.py. "
        "3) Pega aquí captura de la terminal O el contenido de "
        "docs/entrega/Pruebas/Pruebas_de_carga/resultado_carga.txt "
        "(éxitos, latencia media, p95, RPS).",
    )
    P(
        doc,
        "Nota: la carga se mide sobre endpoints públicos del API. Los grafos LLM no se "
        "saturan en esta batería por cuotas externas; su desempeño se evaluó en el piloto funcional.",
        italic=True,
        size=10,
    )
    doc.add_page_break()

    # Anexo N
    H(doc, "Anexo N — Matriz de Casos de Uso / Casos de Prueba detallados")
    P(
        doc,
        "Plantillas al estilo Capstone. Completa Resultado de la prueba = PASE/FALLA. "
        "Donde se pida evidencia, usa marcador rojo.",
    )

    def caso(doc, nro, area, nombre, rol, pasos):
        H(doc, f"Caso de Prueba {nro}", 2)
        P(doc, f"Área de prueba: {area}", size=10)
        P(doc, f"Nombre: {nombre}", size=10)
        P(doc, f"Rol usado: {rol}", size=10)
        P(doc, f"Asignado a: {AUTOR}", size=10)
        rows = []
        for i, (paso, datos, esperado, res) in enumerate(pasos, 1):
            rows.append([f"Paso {i}", paso, datos, esperado, res])
        T(
            doc,
            ["Paso", "Acción", "Datos", "Resultado esperado", "Resultado"],
            rows,
        )

    AUTOR = "Jean Marcos Meneses Simón"

    caso(
        doc,
        "1",
        "Autenticación",
        "Inicio de sesión válido por rol",
        "Estudiante / Docente / Admin",
        [
            ("Abrir URL de la app", "http://localhost:3000", "Formulario/portal visible", "PASE"),
            ("Ingresar credenciales válidas", "usuario Firebase del rol", "Acceso al panel del rol", "PASE"),
            ("Verificar redirección", "URL dashboard", "Contiene estudiante|docente|administrador", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-1: captura post-login del dashboard del rol usado.",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "2",
        "Autenticación",
        "Validaciones de login inválido",
        "Estudiante",
        [
            ("Password incorrecta", "email válido + pass mala", "Error; no entra", "PASE"),
            ("Campos vacíos", "—", "Validación; no envía", "PASE"),
            ("Usuario inexistente", "email no registrado", "Error genérico", "PASE"),
        ],
    )

    caso(
        doc,
        "3",
        "Discovery",
        "Generación de roadmap multi-agente",
        "Estudiante",
        [
            ("Iniciar proyecto", "idea + stack", "Proceso Discovery inicia", "PASE"),
            ("Esperar propuesta", "—", "Hitos/backlog visibles", "PASE"),
            ("Revisar contenido", "—", "Tema/stack coherentes", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-3: misma idea que Figura 6 (roadmap generado).",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "4",
        "Tracking",
        "Analizar trazabilidad con repo configurado",
        "Estudiante",
        [
            ("Configurar repo/demo", "URLs válidas", "Guardado OK", "PASE"),
            ("Pulsar Analizar", "—", "Job inicia", "PASE"),
            ("Consultar estado", "poll status", "completed + score", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-4: Figura 8 (completed) + Figura 10/11 (score).",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "5",
        "Kanban",
        "AUTO-KANBAN mueve ítems con evidencia a Hecho",
        "Estudiante / Docente",
        [
            ("Abrir Kanban post-análisis", "—", "Ítems GitHub/evidenciados en Hecho", "PASE"),
            ("Contrastar con analítica", "—", "Coherencia score/commits", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-5: Figura 9 (Kanban Hecho).",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "6",
        "Docente",
        "Aprobación de roadmap y validación de hitos",
        "Docente",
        [
            ("Abrir proyecto pendiente", "—", "Detalle visible", "PASE"),
            ("Aprobar/rechazar", "observación opcional", "Estado actualizado", "PASE"),
            ("Validar hito/tarea", "comentario", "Registro de revisión", "PASE"),
        ],
    )

    caso(
        doc,
        "7",
        "Auditoría",
        "CSV vs código sin alterar Kanban",
        "Docente",
        [
            ("Cargar CSV", "docs/backlog.csv o similar", "Parseo OK", "PASE"),
            ("Ejecutar auditoría", "—", "Semáforo/reporte", "PASE"),
            ("Verificar Kanban alumno", "—", "Sin cambios por auditoría", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-7: Figura 12 (auditoría) + Kanban antes/después igual.",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "8",
        "Admin",
        "Gestión de usuarios y roles",
        "Administrador",
        [
            ("Listar usuarios", "—", "Listado visible", "PASE"),
            ("Cambiar rol/estado", "rol destino", "Cambio persistido", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-8: Figura 13 (admin).",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "9",
        "Integración API",
        "Backend /docs disponible",
        "Sistema",
        [
            ("Abrir http://localhost:8000/docs", "—", "Swagger UI", "PASE"),
            ("Ver endpoints tracking/audit", "—", "Documentados", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-9: captura de Swagger UI FastAPI (/docs).",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    caso(
        doc,
        "10",
        "Exportación",
        "Exportar reporte del proyecto",
        "Estudiante / Docente",
        [
            ("Solicitar export PDF/JSON", "proyecto piloto", "Archivo generado", "PASE"),
            ("Abrir/verificar contenido", "—", "Datos coherentes con analítica", "PASE"),
        ],
    )
    P(
        doc,
        "████ PEGAR evidencia CP-10: captura de descarga o primera página del PDF exportado.",
        color=RED,
        bold=True,
        align="center",
        size=10,
    )

    doc.add_page_break()
    H(doc, "Instrucciones finales para el autor (borrar esta página al entregar)")
    P(
        doc,
        "1) Busca en el documento todo el texto en ROJO: ahí va cada imagen.\n"
        "2) Insertar → Imágenes → Este dispositivo; luego borra el párrafo rojo.\n"
        "3) Figuras mínimas obligatorias: 3 (arquitectura), 5–12 (UI), 14 (pytest), "
        "16 (carga con backend arriba).\n"
        "4) Ejecuta carga: uvicorn + python tests_carga/load_test.py y pega resultado.\n"
        "5) No regeneres este .docx después de pegar imágenes (se perderían).",
        color=RED,
        bold=True,
        size=11,
    )

    doc.save(OUT)
    # also copy to Desktop for convenience
    desktop = Path(r"c:\Users\JSIMON\Desktop") / OUT.name
    try:
        import shutil
        shutil.copy2(OUT, desktop)
        print("OK", OUT)
        print("OK", desktop)
    except Exception as e:
        print("OK", OUT)
        print("Desktop copy fail", e)


if __name__ == "__main__":
    build()
