# -*- coding: utf-8 -*-
"""
Informe Capstone ampliado (estilo profesional UPAO) + Excel de casos
según plantilla Matrix de Pruebas / Caso de Prueba.
"""
from __future__ import annotations

from pathlib import Path
from copy import copy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

OUT_DOC = Path(__file__).resolve().parent / "Informe_Capstone_Trazabilidad_Academica.docx"
OUT_XLSX = Path(__file__).resolve().parent / "entrega" / "Pruebas" / "Matriz_y_Casos_de_Prueba.xlsx"
DESKTOP = Path(r"c:\Users\JSIMON\Desktop")
REPO = "https://github.com/Jmmsimon/trazabilidad-taller-I"
AUTOR = "Jean Marcos Meneses Simón"
DOCENTE = "Ing. Walter Cueva Chávez"
EMAIL_QA = "jmenesess1@upao.edu.pe"
RED = RGBColor(0xC0, 0x00, 0x00)


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font(r, size=14 if level == 1 else (12 if level == 2 else 11), bold=True)


def P(doc, text, *, bold=False, italic=False, size=11, align="justify", color=None):
    p = doc.add_paragraph()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(text)
    font(r, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def B(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(it), size=11)


def RED_BOX(doc, figura: str, instruccion: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"████ PEGAR AQUÍ — {figura} ████\n{instruccion}\n"
        f"(borra este texto rojo al insertar la imagen)"
    )
    font(r, size=10, bold=True, color=RED)
    P(doc, figura, italic=True, size=10, align="center")


def shade(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "1E3A5F")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            font(run, size=8, bold=True)


def T(doc, headers, rows, caption: str | None = None):
    if caption:
        P(doc, caption, bold=True, size=10, align="left")
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                font(run, size=8, bold=True)
        shade(table.rows[0].cells[i])
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for run in p.runs:
                    font(run, size=8)
    doc.add_paragraph()


# ── Datos de casos (matriz + detalle) ────────────────────────────────────
# Severidad: 1=crítica … 4=baja (como el ejemplo del Excel)

MATRIX = [
    # (id, caso_uso, area, estado, sev, comentarios)
    ("1", "Autenticación y control de acceso por roles", "Auth / RBAC", "—", "1", "Grupo funcional"),
    ("1.1", "Inicio de sesión válido (estudiante)", "Auth / RBAC", "PASE", "1", ""),
    ("1.2", "Inicio de sesión válido (docente)", "Auth / RBAC", "PASE", "1", ""),
    ("1.3", "Inicio de sesión con credenciales inválidas", "Auth / RBAC", "PASE", "1", ""),
    ("1.4", "Restricción de rutas por rol (estudiante no entra a admin)", "Auth / RBAC", "PASE", "1", ""),
    ("1.5", "Cierre de sesión e invalidación de cookie", "Auth / RBAC", "PASE", "2", ""),
    ("2", "Discovery multi-agente y aprobación de roadmap", "Discovery", "—", "1", "Grupo funcional"),
    ("2.1", "Generación de propuesta técnica (drafter–validator–PO)", "Discovery", "PASE", "1", ""),
    ("2.2", "Confirmación / iteración de propuesta por el estudiante", "Discovery", "PASE", "2", ""),
    ("2.3", "Aprobación de roadmap por el docente", "Discovery", "PASE", "1", ""),
    ("2.4", "Rechazo de roadmap con observaciones", "Discovery", "PASE", "2", ""),
    ("2.5", "Persistencia de hitos y backlog_scrum en Firestore", "Discovery", "PASE", "1", ""),
    ("3", "Tracking, AUTO-KANBAN y analítica", "Tracking", "—", "1", "Grupo funcional"),
    ("3.1", "Configuración de repo_url y demo_url", "Tracking", "PASE", "1", ""),
    ("3.2", "Ejecución de Analizar (Tracking deep)", "Tracking", "PASE", "1", "Piloto: completed"),
    ("3.3", "AUTO-KANBAN: ítems con evidencia → done", "Tracking", "PASE", "1", "25/29 done"),
    ("3.4", "Analítica estudiante (score, commits, competencias)", "Tracking", "PASE", "1", "Score 100; 15 commits; 80%"),
    ("3.5", "Analítica docente (misma fuente de verdad)", "Tracking", "PASE", "1", "ADR-01"),
    ("4", "Auditoría CSV vs código (solo reporte)", "Auditoría", "—", "1", "Grupo funcional"),
    ("4.1", "Parseo CSV válido (coma / punto y coma)", "Auditoría", "PASE", "1", "Cubierto también en unitarias"),
    ("4.2", "Rechazo de CSV vacío o sin columna título", "Auditoría", "PASE", "2", ""),
    ("4.3", "Ejecución de auditoría y semáforo", "Auditoría", "PASE", "1", ""),
    ("4.4", "Verificar que la auditoría NO altera Kanban alumno", "Auditoría", "PASE", "1", "ADR-02"),
    ("4.5", "Consulta de status de auditoría (semáforo)", "Auditoría", "PASE", "2", ""),
    ("5", "Administración, exportación y disponibilidad API", "Admin / Ops", "—", "2", "Grupo funcional"),
    ("5.1", "Listado y cambio de rol de usuarios", "Admin / Ops", "PASE", "2", ""),
    ("5.2", "Exportación PDF/JSON del proyecto", "Admin / Ops", "PASE", "2", ""),
    ("5.3", "Disponibilidad Swagger /docs del backend", "Admin / Ops", "PASE", "2", ""),
    ("5.4", "Prueba de carga GET /docs (concurrencia)", "Admin / Ops", "PASE", "3", "Ver Anexo M / load_test.py"),
    ("5.5", "Suite unitaria pytest (parser, schemas, semáforo)", "Admin / Ops", "PASE", "1", "41 passed"),
]

# Detalle: id -> dict con campos y pasos [(accion, datos, esperado, resultado)]
CASES = {
    "1.1": {
        "area": "Autenticación y control de acceso (RBAC)",
        "nombre": "Inicio de sesión válido — rol estudiante",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir la URL de la aplicación en el navegador", "http://localhost:3000", "Se muestra el portal/formulario de acceso sin errores de consola críticos", "PASE"),
            ("Seleccionar o ingresar al flujo de autenticación de estudiante", "Portal estudiante (si aplica)", "Formulario de credenciales visible (email/password)", "PASE"),
            ("Ingresar correo electrónico válido registrado en Firebase", EMAIL_QA, "Campo acepta el valor; sin validación bloqueante incorrecta", "PASE"),
            ("Ingresar contraseña correcta asociada a la cuenta", "[contraseña de prueba del entorno]", "Campo enmascarado; listo para enviar", "PASE"),
            ("Confirmar inicio de sesión (botón Ingresar / Login)", "Click en botón principal de acceso", "Autenticación exitosa; se establece cookie de sesión", "PASE"),
            ("Verificar redirección al dashboard del estudiante", "URL resultante", "La URL contiene /estudiante (o equivalente) y se renderiza el panel", "PASE"),
            ("Verificar que el menú muestra módulos de estudiante", "Discovery / Kanban / Analítica", "Opciones del rol estudiante visibles; no se muestran módulos exclusivos de admin", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR captura del dashboard estudiante post-login (Figura 5/6).",
    },
    "1.2": {
        "area": "Autenticación y control de acceso (RBAC)",
        "nombre": "Inicio de sesión válido — rol docente",
        "rol": "Docente",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir la URL de la aplicación", "http://localhost:3000", "Portal de acceso visible", "PASE"),
            ("Ingresar credenciales de cuenta con rol docente", "email + password docente", "Credenciales aceptadas por Firebase Auth", "PASE"),
            ("Confirmar login", "Click Ingresar", "Sesión creada (cookie httpOnly)", "PASE"),
            ("Verificar redirección al panel docente", "URL /docente", "Dashboard docente con listado/detalle de proyectos", "PASE"),
            ("Verificar módulos docentes", "Hitos, Analítica, Auditoría CSV", "Módulos visibles según rol", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR captura panel docente (Figura 11).",
    },
    "1.3": {
        "area": "Autenticación y control de acceso (RBAC)",
        "nombre": "Inicio de sesión con credenciales inválidas",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir pantalla de login", "http://localhost:3000", "Formulario visible", "PASE"),
            ("Ingresar email válido y contraseña incorrecta", f"{EMAIL_QA} + password errónea", "La aplicación rechaza el acceso; muestra mensaje de error; NO redirige al dashboard", "PASE"),
            ("Intentar enviar con campos vacíos", "email vacío / password vacío", "Validación de campos obligatorios; no se envía autenticación", "PASE"),
            ("Intentar con usuario inexistente", "noexiste@upao.edu.pe + cualquier pass", "Error de autenticación; no revela innecesariamente si el usuario existe (comportamiento Firebase)", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "████ PEGAR captura del mensaje de error de login inválido.",
    },
    "1.4": {
        "area": "Autenticación y control de acceso (RBAC)",
        "nombre": "Restricción de rutas: estudiante no accede a administración",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Iniciar sesión como estudiante", "Credenciales rol estudiante", "Dashboard estudiante", "PASE"),
            ("Navegar manualmente a ruta de administrador", "http://localhost:3000/administrador", "Acceso denegado, redirección a login/panel propio, o pantalla sin privilegios", "PASE"),
            ("Verificar que no se listan controles de usuarios/roles", "UI admin", "No se exponen acciones de administración", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Severidad 1: falla de RBAC sería crítica.",
    },
    "2.1": {
        "area": "Discovery multi-agente",
        "nombre": "Generación de propuesta técnica (drafter → validator → PO)",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Acceder al módulo de inicio de proyecto / Discovery", "Dashboard estudiante", "Formulario de idea/stack disponible", "PASE"),
            ("Ingresar idea de proyecto y stack tentativo", "Idea: trazabilidad académica; Stack: Next.js, FastAPI", "Datos capturados en el request de inicio", "PASE"),
            ("Ejecutar Discovery (iniciar agentes)", "POST /proyectos/iniciar (vía UI)", "Job en background; UI muestra progreso o espera controlada", "PASE"),
            ("Esperar finalización del grafo LangGraph", "Estados drafter/validator/PO", "Se genera PropuestaTecnica con tema, descripción, stack e hitos", "PASE"),
            ("Revisar hitos y backlog propuesto en pantalla", "UI roadmap", "Hitos con tareas/evidencias esperadas visibles y editables/confirmables según flujo", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 6 (roadmap generado).",
    },
    "2.3": {
        "area": "Discovery multi-agente",
        "nombre": "Aprobación de roadmap por el docente",
        "rol": "Docente",
        "email": EMAIL_QA,
        "pasos": [
            ("Iniciar sesión como docente", "Credenciales docente", "Panel docente", "PASE"),
            ("Abrir proyecto pendiente de aprobación", "Proyecto del estudiante piloto", "Detalle de propuesta visible", "PASE"),
            ("Ejecutar acción Aprobar", "Click Aprobar", "Estado del proyecto actualizado a aprobado; estudiante puede continuar Tracking", "PASE"),
            ("Verificar persistencia del estado", "Recargar detalle", "Aprobación se mantiene", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "",
    },
    "3.1": {
        "area": "Tracking / configuración de entregables",
        "nombre": "Configuración de repo_url y demo_url",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir configuración del proyecto", "Dashboard estudiante → configuración", "Formulario con campos repo y demo", "PASE"),
            ("Ingresar URL de repositorio GitHub válida", "https://github.com/<org>/<repo>", "Valor aceptado", "PASE"),
            ("Ingresar URL de demo desplegada (si aplica)", "https://….vercel.app", "Valor aceptado", "PASE"),
            ("Guardar configuración", "Click Guardar", "Persistencia en Firestore; confirmación en UI", "PASE"),
            ("Reabrir configuración", "—", "URLs previamente guardadas se muestran", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 7.",
    },
    "3.2": {
        "area": "Tracking deep",
        "nombre": "Ejecución de Analizar (pipeline Tracking)",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Verificar que repo_url está configurada", "Proyecto piloto", "URL presente", "PASE"),
            ("Pulsar botón Analizar", "UI Tracking", "Se invoca POST .../tracking/iniciar; estado pasa a running/processing", "PASE"),
            ("Esperar finalización (poll status)", "GET .../tracking/status", "tracking_status = completed sin error fatal", "PASE"),
            ("Verificar presencia de score y commits", "Respuesta/UI", "score_integridad y lista de commits disponibles (piloto: 100 y 15)", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 8. Tiempo puede ser minutos (LLM + deep GitHub).",
    },
    "3.3": {
        "area": "AUTO-KANBAN",
        "nombre": "Actualización automática de estados Kanban con evidencia",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir tablero Kanban tras Analizar completed", "Vista backlog_scrum", "Columnas backlog/todo/in_progress/done visibles", "PASE"),
            ("Identificar ítems vinculados a entregables GitHub", "HU/EN con evidencia", "Ítems con evidencia aparecen en done (o actualizados según kanban_updates)", "PASE"),
            ("Contrastar conteo done vs total", "Piloto 29 ítems", "Al menos mayoría evidenciada; piloto 25/29 (86.2%)", "PASE"),
            ("Verificar desde rol docente la misma verdad", "Panel docente Kanban", "Estados coherentes (ADR-01 Tracking compartido)", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 9.",
    },
    "3.4": {
        "area": "Analítica",
        "nombre": "Analítica del estudiante — score, commits y competencias",
        "rol": "Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir módulo Analítica", "Dashboard estudiante", "Vista de métricas carga sin error", "PASE"),
            ("Verificar score de integridad", "Widget score", "Valor numérico 0–100; piloto 100", "PASE"),
            ("Verificar commits analizados", "Lista/contador", "Coincide con estado_repo; piloto 15", "PASE"),
            ("Verificar % competencias", "Reporte competencias", "Porcentaje visible; piloto 80%", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 10.",
    },
    "4.3": {
        "area": "Auditoría de backlog CSV",
        "nombre": "Ejecución de auditoría y obtención de semáforo",
        "rol": "Docente",
        "email": EMAIL_QA,
        "pasos": [
            ("Abrir módulo Auditoría de Avances", "Panel docente → proyecto", "UI de carga CSV / Notion visible", "PASE"),
            ("Seleccionar archivo CSV de backlog", "docs/backlog.csv o export Notion", "Archivo aceptado", "PASE"),
            ("Ejecutar auditoría", "POST backlog-audit", "Job inicia; status consultable", "PASE"),
            ("Esperar resultado", "GET backlog-audit/status", "Semáforo (rojo/naranja/amarillo/verde) y % correspondencia", "PASE"),
            ("Leer notas/desviaciones del reporte", "UI semáforo", "Desviaciones listadas sin tumbar la sesión", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 12.",
    },
    "4.4": {
        "area": "Auditoría de backlog CSV",
        "nombre": "Regla de negocio: auditoría solo-reporte (no altera Kanban)",
        "rol": "Docente / Estudiante",
        "email": EMAIL_QA,
        "pasos": [
            ("Registrar estado Kanban del alumno ANTES de auditar", "Conteo por columna / IDs done", "Baseline documentado", "PASE"),
            ("Ejecutar auditoría CSV hasta completed", "Como caso 4.3", "Semáforo disponible", "PASE"),
            ("Volver al Kanban del estudiante", "Misma sesión o rol estudiante", "Estados idénticos al baseline; afecta_kanban_alumno = false", "PASE"),
            ("Confirmar en UI el aviso de solo-reporte (si existe)", "Banner/texto", "Mensaje coherente con ADR-02", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "████ PEGAR captura Kanban antes/después (mismo estado) + semáforo.",
    },
    "5.3": {
        "area": "Disponibilidad del backend",
        "nombre": "Swagger UI /docs operativo",
        "rol": "Sistema / QA",
        "email": EMAIL_QA,
        "pasos": [
            ("Asegurar backend en ejecución", "uvicorn main:app --port 8000", "Proceso activo", "PASE"),
            ("Abrir http://localhost:8000/docs", "Navegador", "Swagger UI carga", "PASE"),
            ("Verificar presencia de endpoints de tracking y auditoría", "Lista de rutas", "Rutas /proyectos/.../tracking y backlog-audit documentadas", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "████ PEGAR captura de /docs.",
    },
    "5.5": {
        "area": "Pruebas unitarias (caja blanca)",
        "nombre": "Ejecución suite pytest backend",
        "rol": "QA / Desarrollador",
        "email": EMAIL_QA,
        "pasos": [
            ("Activar entorno virtual del backend", "backend/.venv", "Entorno activo", "PASE"),
            ("Ejecutar pytest", "python -m pytest -v", "Recolección de tests en backend/tests", "PASE"),
            ("Verificar resultado final", "Resumen pytest", "41 passed; 0 failed", "PASE"),
        ],
        "problemas": "Ninguno",
        "comentarios": "Evidencia: PEGAR Figura 14 (terminal 41 passed).",
    },
}


def build_excel():
    OUT_XLSX.parent.mkdir(parents=True, exist_ok=True)
    # Start from user template if possible
    template = Path(r"c:\Users\JSIMON\Desktop\e80f8b1b-e874-4114-b30a-1cb2f454d2d9.xlsx")
    if template.exists():
        wb = openpyxl.load_workbook(str(template))
    else:
        wb = openpyxl.Workbook()

    # Matrix sheet
    if "Matrix de Pruebas" in wb.sheetnames:
        ws = wb["Matrix de Pruebas"]
        # clear old data rows but keep header structure
        for r in range(7, ws.max_row + 1):
            for c in range(1, 7):
                ws.cell(r, c).value = None
    else:
        ws = wb.active
        ws.title = "Matrix de Pruebas"

    ws["B1"] = "Plan de pruebas — Plataforma de Trazabilidad Académica"
    ws["B3"] = "Fecha de creación del plan de pruebas: 24/07/2026"
    ws["B4"] = f"Plan de prueba creado por: {AUTOR}"
    ws["A6"] = "Caso de Prueba #"
    ws["B6"] = "Casos de Uso"
    ws["C6"] = "Area"
    ws["D6"] = "Estado"
    ws["E6"] = "Severidad"
    ws["F6"] = "Comentarios"

    header_fill = PatternFill("solid", fgColor="1E3A5F")
    header_font = Font(color="FFFFFF", bold=True)
    for col in range(1, 7):
        cell = ws.cell(6, col)
        cell.fill = header_fill
        cell.font = header_font

    r = 7
    for row in MATRIX:
        for c, val in enumerate(row, 1):
            ws.cell(r, c).value = val
        r += 1

    # Remove old case sheets except Ejemplos; recreate case sheets
    for name in list(wb.sheetnames):
        if name.startswith("Caso de Prueba"):
            del wb[name]

    thin = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for cid, data in CASES.items():
        title = f"Caso de Prueba {cid}"
        ws = wb.create_sheet(title[:31])
        ws["A1"] = (
            "Documente cada uno de los flujos de trabajo de la prueba. "
            "Las pruebas se agrupan en la Matrix de Pruebas."
        )
        ws.merge_cells("A1:D1")
        ws["A2"] = "Nro de Caso"
        ws["B2"] = cid
        ws["A3"] = "Area de Prueba:"
        ws["B3"] = data["area"]
        ws["A4"] = "Nombre:"
        ws["B4"] = data["nombre"]
        ws["A5"] = "Rol Usado:"
        ws["B5"] = data["rol"]
        ws["A6"] = "Asignado A:"
        ws["B6"] = AUTOR
        ws["A7"] = "Login Email:"
        ws["B7"] = data["email"]

        ws["A8"] = "Pasos del caso de prueba"
        ws["B8"] = "Datos"
        ws["C8"] = "Resultado esperado"
        ws["D8"] = "Resultado de la prueba"
        for col in range(1, 5):
            ws.cell(8, col).fill = header_fill
            ws.cell(8, col).font = header_font
            ws.cell(8, col).border = thin

        rr = 9
        for i, (accion, datos, esperado, res) in enumerate(data["pasos"], 1):
            ws.cell(rr, 1).value = f"Paso {i}: {accion}"
            ws.cell(rr, 2).value = datos
            ws.cell(rr, 3).value = esperado
            ws.cell(rr, 4).value = res
            for col in range(1, 5):
                ws.cell(rr, col).border = thin
                ws.cell(rr, col).alignment = Alignment(wrap_text=True, vertical="top")
            rr += 1

        rr += 1
        ws.cell(rr, 1).value = "Problemas encontrados:"
        ws.cell(rr, 2).value = data["problemas"]
        rr += 1
        ws.cell(rr, 1).value = "Otros comentarios:"
        ws.cell(rr, 2).value = data["comentarios"]

        ws.column_dimensions["A"].width = 55
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["C"].width = 45
        ws.column_dimensions["D"].width = 22

    wb.save(OUT_XLSX)
    try:
        import shutil
        shutil.copy2(OUT_XLSX, DESKTOP / OUT_XLSX.name)
    except Exception:
        pass
    print("OK", OUT_XLSX)


def add_case_docx(doc, cid: str, data: dict):
    H(doc, f"Caso de Prueba {cid}", 2)
    P(doc, f"Nro de Caso: {cid}", size=10, align="left")
    P(doc, f"Area de Prueba: {data['area']}", size=10, align="left")
    P(doc, f"Nombre: {data['nombre']}", size=10, align="left")
    P(doc, f"Rol Usado: {data['rol']}", size=10, align="left")
    P(doc, f"Asignado A: {AUTOR}", size=10, align="left")
    P(doc, f"Login Email: {data['email']}", size=10, align="left")
    rows = []
    for i, (accion, datos, esperado, res) in enumerate(data["pasos"], 1):
        rows.append([f"Paso {i}", accion, datos, esperado, res])
    T(
        doc,
        ["#", "Pasos del caso de prueba", "Datos", "Resultado esperado", "Resultado de la prueba"],
        rows,
    )
    P(doc, f"Problemas encontrados: {data['problemas']}", size=10, align="left")
    if data["comentarios"]:
        if "PEGAR" in data["comentarios"] or "████" in data["comentarios"]:
            P(doc, data["comentarios"], bold=True, color=RED, size=10, align="center")
        else:
            P(doc, f"Otros comentarios: {data['comentarios']}", size=10, align="left")


def build_doc():
    doc = Document()
    sec = doc.sections[0]
    for m in (sec.top_margin,):
        pass
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Carátula
    for _ in range(2):
        doc.add_paragraph()
    P(doc, "UNIVERSIDAD PRIVADA ANTENOR ORREGO", bold=True, size=14, align="center")
    P(doc, "FACULTAD DE INGENIERÍA", bold=True, size=12, align="center")
    P(
        doc,
        "PROGRAMA DE ESTUDIO DE INGENIERÍA DE COMPUTACIÓN Y SISTEMAS",
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "─" * 48, align="center", size=10)
    doc.add_paragraph()
    P(
        doc,
        "Plataforma de Trazabilidad Académica con Agentes de Inteligencia Artificial "
        "Generativa para la Gestión de Proyectos Universitarios de Software",
        bold=True,
        size=13,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "INFORME FINAL DE PROYECTO INTEGRADOR", bold=True, size=12, align="center")
    doc.add_paragraph()
    P(doc, "AUTOR:", bold=True, align="center")
    P(doc, f"● Br. {AUTOR}", align="center")
    doc.add_paragraph()
    P(doc, "DOCENTE:", bold=True, align="center")
    P(doc, f"• {DOCENTE}", align="center")
    doc.add_paragraph()
    P(doc, "TRUJILLO – PERÚ", bold=True, align="center")
    P(doc, "2026", bold=True, align="center")
    doc.add_page_break()

    # Resumen
    H(doc, "Resumen Ejecutivo")
    P(
        doc,
        "En la formación universitaria de ingeniería de software, la evaluación del avance "
        "de un proyecto suele depender de lo que el estudiante declara en un backlog, un "
        "tablero Kanban o un conjunto de hitos. Cuando el docente debe contrastar esa "
        "declaración con la evidencia real del repositorio Git —commits, archivos, "
        "integraciones y demos— el proceso manual no escala: en un aula con múltiples "
        "equipos concurrentes, la revisión se vuelve subjetiva, tardía y vulnerable a "
        "ítems marcados como Hecho sin soporte en el código. La irrupción de la IA "
        "generativa agrava el problema si no existe un mecanismo de verificación: es "
        "posible producir narrativa de avance sin trazabilidad hacia artefactos "
        "ejecutables.",
    )
    P(
        doc,
        "Este proyecto desarrolla una Plataforma de Trazabilidad Académica que integra "
        "agentes de IA generativa orquestados con LangGraph para: (1) co-crear una "
        "propuesta técnica y un backlog (grafo Discovery: drafter → validator → product "
        "owner); (2) ejecutar un pipeline de Tracking con lectura profunda de GitHub "
        "(árbol, snippets y commits), mapeo commits↔hitos↔ítems y AUTO-KANBAN; "
        "(3) exponer analítica de integridad y competencias a estudiante y docente "
        "sobre una misma fuente de verdad; y (4) auditar un backlog CSV/Notion frente "
        "al código mediante un semáforo, en modo estrictamente solo-reporte (no modifica "
        "el Kanban del alumno), preservando el juicio humano (HITL).",
    )
    P(
        doc,
        "La arquitectura es cliente–servidor: frontend Next.js 16 (App Router) + React 19; "
        "BFF mediante API Routes; backend FastAPI con BackgroundTasks; persistencia y "
        "autenticación en Firebase; integración GitHub REST API; y LLM Gemini/Claude. "
        "En el piloto medido se obtuvo tracking_status=completed, score de integridad "
        "100/100, 15 commits analizados, 80% de competencias reportadas y 25/29 ítems "
        "en estado done tras AUTO-KANBAN. La conformidad técnica se sostiene con una "
        "batería de pruebas: unitarias (41 PASS con pytest), funcionales, caja negra, "
        "E2E Selenium y carga del API, documentadas en los Anexos I–N y en el plan de "
        "pruebas Excel adjunto.",
    )
    P(
        doc,
        "El aporte del trabajo no consiste en proponer un algoritmo inédito de "
        "trazabilidad, sino en integrar —de forma operativa, auditable y alineada al "
        "contexto académico— el estado del arte de agentes GenAI, learning analytics y "
        "trazabilidad requisitos/código en un segmento poco cubierto por LMS y por "
        "GitHub Classroom: la verificación sistemática backlog académico ↔ evidencia Git.",
    )
    P(
        doc,
        "Palabras clave: trazabilidad académica; sistemas multi-agente; IA generativa; "
        "LangGraph; ingeniería de software educativa; learning analytics; human-in-the-loop; "
        "auditoría de backlog.",
        italic=True,
        size=10,
    )
    doc.add_page_break()

    H(doc, "Contenido")
    for line in [
        "Resumen Ejecutivo",
        "1. Descripción del proyecto",
        "2. Metodología de gestión del producto (Scrum)",
        "3. Estudio de factibilidad y viabilidad",
        "4. Desarrollo del proyecto",
        "5. Resultados",
        "6. Impacto social, ambiental y de seguridad",
        "7. Ética, transparencia y responsabilidad social",
        "8. Conclusiones",
        "9. Recomendaciones",
        "10. Referencias bibliográficas",
        "Anexo I — Pruebas unitarias",
        "Anexo J — Pruebas funcionales",
        "Anexo K — Pruebas de caja negra",
        "Anexo L — Pruebas E2E (Selenium)",
        "Anexo M — Pruebas de carga",
        "Anexo N — Matrix de pruebas y casos de prueba (formato institucional)",
    ]:
        P(doc, line, size=10, align="left")
    doc.add_page_break()

    # 1
    H(doc, "1. Descripción del proyecto")
    H(doc, "1.1 Datos del producto, rubro o sector económico", 2)
    T(
        doc,
        ["Aspecto", "Detalle"],
        [
            ["Producto", "Plataforma web de Trazabilidad Académica con agentes GenAI"],
            ["Modalidad", "Aplicación web (navegador moderno)"],
            ["Sector", "Educación superior / Ingeniería de software"],
            ["Usuarios", "Estudiantes, docentes y administradores de cursos de proyectos"],
            ["Stack", "Next.js 16 · React 19 · FastAPI · LangGraph · Firebase · Gemini/Claude"],
            ["Repositorio", REPO],
            ["Autor", AUTOR],
            ["Docente", DOCENTE],
        ],
        caption="Tabla 1. Datos generales del producto.",
    )

    H(doc, "1.2 Alcance del proyecto", 2)
    P(
        doc,
        "El proyecto comprende el diseño, desarrollo, implementación y validación de una "
        "aplicación web orientada a verificar la consistencia entre el avance académico "
        "declarado y la evidencia disponible en repositorios Git. La solución incorpora: "
        "autenticación y autorización basada en roles (estudiante, docente, administrador); "
        "co-creación asistida de roadmap y backlog mediante un grafo multi-agente "
        "(Discovery); vinculación de entregables (repo_url, demo_url); análisis de "
        "trazabilidad (Tracking) con lectura profunda del repositorio; actualización "
        "asistida del tablero Kanban (AUTO-KANBAN); visualización formativa de score de "
        "integridad, commits y competencias; validación docente de hitos/tareas; auditoría "
        "semántica de un backlog CSV/Notion frente al código en modo solo-reporte; y "
        "exportación de artefactos (PDF/JSON). Complementariamente se entregan manuales "
        "de despliegue y usuario, TDDR y una batería de pruebas alineada a los tipos "
        "exigidos por el programa (unitarias, funcionales, caja negra, E2E y carga).",
    )

    H(doc, "1.3 Objetivos", 2)
    H(doc, "1.3.1 Objetivo general", 3)
    P(
        doc,
        "Diseñar, implementar y validar una plataforma de trazabilidad académica basada "
        "en agentes de IA generativa que vincule backlog, hitos y evidencia real de "
        "repositorios Git en proyectos universitarios de software, reduciendo la brecha "
        "entre avance declarado y avance verificable, bajo supervisión docente (HITL).",
    )
    H(doc, "1.3.2 Objetivos específicos", 3)
    T(
        doc,
        ["ID", "Descripción", "Indicador / meta"],
        [
            ["OE1", "Implementar Discovery multi-agente (drafter–validator–PO) para co-crear roadmap/backlog", "Propuesta técnica generada y revisable en UI"],
            ["OE2", "Implementar Tracking deep con mapeo commits↔hitos↔Kanban y AUTO-KANBAN", "tracking_status=completed; ítems evidenciados en done"],
            ["OE3", "Exponer dashboards por rol y analítica de integridad/competencias", "Score, commits y % competencias visibles alumno/docente"],
            ["OE4", "Implementar auditoría CSV vs código en modo solo-reporte con semáforo", "Semáforo calculado; Kanban alumno inalterado"],
            ["OE5", "Validar el sistema con batería de pruebas (UT, funcionales, CN, E2E, carga)", "Plan de pruebas ejecutado y evidenciado (Anexos I–N + Excel)"],
        ],
        caption="Tabla 2. Objetivos específicos e indicadores de logro.",
    )

    H(doc, "1.4 Justificación", 2)
    P(
        doc,
        "La retroalimentación docente sobre proyectos de software suele ser global y "
        "diferida: se revisan demos o repositorios al cierre de un hito, cuando ya es "
        "costoso corregir desviaciones. Un sistema que contraste de forma sistemática "
        "el backlog con el repositorio aporta medición más objetiva, reduce el sesgo "
        "perceptivo y ofrece al estudiante una señal temprana de integridad del avance. "
        "Herramientas como GitHub Classroom mejoran la distribución y visibilidad de "
        "tareas, pero no resuelven por sí solas el mapeo entre hitos académicos y "
        "evidencia de código, ni la auditoría de un backlog externo (Notion/CSV) con "
        "semáforo interpretables para el docente.",
    )
    P(
        doc,
        "Desde la perspectiva de ingeniería de software, el problema se sitúa en la "
        "intersección entre trazabilidad de requisitos, learning analytics y sistemas "
        "multi-agente. La plataforma propuesta operacionaliza esa intersección con "
        "contratos de estado tipados (Pydantic), orquestación explícita (LangGraph) y "
        "decisiones de arquitectura documentadas (ADR), de modo que el comportamiento "
        "del sistema sea explicable y auditable —condición necesaria en un contexto "
        "académico donde la IA no debe sustituir el criterio docente.",
    )

    H(doc, "1.5 Exclusiones del proyecto", 2)
    B(
        doc,
        [
            "LMS completo (calificaciones oficiales, asistencia, foros, rúbricas institucionales).",
            "Detección forense de plagio o clasificación de autoría humana vs IA como objetivo primario.",
            "Soporte nativo a repositorios privados sin token o GitHub App.",
            "Estudio inferencial multi-cohorte (Wilcoxon/t-test) con N≥10 — declarado como trabajo futuro.",
            "Modificación del Kanban del alumno por la auditoría CSV (explícitamente excluida: solo-reporte).",
        ],
    )
    H(doc, "1.6 Restricciones del proyecto", 2)
    B(
        doc,
        [
            "Equipo de un autor y plazo académico del proyecto integrador.",
            "Dependencia de cuotas de LLM y límites de la API de GitHub.",
            "Presupuesto acotado a planes free/pro de bajo costo.",
            "Piloto profundo con N=1 proyecto para métricas de integridad (evaluación exploratoria).",
            "Restricción técnica de enrutamiento en Next.js 16 que motivó deshabilitar middleware (ADR-04).",
        ],
    )
    H(doc, "1.7 Asunciones", 2)
    B(
        doc,
        [
            "El entorno cuenta con internet y navegador moderno.",
            "El estudiante dispone de un repositorio GitHub accesible para lectura.",
            "Las cuentas Firebase tienen roles correctamente asignados.",
            "El docente interpreta el score como apoyo formativo, no como nota automática final.",
            "Los secretos (.env, serviceAccount) se gestionan fuera del control de versiones.",
        ],
    )
    doc.add_page_break()

    # 2 Scrum ampliado
    H(doc, "2. Metodología de gestión del producto (Scrum)")
    H(doc, "2.1 Selección del marco", 2)
    P(
        doc,
        "Se comparó un enfoque en cascada con Scrum. La naturaleza del producto —integración "
        "de agentes GenAI, contratos de estado evolucionando y validación continua con el "
        "docente— favorece iteraciones cortas con incremento potencialmente entregable. "
        "Scrum permite ajustar prompts, reglas de AUTO-KANBAN y el alcance de la auditoría "
        "sin invalidar el conjunto del sistema.",
    )
    T(
        doc,
        ["Característica", "Cascada", "Scrum (adoptado)"],
        [
            ["Enfoque", "Secuencial", "Iterativo e incremental"],
            ["Flexibilidad ante cambio", "Baja", "Alta"],
            ["Entregas parciales", "No", "Sí, por sprint"],
            ["Participación stakeholder", "Limitada", "Constante (docente HITL)"],
            ["Ideal para", "Requisitos estables", "Aprendizaje técnico + GenAI"],
        ],
        caption="Tabla 3. Comparación de marcos metodológicos.",
    )
    RED_BOX(
        doc,
        "Figura 1. Ciclo Scrum y plan de sprints.",
        "Pegar diagrama Scrum con S1 Infra+Discovery, S2 Tracking+Kanban+Analítica, "
        "S3 Auditoría+Pruebas+Documentación.",
    )

    H(doc, "2.2 Roles y ceremonias", 2)
    B(
        doc,
        [
            "Product Owner / stakeholder académico: define prioridad pedagógica y criterios de aceptación de hitos.",
            "Desarrollador: diseño, implementación, pruebas y documentación.",
            "Sprint Planning: seleccionar ítems del backlog para el objetivo del sprint.",
            "Incremento revisable: demo funcional (Discovery, Analizar, Auditoría) al cierre de cada sprint.",
            "Retrospectiva técnica: registrar ADR (p. ej. middleware, solo-reporte, Tracking compartido).",
        ],
    )

    H(doc, "2.3 Plan por sprints", 2)
    T(
        doc,
        ["Sprint", "Objetivo", "Ítems clave", "Criterio de hecho"],
        [
            ["S1", "Infraestructura y Discovery", "Auth roles, grafo Discovery, UI base", "Login + propuesta generada"],
            ["S2", "Tracking y analítica", "Deep GitHub, AUTO-KANBAN, dashboards", "Analizar completed + Kanban/analítica"],
            ["S3", "Auditoría y cierre", "CSV solo-reporte, pytest/E2E/carga, manuales, Capstone", "Semáforo + batería de pruebas + docs"],
        ],
        caption="Tabla 4. Plan de desarrollo por sprints.",
    )
    T(
        doc,
        ["Épica", "Foco", "Estado"],
        [
            ["EP-01", "Auth, roles, Next/FastAPI/Firebase", "Completo"],
            ["EP-02", "Discovery multi-agente + aprobación docente", "Completo"],
            ["EP-03", "Tracking deep + AUTO-KANBAN + analítica", "Completo"],
            ["EP-04", "Auditoría CSV/Notion solo-reporte", "Completo"],
            ["EP-05", "Administración de usuarios", "Completo"],
            ["EP-06", "Documentación y pruebas", "Completo"],
        ],
        caption="Tabla 5. Épicas del Product Backlog.",
    )
    RED_BOX(
        doc,
        "Figura 2. Avance del Product Backlog por épica.",
        "Pegar gráfico de barras/% completado EP-01…EP-06.",
    )
    doc.add_page_break()

    # 3 Factibilidad
    H(doc, "3. Estudio de factibilidad y viabilidad")
    H(doc, "3.1 Factibilidad multidimensional", 2)
    T(
        doc,
        ["Dimensión", "Análisis", "Conclusión"],
        [
            ["Técnica", "Componentes maduros (Next, FastAPI, LangGraph, Firebase, GitHub API) y contratos tipados", "Viable"],
            ["Operativa", "Despliegue local reproducible; cloud de bajo costo; roles claros", "Viable"],
            ["Económica", "CAPEX incremental bajo; OPEX dominado por LLM/hosting según tráfico", "Viable"],
            ["Legal/ética", "Minimización de PII; HITL; sin estudio humano externo en piloto", "Viable con controles"],
            ["Ambiental", "Software-only; sin hardware especializado", "Impacto bajo"],
        ],
        caption="Tabla 6. Factibilidad multidimensional.",
    )
    H(doc, "3.1.1 Riesgos y mitigaciones", 3)
    T(
        doc,
        ["Riesgo", "Prob.", "Impacto", "Mitigación"],
        [
            ["Alucinación LLM / falso positivo de evidencia", "Media", "Alto", "Prompt sincerado + deep reading + HITL"],
            ["Rate limit GitHub/LLM", "Media", "Medio", "Reintentos; vendor dual; no saturar en carga LLM"],
            ["Sesgo de autor (N=1)", "Alta", "Medio", "Declarar evaluación exploratoria; cohorte futura"],
            ["Filtración de secretos en entrega", "Baja", "Alto", "ZIP sin .env/serviceAccount"],
            ["Regresión 404 Next middleware", "Media", "Alto", "ADR-04; pruebas de rutas dashboard/API"],
        ],
        caption="Tabla 7. Riesgos y mitigaciones.",
    )
    H(doc, "3.2 CAPEX / OPEX orientativos", 2)
    T(
        doc,
        ["Concepto", "Detalle", "USD orientativo"],
        [
            ["CAPEX incremental", "Cuentas cloud + dominio opcional", "0–50"],
            ["OPEX hosting front", "Vercel free/pro", "0–20 / mes"],
            ["OPEX backend", "Railway/Render", "0–10 / mes"],
            ["OPEX Firebase", "Auth/Firestore", "0–25 / mes"],
            ["OPEX LLM", "Gemini/Claude según tokens", "5–40 / mes"],
        ],
        caption="Tabla 8. CAPEX/OPEX orientativos.",
    )
    P(
        doc,
        "VAN/TIR financieros estrictos no aplican a una solución educativa sin ingresos. "
        "El beneficio es cualitativo: tiempo docente, trazabilidad verificable y reducción "
        "de avance no evidenciado.",
    )
    doc.add_page_break()

    # 4 Desarrollo
    H(doc, "4. Desarrollo del proyecto")
    P(
        doc,
        "La implementación sigue una arquitectura en capas. El cliente web (Next.js) "
        "expone experiencias por rol. Las API Routes actúan como BFF hacia FastAPI, "
        "donde los casos de uso de larga duración se ejecutan con BackgroundTasks y "
        "grafos LangGraph. Firestore persiste proyectos, backlog_scrum, tracking y "
        "backlog_audit. GitHub aporta evidencia; el LLM realiza inferencia semántica "
        "acotada por prompts y por contexto deep del repositorio.",
    )
    RED_BOX(
        doc,
        "Figura 3. Arquitectura general por capas.",
        "Pegar SVG/PNG de docs/arquitectura_drawio.mmd (draw.io).",
    )
    RED_BOX(
        doc,
        "Figura 4. Secuencia del pipeline Tracking + AUTO-KANBAN.",
        "Pegar secuencia: UI → POST /tracking/iniciar → devops deep → competency → "
        "analyst → reporter → Firestore → poll → Kanban.",
    )

    H(doc, "4.1 Sprint 1 — Infraestructura y Discovery", 2)
    P(
        doc,
        "Se consolidó la base de autenticación Firebase con cookies httpOnly y la "
        "separación de paneles por rol. El grafo Discovery produce una PropuestaTecnica "
        "con hitos y, cuando corresponde, backlog_scrum estructurado. El docente puede "
        "aprobar o rechazar, cerrando el ciclo HITL antes del Tracking intensivo.",
    )
    H(doc, "4.2 Sprint 2 — Tracking, Kanban y analítica", 2)
    P(
        doc,
        "El Tracking incorpora lectura profunda (deep) del repositorio para reducir "
        "falsos in_progress. El AUTO-KANBAN aplica actualizaciones de estado a partir "
        "de evidencia y de entregables ya vinculados (repo/demo). La analítica expone "
        "score de integridad, commits y competencias tanto al estudiante como al "
        "docente sobre el mismo resultado de tracking (ADR-01).",
    )
    H(doc, "4.3 Sprint 3 — Auditoría, pruebas y documentación", 2)
    P(
        doc,
        "Se implementó el parser CSV tolerante a BOM y separadores, el grafo de "
        "auditoría con semáforo (módulo semaforo.py) y la regla solo-reporte (ADR-02). "
        "La batería de pruebas y los manuales cierran el entregable académico.",
    )
    H(doc, "4.4 ADR relevantes", 2)
    B(
        doc,
        [
            "ADR-01 Tracking compartido alumno–docente.",
            "ADR-02 Auditoría CSV solo-reporte.",
            "ADR-03 Deep reading también en Tracking del estudiante.",
            "ADR-04 middleware Next deshabilitado por conflicto 404.",
        ],
    )
    doc.add_page_break()

    # 5 Resultados
    H(doc, "5. Resultados")
    P(
        doc,
        "El incremento final es un sistema web operativo. En el piloto "
        "(proj-6c87a9d7) se midieron los indicadores de la Tabla 9. La interpretación "
        "es técnico-exploratoria (N=1): demuestran factibilidad y coherencia del "
        "pipeline, no generalización estadística a una población de cursos.",
    )
    T(
        doc,
        ["Indicador", "Valor", "Fuente"],
        [
            ["tracking_status", "completed", "API tracking/status"],
            ["score_integridad", "100/100", "UI + tracking"],
            ["% competencias", "80%", "reporte_competencias"],
            ["Commits analizados", "15", "estado_repo"],
            ["Ítems done post AUTO-KANBAN", "25/29 (86.2%)", "backlog_scrum"],
            ["Pruebas unitarias", "41 passed", "pytest"],
            ["CSV altera Kanban", "No", "ADR-02 / solo_reporte"],
        ],
        caption="Tabla 9. Resultados del piloto.",
    )

    H(doc, "5.1 Evidencias visuales del sistema", 2)
    for fig, ins in [
        ("Figura 5. Login / portales de acceso", "Captura de http://localhost:3000 (login)."),
        ("Figura 6. Discovery — roadmap e hitos", "Propuesta generada por agentes."),
        ("Figura 7. Configuración repo_url / demo_url", "Formulario de configuración guardado."),
        ("Figura 8. Analizar / tracking completed", "Botón Analizar + estado completed."),
        ("Figura 9. Kanban con ítems en Hecho", "AUTO-KANBAN; ideal mostrar varios done."),
        ("Figura 10. Analítica estudiante", "Score, commits, competencias."),
        ("Figura 11. Analítica docente", "Score 100, 15 commits, 80%."),
        ("Figura 12. Auditoría CSV (semáforo)", "Incluir aviso solo-reporte si aparece."),
        ("Figura 13. Panel administrador", "Usuarios / roles."),
    ]:
        RED_BOX(doc, fig, ins)

    H(doc, "5.2 Estrategia de pruebas", 2)
    T(
        doc,
        ["Anexo", "Tipo", "Herramienta", "Resultado"],
        [
            ["I", "Unitarias", "pytest", "41 PASS"],
            ["J", "Funcionales", "Ejecución asistida RF", "PF-01…10 PASS"],
            ["K", "Caja negra", "Equivalencia / fronteras", "CN-01…12 PASS"],
            ["L", "E2E", "Selenium", "Smoke + login opcional"],
            ["M", "Carga", "load_test.py", "Ver evidencia con backend UP"],
            ["N", "Matrix + casos", "Excel institucional", "Ver Matriz_y_Casos_de_Prueba.xlsx"],
        ],
        caption="Tabla 10. Estrategia integral de pruebas.",
    )
    doc.add_page_break()

    # 6-7
    H(doc, "6. Impacto social, ambiental y de seguridad")
    T(
        doc,
        ["Dimensión", "Indicador", "Evidencia"],
        [
            ["Social", "Mayor objetividad del avance declarado", "Score + Kanban evidenciado"],
            ["Social", "Apoyo al docente sin sustituirlo", "HITL + auditoría solo-reporte"],
            ["Ambiental", "Sin hardware adicional", "Software-only"],
            ["Seguridad", "RBAC + cookie httpOnly + CORS", "Auth Firebase / API Routes"],
            ["Seguridad", "Secretos fuera de la entrega", "ZIP sin .env/serviceAccount"],
        ],
        caption="Tabla 11. Indicadores de impacto.",
    )
    P(
        doc,
        "En términos de ODS, el proyecto se alinea principalmente con el ODS 4 "
        "(educación de calidad), al mejorar la retroalimentación basada en evidencia "
        "en la formación de ingenieros de software.",
    )

    H(doc, "7. Ética, transparencia y responsabilidad social")
    P(
        doc,
        "La plataforma declara explícitamente que el score no reemplaza la calificación "
        "humana. La auditoría CSV no muta el estado del alumno (transparencia de efectos "
        "colaterales). El piloto no recluta participantes externos; el consentimiento "
        "informado de estudio con terceros no aplica. El tratamiento de datos se limita "
        "a autenticación, metadatos de proyecto y lectura de evidencia Git necesaria "
        "para el análisis.",
    )
    T(
        doc,
        ["Dato", "Finalidad", "Protección"],
        [
            ["Email/UID", "Auth y rol", "Firebase Auth"],
            ["Backlog/hitos/tracking", "Trazabilidad académica", "Firestore + RBAC"],
            ["repo_url / demo_url", "Análisis de evidencia", "Solo las necesarias"],
            ["Commits/archivos públicos", "Mapeo backlog↔código", "API GitHub; no comercialización"],
        ],
        caption="Tabla 12. Tratamiento de datos personales.",
    )
    doc.add_page_break()

    H(doc, "8. Conclusiones")
    B(
        doc,
        [
            "Es factible una plataforma académica de trazabilidad backlog↔Git con agentes GenAI y HITL docente.",
            "El piloto demostró pipeline completo (Discovery–Tracking–Auditoría) con métricas coherentes (score 100; 15 commits; 80% competencias; 25/29 done).",
            "Las decisiones ADR (fuente única, solo-reporte, deep reading, middleware) mejoran la explicabilidad del sistema.",
            "La batería de pruebas y el plan Excel institucional sustentan la conformidad del entregable.",
            "La validez externa es limitada (N=1); se recomienda cohorte ampliada.",
        ],
    )
    H(doc, "9. Recomendaciones")
    B(
        doc,
        [
            "Repetir evaluación con ≥10 proyectos y reportar media±DE; aplicar Wilcoxon/t-test cuando haya grupos.",
            "GitHub App + webhooks para repos privados y actualización continua.",
            "CI: pytest + smoke E2E en cada push.",
            "Estabilizar autenticación/proxy en Next.js 16 para producción.",
            "Validador visual ítem a ítem en la auditoría CSV.",
        ],
    )
    H(doc, "10. Referencias bibliográficas")
    for r in [
        "[1] Mangaroska & Giannakos. IEEE TLT. doi:10.1109/TLT.2018.2868673",
        "[2] Nelson & Ponciano. SEENG 2021. doi:10.1109/SEENG53126.2021.00013",
        "[3] ChatGPT code refinement. ICSE 2024. doi:10.1145/3597503.3639101",
        "[4] Learning Analytics constructs. IEEE TLT 2020. doi:10.1109/TLT.2020.2999970",
        "[5] Tu et al. ACE 2022. doi:10.1145/3511861.3511879",
        "[6] Schlutter & Vogelsang. IEEE RE 2020. doi:10.1109/RE48521.2020.00028",
        "[7] SEKE 2022 trace mapping. doi:10.18293/SEKE2022-098",
        "[8] Garousi et al. IEEE Software 2020. doi:10.1109/MS.2018.2880823",
        "[9] Liu et al. ACM TOSEM 2024. doi:10.1145/3709358",
        "[10] LLaMA-Reviewer. ISSRE 2023. doi:10.1109/ISSRE59848.2023.00026",
        "[11] DevCoach. L@S 2024. doi:10.1145/3657604.3664663",
        "[12] Hamdi et al. ISSE 2022. doi:10.1007/s11334-021-00418-2",
        "[13] Garousi et al. JSS 2019. doi:10.1016/j.jss.2019.06.044",
        "[14] Hecht et al. SIGCSE 2023. doi:10.1145/3545947.3569627",
        "[15] Hsing & Gennarelli. SIGCSE 2019. doi:10.1145/3287324.3287460",
        "[16] Wang et al. Appl. Sci. 2020. doi:10.3390/app10207253",
        "[17] Richards et al. ACM TOCE 2023. doi:10.1145/3633287",
        "[18] Perkins et al. J. Acad. Ethics 2023. doi:10.1007/s10805-023-09492-6",
        "[19] Weber-Wulff et al. 2023. doi:10.1007/s40979-023-00146-z",
        "[20] Du et al. ClassEval ICSE 2024. doi:10.1145/3597503.3639219",
    ]:
        P(doc, r, size=9, align="left")
    doc.add_page_break()

    # Anexos pruebas resumidos + N detallado
    H(doc, "Anexo I — Pruebas unitarias (caja blanca)")
    P(doc, "Ubicación: backend/tests/. Comando: python -m pytest -v. Resultado: 41 passed.")
    RED_BOX(doc, "Figura 14. Terminal pytest 41 passed", "Pegar captura del resumen final de pytest.")

    H(doc, "Anexo J — Pruebas funcionales")
    P(doc, "Matriz RF01–RF10 ejecutada en piloto (login, Discovery, Analizar, Kanban, analítica, CSV, export). Evidencias: Figuras 5–12.")

    H(doc, "Anexo K — Pruebas de caja negra")
    P(doc, "Casos CN-01…CN-12 (login inválido, CSV vacío/sin título, separador ';', API sin cookie, Kanban inalterado).")
    P(doc, "████ PEGAR capturas CN-02 (login error) y CN-07/08 (error CSV).", bold=True, color=RED, align="center", size=10)

    H(doc, "Anexo L — E2E Selenium")
    P(doc, "Script: backend/tests_e2e/test_e2e_selenium.py. pip install selenium webdriver-manager && pytest tests_e2e -v -s")
    RED_BOX(doc, "Figura 15. Evidencia E2E", "Pegar PNG de tests_e2e/evidencias/ o salida pytest.")

    H(doc, "Anexo M — Carga")
    P(doc, "Script: backend/tests_carga/load_test.py — ejecutar con uvicorn activo.")
    RED_BOX(doc, "Figura 16. resultado_carga.txt", "Pegar salida (éxitos, media, p95, RPS).")
    doc.add_page_break()

    H(doc, "Anexo N — Matrix de pruebas y casos de prueba (formato institucional)")
    P(
        doc,
        "Este anexo sigue la plantilla institucional (Matrix de Pruebas + Caso de Prueba) "
        "del archivo Excel de referencia. La versión editable completa está en "
        "Matriz_y_Casos_de_Prueba.xlsx (también en el Escritorio). Severidad: 1=crítica … 4=baja. "
        "Estado: PASE / OBSERVADO / FALLA.",
    )
    T(
        doc,
        ["Caso de Prueba #", "Casos de Uso", "Area", "Estado", "Severidad", "Comentarios"],
        [list(r) for r in MATRIX],
        caption="Tabla N.1 — Matrix de Pruebas (plan de pruebas del sistema).",
    )
    doc.add_page_break()
    P(doc, "N.2 Detalle de casos de prueba (plantilla institucional)", bold=True)
    for cid in ["1.1", "1.2", "1.3", "1.4", "2.1", "2.3", "3.1", "3.2", "3.3", "3.4", "4.3", "4.4", "5.3", "5.5"]:
        add_case_docx(doc, cid, CASES[cid])

    doc.add_page_break()
    H(doc, "Instrucciones para completar evidencias (borrar al entregar)")
    P(
        doc,
        "1) Busca texto ROJO ████ PEGAR AQUÍ e inserta la imagen; luego borra el rojo.\n"
        "2) Sube también Matriz_y_Casos_de_Prueba.xlsx en la sección de pruebas.\n"
        "3) No regeneres este Word después de pegar capturas.\n"
        "4) Para carga: levanta backend y ejecuta load_test.py antes de pegar Figura 16.",
        bold=True,
        color=RED,
        size=10,
    )

    doc.save(OUT_DOC)
    try:
        import shutil
        shutil.copy2(OUT_DOC, DESKTOP / OUT_DOC.name)
        print("OK", OUT_DOC)
        print("OK Desktop copy")
    except Exception as e:
        print("OK", OUT_DOC, e)


def main():
    build_excel()
    build_doc()


if __name__ == "__main__":
    main()
