# -*- coding: utf-8 -*-
"""Actas de Conformidad al estilo institucional UPAO (firmas de entregables)."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

OUT_DIR = Path(__file__).resolve().parent / "entrega" / "Actas"
DESKTOP = Path(r"c:\Users\JSIMON\Desktop")

PROYECTO = (
    "Plataforma de Trazabilidad Académica con Agentes de Inteligencia Artificial "
    "Generativa para la Gestión de Proyectos Universitarios de Software"
)
FECHA = "24/07/2026"
FECHA_LARGA = "jueves 24 de julio de 2026"
CLIENTE = "Walter Cueva Chávez"
AUTOR = "Jean Marcos Meneses Simón"


def font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def P(doc, text, *, bold=False, italic=False, size=11, align="justify"):
    p = doc.add_paragraph()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(text)
    font(r, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(8)
    return p


def shade(cell):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), "D9E2F3")
    sh.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(sh)


def firma_table(doc, incluir_proyecto=True):
    rows = 2 if incluir_proyecto else 1
    table = doc.add_table(rows=1 + rows, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["", "Nombre", "Firma"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                font(r, size=10, bold=True)
        shade(table.rows[0].cells[i])
    table.rows[1].cells[0].text = "Representante del Cliente"
    table.rows[1].cells[1].text = CLIENTE
    table.rows[1].cells[2].text = ""
    if incluir_proyecto:
        table.rows[2].cells[0].text = "Representante del Proyecto"
        table.rows[2].cells[1].text = AUTOR
        table.rows[2].cells[2].text = ""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    font(r, size=10)
            # altura visual
            cell.paragraphs[0].paragraph_format.space_before = Pt(10)
            cell.paragraphs[0].paragraph_format.space_after = Pt(10)
    doc.add_paragraph()


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)
    return doc


def acta_documento(titulo_doc: str, filename: str, cuerpo_extra: str = ""):
    """Acta de Conformidad del Documento de X."""
    doc = new_doc()
    P(doc, f"Acta de Conformidad del Documento de {titulo_doc}", bold=True, size=16, align="center")
    P(doc, f"“{PROYECTO}”", bold=True, size=12, align="center")
    doc.add_paragraph()
    P(doc, f"Fecha de entrega del acta: {FECHA}", size=11, align="left")
    doc.add_paragraph()
    P(
        doc,
        f"En la fecha {FECHA_LARGA}, se tiene constatado que el documento de "
        f"{titulo_doc} correspondiente al proyecto “{PROYECTO}”, ha sido aceptado "
        f"y aprobado por el cliente, cumpliendo con los términos de referencia en "
        f"el alcance del proyecto y las especificaciones dentro del Project Charter; "
        f"por lo que se emite la presente ACTA DE CONFORMIDAD DEL DOCUMENTO.",
    )
    if cuerpo_extra:
        P(doc, cuerpo_extra)
    doc.add_paragraph()
    firma_table(doc, incluir_proyecto=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    try:
        import shutil
        shutil.copy2(path, DESKTOP / filename)
    except Exception:
        pass
    print("OK", path)
    return path


def acta_proyecto():
    doc = new_doc()
    P(doc, "Acta de Conformidad del Proyecto", bold=True, size=16, align="center")
    P(doc, f"“{PROYECTO}”", bold=True, size=12, align="center")
    doc.add_paragraph()
    P(doc, f"Fecha de entrega del acta: {FECHA}", align="left")
    doc.add_paragraph()
    P(
        doc,
        f"En la fecha {FECHA_LARGA}, se tiene constatado que los resultados obtenidos "
        f"correspondientes al proyecto “{PROYECTO}”, han sido aceptados y aprobados "
        f"por el cliente, cumpliendo con los términos de referencia en el alcance del "
        f"proyecto y las especificaciones dentro del Project Charter; por lo que se "
        f"emite la presente ACTA DE CONFORMIDAD.",
    )
    doc.add_paragraph()
    P(doc, "Productos principales aceptados:", bold=True)
    for it in [
        "Plataforma web operativa (frontend Next.js + backend FastAPI/LangGraph).",
        "Módulos Discovery, Tracking + AUTO-KANBAN, Analítica y Auditoría CSV solo-reporte.",
        "Documentación técnica (TDDR / Informe Capstone), manuales y suite de pruebas.",
        "Repositorio: https://github.com/Jmmsimon/trazabilidad-taller-I",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(it), size=11)
    doc.add_paragraph()
    firma_table(doc, incluir_proyecto=True)
    path = OUT_DIR / "Acta_Conformidad_Proyecto.docx"
    doc.save(path)
    try:
        import shutil
        shutil.copy2(path, DESKTOP / path.name)
    except Exception:
        pass
    print("OK", path)


def acta_pruebas():
    doc = new_doc()
    P(doc, "Acta de Conformidad de pruebas", bold=True, size=16, align="center")
    P(doc, f"“{PROYECTO}”", bold=True, size=12, align="center")
    doc.add_paragraph()
    P(doc, f"Fecha de entrega del acta: {FECHA}", align="left")
    doc.add_paragraph()
    P(
        doc,
        f"En la fecha {FECHA_LARGA}, se tiene constatado que el documento de pruebas "
        f"correspondiente al proyecto “{PROYECTO}”, ha sido aceptado y aprobado por el "
        f"cliente, cumpliendo con los términos de referencia en el alcance del proyecto "
        f"y las especificaciones dentro del Project Charter; por lo que se emite la "
        f"presente ACTA DE CONFORMIDAD DEL DOCUMENTO.",
    )
    doc.add_paragraph()
    P(
        doc,
        "La batería de pruebas aceptada incluye los tipos exigidos por el programa:",
        bold=True,
    )
    for it in [
        "Pruebas unitarias (caja blanca) — pytest, 41 casos PASS (Anexo I del Informe Capstone).",
        "Pruebas funcionales — matriz RF / casos de uso (Anexo J).",
        "Pruebas de caja negra — equivalencia y fronteras (Anexo K).",
        "Pruebas end-to-end (E2E) — Selenium WebDriver (Anexo L).",
        "Pruebas de carga / rendimiento — load_test.py (Anexo M).",
        "Matrix de Pruebas y casos detallados — Matriz_y_Casos_de_Prueba.xlsx (Anexo N).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(it), size=11)
    doc.add_paragraph()
    P(
        doc,
        "Nota: esta acta certifica la conformidad del paquete documental y de "
        "ejecución de pruebas del sistema; no sustituye la firma de cada caso "
        "individual (1.1, 1.2, …), los cuales constan en la Matrix de Pruebas con "
        "estado PASE/OBSERVADO/FALLA.",
        italic=True,
        size=10,
    )
    doc.add_paragraph()
    firma_table(doc, incluir_proyecto=True)
    path = OUT_DIR / "Acta_Conformidad_Pruebas.docx"
    doc.save(path)
    try:
        import shutil
        shutil.copy2(path, DESKTOP / path.name)
    except Exception:
        pass
    print("OK", path)


def acta_indicadores():
    doc = new_doc()
    P(doc, "Acta de Conformidad de indicadores de éxito", bold=True, size=16, align="center")
    P(doc, f"“{PROYECTO}”", bold=True, size=12, align="center")
    doc.add_paragraph()
    P(doc, f"Fecha de entrega del acta: {FECHA}", align="left")
    doc.add_paragraph()
    P(
        doc,
        f"En la fecha {FECHA_LARGA}, se tiene constatado que los indicadores mostrados "
        f"a continuación, pertenecientes al proyecto “{PROYECTO}”, han sido aceptados "
        f"y aprobados por el cliente, cumpliendo con los términos de referencia en el "
        f"alcance del proyecto y las especificaciones dentro del Project Charter; "
        f"por lo que se emite la presente ACTA DE CONFORMIDAD.",
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Identificador de Indicador"
    table.rows[0].cells[1].text = "Descripción"
    for cell in table.rows[0].cells:
        shade(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                font(r, size=10, bold=True)

    indicadores = [
        (
            "IND-01",
            "Arquitectura híbrida implementada: frontend Next.js 16, backend FastAPI, "
            "orquestación LangGraph (Discovery, Tracking, Backlog Audit), Firebase Auth/Firestore "
            "e integración GitHub + LLM (Gemini/Claude).",
        ),
        (
            "IND-02",
            "Pipeline de Tracking operativo en piloto: tracking_status=completed, "
            "score de integridad 100/100, 15 commits analizados y 80% de competencias reportadas.",
        ),
        (
            "IND-03",
            "AUTO-KANBAN evidenciado: 25/29 ítems (86.2%) en estado done tras el análisis, "
            "con fuente de verdad compartida alumno–docente (ADR-01).",
        ),
        (
            "IND-04",
            "Auditoría CSV vs código en modo solo-reporte (ADR-02): genera semáforo sin "
            "alterar el Kanban del estudiante.",
        ),
        (
            "IND-05",
            "Batería de pruebas ejecutada y documentada: unitarias (41 PASS), funcionales, "
            "caja negra, E2E Selenium y carga; Matrix de Pruebas institucional adjunta.",
        ),
        (
            "IND-06",
            "Entregables de cierre: Informe Capstone, manual de despliegue, manual de usuario, "
            "código comprimido y actas de conformidad.",
        ),
    ]
    for ide, desc in indicadores:
        row = table.add_row().cells
        row[0].text = ide
        row[1].text = desc
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    font(r, size=9)

    doc.add_paragraph()
    firma_table(doc, incluir_proyecto=True)
    path = OUT_DIR / "Acta_Conformidad_Indicadores_Exito.docx"
    doc.save(path)
    try:
        import shutil
        shutil.copy2(path, DESKTOP / path.name)
    except Exception:
        pass
    print("OK", path)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    acta_documento(
        "Project Charter",
        "Acta_Conformidad_Project_Charter.docx",
        "El Project Charter / alcance del producto define el problema de trazabilidad "
        "académica backlog↔Git, objetivos OE1–OE5, exclusiones, restricciones y supuestos "
        "documentados en el Informe Capstone (Sección 1).",
    )
    acta_proyecto()
    acta_pruebas()
    acta_documento(
        "manual de usuario",
        "Acta_Conformidad_Manual_Usuario.docx",
        "Documento de referencia: Manuales/Guia_de_Usuario/Manual_Usuario.docx "
        "(flujos por rol estudiante, docente y administrador).",
    )
    acta_documento(
        "manual de despliegue",
        "Acta_Conformidad_Manual_Despliegue.docx",
        "Documento de referencia: Manuales/Despliegue/Manual_Despliegue.docx "
        "(requisitos, variables de entorno, arranque local y producción).",
    )
    acta_indicadores()
    # README
    (OUT_DIR / "README_ACTAS.txt").write_text(
        """ACTAS DE CONFORMIDAD — qué son y qué no son

NO son un acta por cada caso de prueba (1.1, 1.2, 3.2…).
SÍ son firmas de aceptación de ENTREGABLES grandes:

1. Acta_Conformidad_Project_Charter.docx
2. Acta_Conformidad_Proyecto.docx
3. Acta_Conformidad_Pruebas.docx          ← cubre unitarias + E2E + funcionales + caja negra + carga
4. Acta_Conformidad_Manual_Usuario.docx
5. Acta_Conformidad_Manual_Despliegue.docx
6. Acta_Conformidad_Indicadores_Exito.docx

Los casos de prueba individuales van en:
  Matriz_y_Casos_de_Prueba.xlsx  (estado PASE en cada fila)
  Informe Capstone — Anexo N

E2E está en el Capstone: Anexo L — Pruebas End-to-End (Selenium).
""",
        encoding="utf-8",
    )
    print("DIR", OUT_DIR)


if __name__ == "__main__":
    main()
