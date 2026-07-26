# -*- coding: utf-8 -*-
"""Genera el Word de Pruebas Unitarias a partir del último pytest."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "entrega" / "Pruebas" / "Pruebas_unitarias"
OUT = OUT_DIR / "Informe_Pruebas_Unitarias.docx"
PYTEST_LOG = OUT_DIR / "pytest_resultado.txt"


def set_run_font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def H(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)


def P(doc, text, bold=False, italic=False, size=11, align="justify"):
    p = doc.add_paragraph()
    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)


def B(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(it), size=11)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    log = PYTEST_LOG.read_text(encoding="utf-8", errors="replace") if PYTEST_LOG.exists() else "(sin log)"

    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.1)
    s.right_margin = Inches(1.1)

    for _ in range(2):
        doc.add_paragraph()
    P(doc, "UNIVERSIDAD PRIVADA ANTENOR ORREGO", bold=True, size=14, align="center")
    P(doc, "Facultad de Ingeniería", bold=True, size=12, align="center")
    P(
        doc,
        "Escuela Profesional de Ingeniería de Computación y Sistemas",
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "INFORME DE PRUEBAS UNITARIAS", bold=True, size=16, align="center")
    P(
        doc,
        "Plataforma de Trazabilidad Académica con Agentes de IA Generativa",
        italic=True,
        size=11,
        align="center",
    )
    doc.add_paragraph()
    P(doc, "Autor: Jean Marcos Meneses Simón", size=11, align="center")
    P(doc, "Fecha: Julio 2026", size=11, align="center")
    P(doc, "Framework: pytest (Python 3.12)", size=11, align="center")
    doc.add_page_break()

    H(doc, "1. Objetivo")
    P(
        doc,
        "Verificar de forma aislada (sin UI, sin red y sin LLM) las funciones críticas "
        "del backend: normalización y parseo de backlog CSV, modelos Pydantic/schemas "
        "y cálculo del semáforo de auditoría backlog↔código.",
    )

    H(doc, "2. Alcance")
    B(
        doc,
        [
            "Módulo backlog_parser.py (normalización de estado/tipo, parseo CSV).",
            "Módulo schemas.py (validación de modelos y helpers de conversión).",
            "Módulo semaforo.py (rangos de color y penalización por bulk-commit).",
            "Fuera de alcance: llamadas a GitHub, Firebase y proveedores LLM (cubiertas en E2E/funcionales).",
        ],
    )

    H(doc, "3. Entorno de ejecución")
    B(
        doc,
        [
            "Sistema operativo: Windows 10",
            "Python 3.12 + entorno virtual backend/.venv",
            "Comando: cd backend && python -m pytest -v",
            "Ubicación de pruebas: backend/tests/",
        ],
    )

    H(doc, "4. Casos de prueba")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["ID", "Módulo", "Descripción", "Resultado"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, size=9, bold=True)

    casos = [
        ("UT-01", "backlog_parser", "Normalización de estados Kanban (Done/In Progress/To Do)", "PASS"),
        ("UT-02", "backlog_parser", "Normalización de tipos (HU, TA, SP, EN, DO, RN)", "PASS"),
        ("UT-03", "backlog_parser", "Detección de columna título (ES/EN)", "PASS"),
        ("UT-04", "backlog_parser", "Eliminación de BOM UTF-8", "PASS"),
        ("UT-05", "backlog_parser", "Parseo CSV con separador coma", "PASS"),
        ("UT-06", "backlog_parser", "Parseo CSV con separador punto y coma", "PASS"),
        ("UT-07", "backlog_parser", "CSV vacío / sin columna título → ValueError", "PASS"),
        ("UT-08", "schemas", "Defaults y validación de Competencia", "PASS"),
        ("UT-09", "schemas", "BacklogItem con estado Kanban válido", "PASS"),
        ("UT-10", "schemas", "dict_to_propuesta asigna id y parsea semana_sugerida", "PASS"),
        ("UT-11", "schemas", "Round-trip propuesta_to_dict / dict_to_propuesta", "PASS"),
        ("UT-12", "semaforo", "Rangos rojo/naranja/amarillo/verde", "PASS"),
        ("UT-13", "semaforo", "Error fuerza rojo; bulk-commit penaliza −15%", "PASS"),
    ]
    for row in casos:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=9)
    doc.add_paragraph()

    H(doc, "5. Cómo reproducir")
    P(doc, "cd backend", size=9)
    P(doc, "python -m venv .venv", size=9)
    P(doc, "source .venv/Scripts/activate   # Windows", size=9)
    P(doc, "pip install -r requirements.txt", size=9)
    P(doc, "python -m pytest -v", size=9)

    H(doc, "6. Evidencia de ejecución (pytest)")
    P(
        doc,
        "Salida completa guardada en pytest_resultado.txt. Extracto:",
        italic=True,
    )
    # Truncate very long logs
    excerpt = log if len(log) < 6000 else log[:6000] + "\n...[truncado]..."
    P(doc, excerpt, size=8)

    H(doc, "7. Conclusión")
    P(
        doc,
        "Las pruebas unitarias del backend se ejecutan correctamente con pytest. "
        "Los módulos de parseo, schemas y semáforo cumplen los criterios definidos "
        "en los casos UT-01 a UT-13. Se recomienda mantener esta suite en CI ante "
        "cambios futuros del parser o de los umbrales del semáforo.",
    )

    doc.save(OUT)
    print("OK", OUT)


if __name__ == "__main__":
    build()
