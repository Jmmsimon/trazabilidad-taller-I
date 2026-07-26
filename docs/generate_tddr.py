# -*- coding: utf-8 -*-
"""Genera el Technical Design & Development Report (TDDR) en Word."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent / "TDDR_Trazabilidad_Academica.docx"


def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, italic=False, size=11, align="justify", space_after=6):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("[REVISAR / COMPLETAR] " + text)
    set_run_font(run, size=10, italic=True, color=RGBColor(0xB4, 0x5F, 0x06))
    p.paragraph_format.space_after = Pt(8)


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        set_run_font(run, size=11)


def shade_header(cell):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "1E3A5F")
    shading.set(qn("w:val"), "clear")
    cell._tePr = cell._tc.get_or_add_tcPr()
    cell._tc.get_or_add_tcPr().append(shading)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            set_run_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))
        shade_header(hdr[i])
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)

    # ═══════════════ PORTADA ═══════════════
    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, "UNIVERSIDAD PRIVADA ANTENOR ORREGO", bold=True, size=14, align="center")
    add_para(doc, "Facultad de Ingeniería", bold=True, size=12, align="center")
    add_para(doc, "Escuela Profesional de Ingeniería de Computación y Sistemas", size=11, align="center")
    doc.add_paragraph()
    add_para(
        doc,
        "TECHNICAL DESIGN & DEVELOPMENT REPORT (TDDR)",
        bold=True,
        size=16,
        align="center",
    )
    add_para(
        doc,
        "Plataforma de Trazabilidad Académica con Agentes de IA Generativa\n"
        "para Gestión de Proyectos Universitarios de Software",
        bold=True,
        size=13,
        align="center",
    )
    doc.add_paragraph()
    add_para(doc, "Informe Técnico de Diseño, Desarrollo y Validación", italic=True, size=11, align="center")
    doc.add_paragraph()
    add_para(doc, "Autor: Jean Marcos Meneses Simón", size=11, align="center")
    add_para(doc, "Asesor / Docente revisor: Ing. Walter Cueva", size=11, align="center")
    add_para(doc, "Fecha: Julio 2026", size=11, align="center")
    add_note(doc, "Completar nombre exacto de la universidad/facultad si difiere, y datos del asesor.")

    doc.add_page_break()

    # ═══════════════ SECCIÓN 0 ═══════════════
    add_heading_custom(doc, "SECCIÓN 0 — Portada y metadatos del proyecto", 1)

    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["Título del proyecto", "Plataforma de Trazabilidad Académica con Agentes de IA Generativa"],
            ["Tipo de solución", "Híbrida: Web Application + IA Generativa (multi-agente)"],
            ["Dominio de aplicación", "Educación superior / Ingeniería de software / Evaluación académica"],
            [
                "Palabras clave (ACM/IEEE)",
                "Academic Traceability; Multi-Agent Systems; Generative AI; "
                "Software Engineering Education; Learning Analytics; "
                "Human-in-the-Loop; Kanban; Source Code Analysis",
            ],
            ["Repositorio del código", "https://github.com/Jmmsimon/trazabilidad-taller-I"],
            ["Dataset", "No aplica dataset público externo. Datos de proyecto: commits GitHub públicos, "
             "backlogs CSV/Notion del estudiante, estados Firestore del proyecto académico."],
            ["Stack principal", "Next.js 16 · React 19 · FastAPI · LangGraph · Firebase Firestore · Gemini/Claude"],
            ["Versión documentada", "Commit ddf1890 (main) — Julio 2026"],
        ],
    )

    # ═══════════════ SECCIÓN 1 ═══════════════
    add_heading_custom(doc, "SECCIÓN 1 — Problema y motivación técnica", 1)

    add_heading_custom(doc, "1.1 Descripción del problema real", 2)
    add_para(
        doc,
        "En la formación universitaria de ingeniería de software, los docentes deben evaluar no solo "
        "entregables finales (documentos o demos), sino el avance real del código a lo largo del ciclo "
        "académico. En la práctica, esta evaluación se basa frecuentemente en reportes auto-declarados "
        "del estudiante (backlogs en Notion/CSV, presentaciones, checklists), sin una verificación "
        "sistemática contra el repositorio Git. Esto genera asimetría de información: el docente no "
        "puede auditar, a escala y con consistencia, si lo marcado como “Done” existe realmente en el código.",
    )
    add_para(
        doc,
        "El problema se agrava en cursos con múltiples proyectos concurrentes (talleres, tesis, "
        "capstone), donde la carga de revisión manual de commits, ramas, archivos y criterios de "
        "aceptación supera el tiempo disponible del docente. Como consecuencia aparecen: (i) "
        "evaluación subjetiva, (ii) desviaciones no detectadas (ítems “Done” sin evidencia), "
        "(iii) baja trazabilidad entre hitos académicos y commits, y (iv) feedback tardío al estudiante.",
    )
    add_note(
        doc,
        "Agregar evidencia cuantitativa local: p.ej. #proyectos por docente, horas de revisión, "
        "porcentaje de ítems Done sin evidencia en una muestra piloto, o citas de estudios sobre "
        "evaluación de proyectos de software en educación (ACM/IEEE).",
    )

    add_heading_custom(doc, "1.2 Brecha tecnológica identificada", 2)
    add_para(
        doc,
        "Existen herramientas de gestión (Jira, Notion, GitHub Projects) y de analítica de código "
        "(SonarQube, GitHub Insights), pero no una solución académica integrada que: (1) co-cree "
        "roadmap/backlog con agentes de IA bajo supervisión docente; (2) lea el repositorio en "
        "profundidad (árbol, snippets, diffs); (3) alinee commits ↔ hitos ↔ ítems Kanban de forma "
        "explícita; y (4) ofrezca al docente un semáforo de auditoría CSV/backlog vs código sin "
        "reescribir el estado del estudiante de manera opaca.",
    )
    add_para(
        doc,
        "La brecha es por tanto de trazabilidad académica automatizada con agentes generativos "
        "sincerados (evidence-based), human-in-the-loop y una sola fuente de verdad del proyecto "
        "compartida entre alumno y docente.",
    )

    add_heading_custom(doc, "1.3 Pregunta de investigación técnica", 2)
    add_para(
        doc,
        "¿En qué medida una plataforma web con orquestación multi-agente (LangGraph) y lectura "
        "profunda de repositorios GitHub permite mejorar la trazabilidad entre backlog/hitos "
        "académicos y evidencia de código (commits + archivos), reduciendo desviaciones de avance "
        "declarado versus avance real, en el contexto de proyectos universitarios de software?",
        italic=True,
    )

    add_heading_custom(doc, "1.4 Objetivo general y objetivos específicos", 2)
    add_para(doc, "Objetivo general", bold=True)
    add_para(
        doc,
        "Diseñar, implementar y validar una plataforma de trazabilidad académica basada en agentes "
        "de IA generativa que vincule backlog, hitos y repositorio Git, apoyando la evaluación "
        "docente y el autocontrol del estudiante.",
    )
    add_para(doc, "Objetivos específicos", bold=True)
    add_bullets(
        doc,
        [
            "OE1. Especificar e implementar un flujo Discovery multi-agente (drafter–validator–PO) para co-crear propuesta técnica, hitos y backlog Scrum.",
            "OE2. Implementar un pipeline Tracking (DevOps → Competency → Analyst → Reporter) con lectura profunda de GitHub y mapeo commits↔hitos↔Kanban.",
            "OE3. Desarrollar dashboards diferenciados (estudiante/docente/admin) con autenticación por roles (Firebase) y analítica de integridad/competencias.",
            "OE4. Implementar auditoría docente de backlog (CSV/Notion) vs código como reporte (semáforo), sin alterar indebidamente el Kanban del alumno.",
            "OE5. Evaluar la solución mediante escenarios de prueba funcionales, coherencia del auto-Kanban y usabilidad percibida en el flujo académico.",
        ],
    )

    add_heading_custom(doc, "1.5 Alcance y limitaciones declaradas", 2)
    add_para(doc, "Incluye:", bold=True)
    add_bullets(
        doc,
        [
            "Aplicación web (Next.js) + API (FastAPI) + agentes LangGraph.",
            "Roles: estudiante, docente, administrador.",
            "Integración GitHub (API pública): commits, árbol, snippets.",
            "Auto-actualización Kanban desde tracking compartido.",
            "Auditoría CSV/Notion → semáforo (solo reporte).",
            "Exportación PDF/JSON de reportes.",
        ],
    )
    add_para(doc, "Excluye / limitaciones:", bold=True)
    add_bullets(
        doc,
        [
            "Repositorios privados sin token/configuración avanzada (limitado a repos públicos o accesibles por la API configurada).",
            "No es un LMS completo (no reemplaza Moodle/Canvas).",
            "No garantiza detección de plagio formal (solo señales de integridad: bulk-commit, mensajes no constructivos).",
            "La calidad del mapeo semántico depende del LLM (Gemini/Claude) y de la calidad del backlog/CSV.",
            "Evaluación experimental a escala masiva (cientos de cursos) queda como trabajo futuro.",
        ],
    )

    add_heading_custom(doc, "1.6 Contribución técnica principal", 2)
    add_para(
        doc,
        "La contribución principal es una arquitectura híbrida de trazabilidad académica que combina: "
        "(a) orquestación multi-agente determinística (LangGraph) con estados tipados; "
        "(b) lectura profunda del repositorio para evidencia; "
        "(c) mapeo sincero commits↔hitos↔Kanban compartido alumno–docente; y "
        "(d) auditoría externa CSV/Notion como reporte docente con semáforo, desacoplada del Kanban. "
        "Esta combinación aplicada al dominio de evaluación de proyectos universitarios no aparece "
        "integrada de forma equivalente en herramientas genéricas de project management o code analytics.",
    )

    # ═══════════════ SECCIÓN 2 ═══════════════
    add_heading_custom(doc, "SECCIÓN 2 — Revisión de literatura técnica", 1)

    add_heading_custom(doc, "2.1 Marco conceptual técnico", 2)
    add_bullets(
        doc,
        [
            "Trazabilidad de requisitos/software: vínculo entre artefactos (requisitos, tareas, código, evidencias).",
            "Sistemas multi-agente y grafos de estado (LangGraph): flujos cíclicos con human-in-the-loop.",
            "IA generativa para análisis de código y mensajes de commit (LLM-as-judge con restricciones).",
            "Learning analytics / academic analytics: métricas de avance y competencias.",
            "Scrum/Kanban educativo: backlog, sprints, estados backlog|todo|in_progress|done.",
            "Human-in-the-loop: aprobación docente de roadmap, validación de hitos y revisión de HU.",
        ],
    )
    add_note(doc, "Expandir cada concepto con 1–2 citas peer-reviewed (IEEE/ACM).")

    add_heading_custom(doc, "2.2 Estado del arte de soluciones similares", 2)
    add_para(
        doc,
        "Tabla resumen (plantilla). Completar/ajustar referencias reales con DOI antes de entrega.",
    )
    add_table(
        doc,
        ["Ref", "Año", "Tipo", "Técnica/Tecnología", "Contexto", "Métrica", "Limitación"],
        [
            ["[1]", "2022", "LMS + analytics", "Dashboards", "Educación superior", "Engagement", "Sin vínculo a código Git"],
            ["[2]", "2021", "Git education", "GitHub Classroom", "Cursos de prog.", "Entrega a tiempo", "Evaluación superficial de commits"],
            ["[3]", "2023", "Code review AI", "LLM", "Software eng.", "Precisión review", "No Kanban académico"],
            ["[4]", "2020", "Learning analytics", "Process mining", "MOOC", "Dropout pred.", "No proyectos software"],
            ["[5]", "2024", "Multi-agent LLM", "AutoGen/LangGraph", "Agentes genéricos", "Task success", "Poco dominio educativo"],
            ["[6]", "2022", "Requirements tracing", "IR + ML", "Industria", "F-measure", "No UI docente universitaria"],
            ["[7]", "2023", "SE education", "Capstone platforms", "Universidades", "Satisfacción", "Manual audit of repos"],
            ["[8]", "2021", "DevOps education", "CI pipelines", "Cursos", "Build pass", "Sin mapeo a hitos"],
            ["[9]", "2024", "LLM code analysis", "GPT/Claude", "OSS", "Bug detect.", "Alucinaciones"],
            ["[10]", "2022", "Kanban education", "Agile tools", "Aulas", "Throughput", "Auto-declarado"],
            ["[11]", "2023", "Competence frameworks", "Rubrics", "ABET/IEEE", "Coverage", "Sin evidencia Git"],
            ["[12]", "2020", "Firestoreability tools", "Req→code", "Industria", "Link density", "Costoso/setup"],
            ["[13]", "2024", "AI tutors", "RAG agents", "Educación", "Learning gain", "No audit backlog"],
            ["[14]", "2021", "Firebase apps", "BaaS", "Web edu", "Time-to-MVP", "Vendor lock-in"],
            ["[15]", "2025", "GenAI SE", "Agents", "Dev tools", "Productivity", "Integridad académica débil"],
        ],
    )
    add_note(doc, "Reemplazar [1]–[15] por papers reales Q1/Q2 con DOI. Ampliar a ≥15 con tus lecturas.")

    add_heading_custom(doc, "2.3 Análisis comparativo de gaps", 2)
    add_table(
        doc,
        ["Capacidad", "LMS", "GitHub Classroom", "Jira/Notion", "Code AI tools", "Esta solución"],
        [
            ["Co-creación backlog con IA", "No", "No", "Parcial", "Parcial", "Sí (Discovery)"],
            ["Lectura profunda repo", "No", "Básica", "No", "Sí", "Sí (deep)"],
            ["Commits ↔ hitos académicos", "No", "No", "No", "No", "Sí"],
            ["Auto-Kanban evidence-based", "No", "No", "Manual", "No", "Sí"],
            ["Auditoría CSV vs código", "No", "No", "No", "No", "Sí (semáforo)"],
            ["HITL docente", "Sí", "Parcial", "Sí", "No", "Sí"],
            ["Roles académicos", "Sí", "Parcial", "No", "No", "Sí"],
        ],
    )
    add_para(
        doc,
        "Posicionamiento: la solución se sitúa en la intersección de learning analytics, "
        "ingeniería de software educativa y agentes generativos, cubriendo el gap de verificación "
        "evidenciada del avance declarado.",
    )

    add_heading_custom(doc, "2.4 Justificación de la elección tecnológica", 2)
    add_para(
        doc,
        "La selección se basó en determinismo de flujos multi-agente, tipado, velocidad de iteración "
        "académica y capacidad de persistir estados complejos (propuesta, backlog, tracking). "
        "Detalle expandido en docs/benchmark_tecnologico.md del repositorio.",
    )
    add_table(
        doc,
        ["Capa", "Elegida", "Descartada", "Motivo principal"],
        [
            ["Orquestación agentes", "LangGraph", "AutoGen / LangChain lineal", "Grafos cíclicos + estado explícito + HITL"],
            ["Backend", "FastAPI", "Django / Express", "Async nativo, Pydantic, BackgroundTasks para LLM"],
            ["Frontend", "Next.js 16 + React 19", "Vue / SPA CRA", "App Router, SSR/CSR híbrido, ecosistema TS"],
            ["Persistencia", "Firestore", "PostgreSQL (fase actual)", "Documentos flexibles para backlog/tracking"],
            ["Auth", "Firebase Auth", "Auth propia completa", "Roles + Google/email + cookies de sesión"],
            ["LLM", "Gemini / Claude vía cliente unificado", "Un solo proveedor", "Resiliencia y costos"],
        ],
    )

    # ═══════════════ SECCIÓN 3 ═══════════════
    add_heading_custom(doc, "SECCIÓN 3 — Diseño de la solución tecnológica", 1)

    add_heading_custom(doc, "3.1 Visión general de la arquitectura", 2)
    add_para(
        doc,
        "Arquitectura en capas: (1) Cliente Web Next.js (dashboards estudiante/docente/admin); "
        "(2) API Routes Next como BFF/proxy; (3) Backend FastAPI con grafos LangGraph "
        "(Discovery, Tracking, BacklogAudit); (4) Firebase Auth + Firestore; (5) GitHub API "
        "como fuente de evidencia de código.",
    )
    add_para(doc, "Componentes y responsabilidades:", bold=True)
    add_bullets(
        doc,
        [
            "Frontend estudiante: Fase A–D (idea → discovery → aprobación → tracking/Kanban/analítica).",
            "Frontend docente: aprobación roadmap, validación hitos, revisión HU, Control Git, Analítica, Auditoría CSV.",
            "API FastAPI: /proyectos/*, /profesor/*, /admin/*, /github/webhook, tracking y backlog-audit.",
            "Discovery graph: drafter → validator → po (co-creación).",
            "Tracking graph: devops(deep) → competency → analyst → reporter + auto-Kanban.",
            "Backlog audit graph: parser CSV/Notion → github deep → semantic match → semáforo (solo reporte).",
        ],
    )
    add_note(
        doc,
        "Insertar Figura 1: diagrama de arquitectura (exportar desde draw.io/Mermaid a SVG/PDF vectorial). "
        "Sugerencia Mermaid en Anexo A.",
    )

    add_heading_custom(doc, "3.2 Especificación de requerimientos técnicos", 2)
    add_heading_custom(doc, "3.2.1 Requerimientos funcionales", 3)
    add_table(
        doc,
        ["ID", "Requerimiento", "Prioridad", "Objetivo"],
        [
            ["RF01", "Registrar/iniciar sesión por rol (estudiante, docente, admin)", "Alta", "OE3"],
            ["RF02", "Co-crear propuesta, hitos y backlog Scrum con agentes Discovery", "Alta", "OE1"],
            ["RF03", "Docente aprueba/rechaza roadmap con comentario", "Alta", "OE3"],
            ["RF04", "Estudiante vincula repo_url y demo_url", "Alta", "OE2"],
            ["RF05", "Ejecutar tracking con lectura profunda GitHub", "Alta", "OE2"],
            ["RF06", "Mapear commits a hitos e ítems y actualizar Kanban", "Alta", "OE2"],
            ["RF07", "Mostrar analítica (score, competencias, commits↔hitos)", "Alta", "OE3"],
            ["RF08", "Docente valida hitos y revisa historias de usuario", "Alta", "OE3"],
            ["RF09", "Auditoría CSV/Notion vs código con semáforo (solo reporte)", "Alta", "OE4"],
            ["RF10", "Exportar reporte PDF/JSON", "Media", "OE5"],
            ["RF11", "Webhook GitHub para re-lanzar tracking", "Media", "OE2"],
            ["RF12", "Admin invita usuarios y gestiona roles/estado", "Media", "OE3"],
        ],
    )

    add_heading_custom(doc, "3.2.2 Requerimientos no funcionales", 3)
    add_bullets(
        doc,
        [
            "Rendimiento: respuestas de API de consulta < 2 s; jobs LLM en BackgroundTasks sin bloquear UI.",
            "Escalabilidad: diseño async ASGI; proyección de concurrencia documentada en benchmark interno.",
            "Seguridad: auth Firebase + cookie httpOnly de sesión; roles; CORS restringido a localhost frontend en dev.",
            "Usabilidad: dashboards por rol; botones Analizar unificados; edición de hitos bajo modo Editar.",
            "Mantenibilidad: tipado TS/Pydantic; módulos por grafo; proxies Next hacia backend.",
            "Reproducibilidad: repositorio GitHub + requirements.txt + package.json.",
        ],
    )

    add_heading_custom(doc, "3.3 Modelado del sistema", 2)
    add_para(doc, "Casos de uso principales (actores: Estudiante, Docente, Admin, Sistema/Agentes):", bold=True)
    add_bullets(
        doc,
        [
            "UC1 Iniciar proyecto y co-crear backlog",
            "UC2 Aprobar/rechazar roadmap",
            "UC3 Analizar repositorio (tracking)",
            "UC4 Actualizar/consultar Kanban",
            "UC5 Validar hito / observar tareas",
            "UC6 Auditar CSV vs código",
            "UC7 Exportar reporte",
        ],
    )
    add_para(
        doc,
        "Modelo de datos (Firestore, documentos principales): usuarios{uid, rol, email, nombre, estado}; "
        "proyectos{proyectoId, alumnoId, status, propuesta/hitos, backlog_scrum, repo_url, demo_url, "
        "tracking, tracking_history, backlog_audit, tracking_status, …}.",
    )
    add_para(
        doc,
        "Flujo crítico Tracking (secuencia): UI → POST /tracking/iniciar → BackgroundTask → "
        "devops_node(deep) → competency_node(mapeo_tareas) → analyst_node → reporter_node → "
        "AUTO-KANBAN → Firestore → UI poll /tracking/status.",
    )
    add_note(doc, "Adjuntar: casos de uso UML, ER/documento, secuencia, ≥5 mockups (Anexo).")

    add_heading_custom(doc, "3.4 Stack tecnológico justificado", 2)
    add_table(
        doc,
        ["Capa", "Tecnología", "Versión", "Justificación", "Alternativa descartada"],
        [
            ["Frontend", "Next.js", "16.2.6", "App Router, proxies API, TS", "CRA / Remix"],
            ["UI", "React + Tailwind + Recharts", "19.x / 4.x / 3.x", "Dashboards y charts", "Mantine completo"],
            ["Backend", "FastAPI + Uvicorn", "actual en requirements", "Async + BackgroundTasks", "Django"],
            ["Agentes", "LangGraph", "1.x", "Grafos con estado", "AutoGen"],
            ["LLM", "Gemini / Claude", "vía langchain-*", "Análisis semántico", "Solo un vendor"],
            ["Auth/DB", "Firebase Auth + Firestore", "12/13 admin", "Roles + docs flexibles", "Postgres+JWT propio"],
            ["Validación", "Pydantic / Zod", "2.x / 4.x", "Contratos tipados", "Validación ad hoc"],
            ["PDF", "Reportlab", "-", "Reporte firmable SHA-256", "WeasyPrint"],
        ],
    )

    add_heading_custom(doc, "3.5 Decisiones de diseño críticas (ADR)", 2)
    add_para(doc, "ADR-01: Una sola verdad del tracking compartida alumno–docente", bold=True)
    add_para(
        doc,
        "Contexto: riesgo de dos analíticas divergentes. Decisión: mismo endpoint run_tracking_task. "
        "Alternativa: tracking_profesor separado. Consecuencia: el Analizar docente actualiza score/Kanban "
        "visible al alumno (coherente con trazabilidad).",
    )
    add_para(doc, "ADR-02: Auditoría CSV como solo-reporte", bold=True)
    add_para(
        doc,
        "Contexto: CSV externo puede desalinearse del backlog Scrum interno. Decisión: backlog_audit "
        "no escribe backlog_scrum (afecta_kanban_alumno=false). Alternativa: sync automático CSV→Kanban. "
        "Consecuencia: semáforo docente sin corromper Kanban del estudiante.",
    )
    add_para(doc, "ADR-03: Lectura profunda (deep) en tracking del estudiante", bold=True)
    add_para(
        doc,
        "Contexto: shallow commits no bastan para sinceridad. Decisión: fetch_github_data(deep=True) "
        "también en tracking. Alternativa: deep solo en auditoría docente. Consecuencia: mayor latencia/tokens, "
        "mejor evidencia para auto-Kanban.",
    )
    add_para(doc, "ADR-04: Deshabilitar middleware Next 16 en desarrollo", bold=True)
    add_para(
        doc,
        "Contexto: bug Next 16/Turbopack: middleware/proxy provoca 404 en rutas válidas. "
        "Decisión: middleware.ts.disabled. Alternativa: migrar a proxy.ts inestable. "
        "Consecuencia: auth de rutas dashboard se maneja en cliente (AuthGuard) en esta fase.",
    )

    add_heading_custom(doc, "3.6 Modelo de seguridad y privacidad", 2)
    add_bullets(
        doc,
        [
            "Autenticación: Firebase Auth (Google/email); cookie httpOnly token para sesión web.",
            "Autorización: roles en Firestore (estudiante/profesor/administrador); guards de UI y endpoints por dominio.",
            "Datos: proyectos ligados a alumnoId; docente lista proyectos de su contexto.",
            "Privacidad: correos y nombres académicos; no almacenar secretos de GitHub en cliente; .env fuera de git.",
            "Integridad de reportes: PDF con hash SHA-256 (diseño HU-007).",
            "CORS: origen frontend local en desarrollo.",
        ],
    )
    add_note(doc, "Si hay estudio con usuarios, incluir consentimiento (Anexo G) y minimización de PII.")

    # ═══════════════ SECCIÓN 4 ═══════════════
    add_heading_custom(doc, "SECCIÓN 4 — Desarrollo e implementación", 1)

    add_heading_custom(doc, "4.1 Metodología de desarrollo", 2)
    add_para(
        doc,
        "Se aplicó un enfoque híbrido de prototipado evolutivo + prácticas Scrum/Kanban: "
        "sprints semanales alineados a hitos académicos del curso, con entregables demos "
        "(Discovery, Tracking, Auditoría, UX). Design Science Research (DSR) guía el artefacto "
        "(plataforma) y su evaluación en un entorno real de taller/tesis.",
    )
    add_table(
        doc,
        ["Fase / Sprint", "Entregable", "Estado"],
        [
            ["S1–S2", "Auth, roles, landing, scaffold Next+FastAPI", "Completado"],
            ["S3–S4", "Discovery multi-agente + aprobación docente", "Completado"],
            ["S5–S6", "Tracking + Kanban + competencias", "Completado"],
            ["S7", "Agentes sincerados deep + analítica commits↔hitos", "Completado"],
            ["S8", "Auditoría CSV robusta + proxy Next + UX docente", "Completado"],
            ["S9+", "Evaluación formal, carga, paper-ready metrics", "En curso / completar"],
        ],
    )

    add_heading_custom(doc, "4.2 Módulos implementados", 2)
    add_para(doc, "Módulo Discovery (backend/discovery_graph.py)", bold=True)
    add_para(
        doc,
        "Función: transformar idea del alumno en propuesta técnica, hitos y backlog Scrum mediante "
        "agentes drafter/validator/PO. Persistencia en Firestore (_propuesta_raw, backlog_scrum).",
    )
    add_para(doc, "Módulo Tracking (backend/tracking_graph.py + main.run_tracking_task)", bold=True)
    add_para(
        doc,
        "Función: evidenciar avance real. devops_node lee GitHub en deep; competency_node produce "
        "commits_analizados, mapeo_tareas y competencias; analyst/reporter calculan score y resumen; "
        "AUTO-KANBAN aplica estados backlog|todo|in_progress|done con boost por entregables configurados.",
    )
    add_para(doc, "Fragmento representativo (concepto):", italic=True)
    add_para(
        doc,
        'github_data = fetch_github_data(repo_url, deep=True)\n'
        "# → tree_files, snippets, commits_with_files\n"
        'mapeo_tareas = competency_node(...)\n'
        "# → AUTO-KANBAN actualiza backlog_scrum.estado",
        size=9,
    )

    add_para(doc, "Módulo Auditoría Backlog (backlog_parser + backlog_audit_graph)", bold=True)
    add_para(
        doc,
        "Función: parsear CSV/Notion flexible (BOM, ;/, estados Kanban), leer repo deep, match semántico "
        "ítem↔código, semáforo. Guarda backlog_audit con modo solo_reporte.",
    )

    add_para(doc, "Módulo Frontend estudiante/docente", bold=True)
    add_para(
        doc,
        "Estudiante: CommitsHitosPanel, Kanban, botón Analizar, edición de hitos bajo modo Editar. "
        "Docente: Analítica con score/commits/competencias, poll real de agentes, BacklogAuditor vía "
        "/api/profesor/.../backlog-audit.",
    )

    add_heading_custom(doc, "4.3 Gestión de datos", 2)
    add_heading_custom(doc, "4.3.1 Fuentes de datos", 3)
    add_bullets(
        doc,
        [
            "GitHub REST API: commits, contents, git trees, commit files (repositorios del estudiante).",
            "Firestore: estado del proyecto académico, tracking, auditoría.",
            "CSV/Notion del backlog declarado por el alumno (auditoría docente).",
            "Entradas de usuario: idea, stack, URLs, correcciones de hitos.",
        ],
    )
    add_heading_custom(doc, "4.3.2 Preprocesamiento", 3)
    add_para(
        doc,
        "Normalización de estados CSV; truncado de snippets (líneas/archivos); selección de extensiones "
        "de código; detección bulk-commit; serialización JSON de hitos/backlog para prompts LLM; "
        "limpieza de respuesta JSON del modelo (clean_json_response).",
    )
    add_heading_custom(doc, "4.3.3 Partición de datos", 3)
    add_para(
        doc,
        "No aplica train/test clásico de ML supervisado. La evaluación es de sistema (casos de prueba "
        "funcionales y escenarios de repositorio real). Si se entrena/afinara un clasificador propio "
        "en el futuro, se documentará aquí.",
    )

    add_heading_custom(doc, "4.4 Entorno de desarrollo y producción", 2)
    add_bullets(
        doc,
        [
            "Dev: Windows, Python 3.12 + venv, Node.js, Next dev :3000, Uvicorn :8000.",
            "Dependencias: backend/requirements.txt; package.json (Next 16.2.6, React 19.2.4).",
            "Secrets: .env / .env.local (Firebase, LLM keys) — no versionados.",
            "Prod (ejemplo del alumno): frontend en Vercel; backend según despliegue institucional/cloud.",
        ],
    )

    add_heading_custom(doc, "4.5 Control de versiones y trazabilidad", 2)
    add_para(
        doc,
        "Repositorio: https://github.com/Jmmsimon/trazabilidad-taller-I — rama main. "
        "Al momento de este informe: ~60 commits. Estrategia práctica feature-on-main / commits "
        "atómicos por capacidad (tracking, UI, auditoría). Release documentada: ddf1890 "
        "(agentes sincerados + analítica + CSV audit).",
    )

    # ═══════════════ SECCIÓN 5 ═══════════════
    add_heading_custom(doc, "SECCIÓN 5 — Evaluación y validación", 1)

    add_heading_custom(doc, "5.1 Estrategia de evaluación", 2)
    add_para(
        doc,
        "Evaluación mixta: (i) validación técnica funcional de flujos críticos; "
        "(ii) validación de coherencia del mapeo evidencia→Kanban/auditoría; "
        "(iii) evaluación de usabilidad con usuarios del taller (docente/estudiante). "
        "No se pretende benchmark de accuracy tipo ImageNet, sino trazabilidad evidenciable.",
    )

    add_heading_custom(doc, "5.2 Métricas de evaluación definidas", 2)
    add_table(
        doc,
        ["Métrica", "Definición", "Justificación"],
        [
            ["Score de integridad", "0–100 (analyst_node)", "Salud del avance y riesgos DevOps"],
            ["% correspondencia auditoría", "Promedio scores ítem CSV↔código", "Desviación declared vs real"],
            ["Precisión auto-Kanban (manual)", "% ítems cuyo estado IA coincide con juicio experto", "Sinceridad del agente"],
            ["Cobertura commits alineados", "% commits alineado=true con hito_ref/item_ids", "Trazabilidad commits↔hitos"],
            ["Latencia percibida análisis", "Tiempo iniciar→completed", "Viabilidad en aula"],
            ["SUS (usabilidad)", "Cuestionario 10 ítems", "Aceptación docente/estudiante"],
            ["Tasa error funcional", "#fallos / #casos de prueba", "Calidad de implementación"],
            ["Desviaciones críticas CSV", "#Done sin evidencia", "Impacto docente"],
        ],
    )
    add_note(doc, "Completar con valores medidos reales de tus pruebas (media ± DE, N).")

    add_heading_custom(doc, "5.3 Diseño experimental", 2)
    add_bullets(
        doc,
        [
            "Ambiente: localhost (Next+FastAPI) + repo GitHub público de prueba del autor.",
            "Casos funcionales: login roles; discovery; aprobación; analizar; Kanban; auditoría CSV (columnas mínimas y completas); export PDF.",
            "Casos de sensibilidad: ítems GitHub con repo ya vinculado → estado done; CSV con menos columnas; commits no constructivos.",
            "Usuarios (piloto): docente del curso + estudiante autor (ampliar N si hay cohorte).",
        ],
    )

    add_heading_custom(doc, "5.4 Resultados obtenidos", 2)
    add_para(
        doc,
        "Resultados preliminares del escenario de demostración (proyecto taller del autor): "
        "tracking completado con score de integridad alto cuando existen commits descriptivos, "
        "CI pass y demo_url; 15 commits analizados en una corrida; auto-Kanban movió la mayoría "
        "de ítems con evidencia (p.ej. vinculación GitHub → done). La analítica docente carga "
        "score, % competencias, commits y resumen ejecutivo tras el mismo pipeline.",
    )
    add_table(
        doc,
        ["Escenario", "Resultado observado (preliminar)", "Evidencia"],
        [
            ["Analizar con repo+demo", "tracking_status=completed; score≈100 en demo controlada", "UI docente/estudiante"],
            ["Ítems GitHub en Kanban", "EN/HU de GitHub pasan a done con boost entregable", "kanban_updates"],
            ["CSV con columnas mínimas", "Parser acepta titulo+estado (;/, BOM)", "Tests parser"],
            ["Auditoría CSV", "Semáforo + desviaciones; Kanban alumno intacto", "backlog_audit.modo"],
            ["Middleware Next 16", "404 resuelto al deshabilitar middleware", "Logs Next"],
        ],
    )
    add_note(doc, "Reemplazar ≈ y preliminares por tablas con N corridas, DE y capturas en Anexo F.")

    add_heading_custom(doc, "5.5 Comparación con línea base", 2)
    add_table(
        doc,
        ["Método", "Trazabilidad commits↔hitos", "Auto-Kanban evidenciado", "Auditoría CSV vs código", "HITL docente"],
        [
            ["Revisión manual docente", "Baja", "No", "Manual", "Sí"],
            ["GitHub Insights + Notion", "Media", "No", "No integrada", "Parcial"],
            ["Solo LLM chat sin grafo", "Variable", "Ad hoc", "Ad hoc", "Débil"],
            ["Método propuesto", "Alta (diseño)", "Sí", "Sí (semáforo)", "Sí"],
        ],
    )

    add_heading_custom(doc, "5.6 Análisis estadístico", 2)
    add_note(
        doc,
        "Cuando tengas N≥2 grupos o ≥10 proyectos: aplicar Wilcoxon/t-test sobre precisión de mapeo "
        "o SUS pre/post; reportar p y effect size. Mientras tanto declarar evaluación exploratoria.",
    )

    add_heading_custom(doc, "5.7 Discusión de resultados", 2)
    add_para(
        doc,
        "Los resultados preliminares sugieren que la lectura profunda + prompt sincerado reduce el "
        "sesgo a dejar ítems críticos (p.ej. GitHub) en in_progress cuando el entregable ya existe. "
        "La separación ADR-02 evita que un CSV ruidoso corrompa el Kanban. Limitaciones: dependencia "
        "de LLM, rate limits GitHub, y repos privados. Casos de falla esperables: backlogs ambiguos, "
        "commits vacíos de significado, o CSV sin columna título.",
    )

    # ═══════════════ SECCIÓN 6 ═══════════════
    add_heading_custom(doc, "SECCIÓN 6 — Discusión integradora", 1)
    add_heading_custom(doc, "6.1 Criterio de respuesta a la pregunta de investigación", 2)
    add_para(
        doc,
        "La plataforma permite, en el contexto evaluado, vincular de forma automatizada y "
        "supervisable el backlog/hitos con evidencia de repositorio, exponiendo score, commits "
        "alineados y un semáforo de auditoría. El grado exacto de mejora cuantitativa frente a "
        "línea base manual debe cerrarse con el estudio formal de la Sección 5.",
    )
    add_heading_custom(doc, "6.2 Contribuciones técnicas verificadas", 2)
    add_bullets(
        doc,
        [
            "Arquitectura multi-agente académica Discovery + Tracking + Audit.",
            "Auto-Kanban evidence-based compartido alumno–docente.",
            "Analítica commits↔hitos en ambos dashboards.",
            "Auditoría CSV flexible y solo-reporte.",
        ],
    )
    add_heading_custom(doc, "6.3 Limitaciones", 2)
    add_bullets(
        doc,
        [
            "Evaluación estadística formal aún parcial.",
            "Dependencia de APIs LLM y GitHub.",
            "Middleware Next deshabilitado (auth edge pendiente de rediseño).",
            "CSV Excel .xlsx no parseado nativamente (requiere export CSV UTF-8).",
        ],
    )
    add_heading_custom(doc, "6.4 Amenazas a la validez", 2)
    add_bullets(
        doc,
        [
            "Interna: sesgo del autor como evaluador del propio sistema.",
            "Externa: un dominio/curso; generalización limitada.",
            "Constructo: “integridad” operacionalizada vía heurísticas+LLM, no auditoría forense.",
            "Estadística: N pequeño en piloto.",
        ],
    )
    add_heading_custom(doc, "6.5 Trabajo futuro", 2)
    add_bullets(
        doc,
        [
            "Estudio controlado con cohorte ≥20 proyectos y métricas de precisión del mapeo.",
            "Soporte repos privados con GitHub App y webhooks firmados en producción.",
            "Reintroducir auth edge/proxy estable en Next 16 y pruebas de carga reales (k6/Locust).",
            "Exportadores de plantilla CSV y validador visual de columnas críticas/opcionales.",
        ],
    )

    # ═══════════════ SECCIÓN 7 ═══════════════
    add_heading_custom(doc, "SECCIÓN 7 — Conclusiones", 1)
    add_para(
        doc,
        "Se diseñó e implementó una plataforma web de trazabilidad académica con agentes de IA "
        "generativa que integra co-creación de backlog, análisis profundo de repositorios, "
        "actualización evidenciable del Kanban y auditoría docente CSV/Notion. El repositorio "
        "público permite reproducir la solución. Los resultados preliminares muestran viabilidad "
        "técnica y utilidad potencial para reducir la brecha entre avance declarado y avance real; "
        "la validación estadística ampliada consolidará el aporte para publicación/tesis.",
    )

    # ═══════════════ SECCIÓN 8 ═══════════════
    add_heading_custom(doc, "SECCIÓN 8 — Referencias (plantilla — completar con DOI reales)", 1)
    refs = [
        "[1] A. Author et al., “Learning analytics in higher education: A review,” IEEE Trans. Learn. Technol., 2022.",
        "[2] B. Author et al., “GitHub Classroom experiences,” ACM SIGCSE, 2021.",
        "[3] C. Author et al., “Large language models for code review,” ICSE, 2023.",
        "[4] D. Author et al., “Process mining for education,” Elsevier Computers & Education, 2020.",
        "[5] E. Author et al., “LangGraph/multi-agent orchestration,” arXiv/peer-reviewed venue, 2024.",
        "[6] F. Author et al., “Requirements-to-code traceability,” Requirements Eng. J., 2022.",
        "[7] G. Author et al., “Capstone project assessment,” IEEE EDUCON, 2023.",
        "[8] H. Author et al., “DevOps in SE education,” Computer Applications in Eng. Education, 2021.",
        "[9] I. Author et al., “LLM-based program analysis,” ASE, 2024.",
        "[10] J. Author et al., “Kanban in software education,” Journal of Systems and Software, 2022.",
        "[11] K. Author et al., “Competency frameworks in computing,” ACM TOCE, 2023.",
        "[12] L. Author et al., “Software traceability tools survey,” IEEE TSE, 2020.",
        "[13] M. Author et al., “Generative AI tutors,” CHI/LAK, 2024.",
        "[14] N. Author et al., “Serverless educational backends,” IEEE Cloud, 2021.",
        "[15] O. Author et al., “Human-in-the-loop ML systems,” ACM Computing Surveys, 2022.",
        "[16] P. Author et al., “Agile metrics for student teams,” IEEE Software, 2023.",
        "[17] Q. Author et al., “Academic integrity and AI,” Computers & Education: AI, 2024.",
        "[18] R. Author et al., “Firebase/NoSQL for rapid edu apps,” Softw. Pract. Exp., 2021.",
        "[19] S. Author et al., “FastAPI and async web services,” (sustituir por paper/libro peer-reviewed), 2022.",
        "[20] T. Author et al., “Evaluation methods for SE education tools,” ACM TOCE, 2025.",
    ]
    for r in refs:
        add_para(doc, r, size=10)
    add_note(doc, "Obligatorio: ≥20 referencias reales, ≥70% últimos 5 años, ≥60% Q1/Q2. Eliminar placeholders.")

    # ═══════════════ SECCIÓN 9 ═══════════════
    add_heading_custom(doc, "SECCIÓN 9 — Anexos técnicos", 1)
    add_para(doc, "Anexo A — Diagrama de arquitectura", bold=True)
    add_para(
        doc,
        "Insertar figura vectorial. Borrador textual del flujo:\n"
        "Estudiante/Docente UI (Next.js) → API Routes → FastAPI → LangGraph "
        "(Discovery | Tracking | BacklogAudit) → Firestore / GitHub API / LLM.",
    )
    add_para(doc, "Anexo B — Especificación de API (extracto)", bold=True)
    add_table(
        doc,
        ["Método", "Endpoint", "Descripción"],
        [
            ["POST", "/proyectos/iniciar", "Inicia Discovery"],
            ["GET", "/proyectos/{id}/status", "Estado discovery"],
            ["POST", "/proyectos/{id}/tracking/iniciar", "Inicia tracking sincerado"],
            ["GET", "/proyectos/{id}/tracking/status", "Poll tracking + backlog_scrum"],
            ["POST", "/proyectos/{id}/configuracion", "Guarda repo_url/demo_url"],
            ["POST", "/profesor/proyectos/{id}/aprobar-roadmap", "Aprueba roadmap"],
            ["POST", "/profesor/proyectos/{id}/backlog-audit", "Inicia auditoría CSV/Notion"],
            ["GET", "/profesor/proyectos/{id}/backlog-audit/status", "Estado/resultado auditoría"],
            ["GET", "/proyectos/{id}/reporte-pdf", "Descarga PDF"],
            ["POST", "/github/webhook", "Re-dispara tracking"],
        ],
    )
    add_para(doc, "Anexo C — Diccionario de datos (extracto proyectos)", bold=True)
    add_table(
        doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["proyectoId", "string", "ID del proyecto"],
            ["alumnoId", "string", "UID Firebase del estudiante"],
            ["status", "string", "processing|pending_approval|active|rejected|error"],
            ["backlog_scrum", "object", "Épicas e ítems Kanban"],
            ["tracking", "object", "Score, commits, competencias, kanban_updates"],
            ["backlog_audit", "object", "Semáforo y resultados (solo reporte)"],
            ["repo_url / demo_url", "string", "Entregables"],
        ],
    )
    add_para(doc, "Anexo D — Manual de instalación (resumen)", bold=True)
    add_para(
        doc,
        "1) Clonar repo. 2) Backend: cd backend && python -m venv .venv && activate && "
        "pip install -r requirements.txt && configurar .env && uvicorn main:app --reload --port 8000. "
        "3) Frontend: npm install && configurar .env.local && npm run dev. "
        "4) Abrir http://localhost:3000. Nota: middleware deshabilitado (middleware.ts.disabled).",
    )
    add_para(doc, "Anexo E — Dataset / acceso", bold=True)
    add_para(doc, "Repos públicos del estudiante + CSV de ejemplo docs/backlog.csv. Sin dataset ML externo.")
    add_para(doc, "Anexo F — Resultados completos de pruebas", bold=True)
    add_note(doc, "Pegar matrices de casos (OK/FAIL), capturas y corridas negativas.")
    add_para(doc, "Anexo G — Consentimiento informado", bold=True)
    add_note(doc, "Si aplica evaluación con terceros, adjuntar formato institucional.")

    # Checklist
    doc.add_page_break()
    add_heading_custom(doc, "Checklist de autoevaluación (rúbrica)", 1)
    add_table(
        doc,
        ["Criterio", "Estado actual del borrador", "Acción pendiente del autor"],
        [
            ["Problema con evidencia cuantitativa", "Cualitativo + placeholders", "Añadir cifras y fuentes primarias"],
            ["Gap tecnológico explícito", "Tabla comparativa incluida", "Citar papers reales en celdas"],
            ["Arquitectura documentada", "Descrita + ADR", "Insertar diagrama vectorial"],
            ["Stack justificado", "Tabla + benchmark interno", "Cruzar con refs"],
            ["Métricas apropiadas", "≥4 métricas definidas", "Medir y reportar valores"],
            ["Comparación estado del arte", "Tabla gap + baseline", "≥5 trabajos con números"],
            ["Análisis estadístico", "Declarado pendiente", "Correr test con N adecuado"],
            ["Reproducibilidad", "Repo público + instrucciones", "README instalación pulido"],
            ["Referencias Q1", "Plantilla 20 items", "Sustituir por DOI reales ≥60% Q1/Q2"],
            ["Limitaciones y amenazas", "Incluidas", "Refinar tras experimentos"],
        ],
    )

    add_para(
        doc,
        "Fin del TDDR (borrador técnico generado a partir del repositorio trazabilidad-taller-I). "
        "Revisar marcas [REVISAR / COMPLETAR] antes de entrega formal.",
        italic=True,
        align="center",
    )

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    build()
